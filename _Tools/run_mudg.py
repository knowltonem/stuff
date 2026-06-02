#!/usr/bin/env python3
"""
run_mudg.py — Parameterized Military Unique Deployment Guide runner.

Replaces all gen_dtr###_mudg.py one-off scripts.
Invoked by the newd prompt sequence after Product Category, CTN, and Document Type
have been selected.

Usage (direct):
    python3 _Tools/run_mudg.py
"""

import copy
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from runner_core import run_validate, append_draft_log, BASE, mudg_draft_path, mudg_example_path, normalize_month_abbr, get_git_username

from docx import Document


# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Helpers (normalize_month_abbr imported from runner_core)
# ---------------------------------------------------------------------------


def collapse_empty_paragraphs(doc: Document, max_consecutive: int = 1):
    """Remove excess consecutive empty paragraphs, keeping at most max_consecutive."""
    body = doc.element.body
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    paras = [el for el in body if el.tag == f'{{{ns}}}p']

    consecutive_empties = []
    removed = 0
    for p in paras:
        text = "".join(t.text or "" for t in p.findall(f".//{{{ns}}}t")).strip()
        # Also check for images/drawings — don't treat image paragraphs as empty
        has_drawing = len(p.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
        if not text and not has_drawing:
            consecutive_empties.append(p)
        else:
            if len(consecutive_empties) > max_consecutive:
                for excess in consecutive_empties[max_consecutive:]:
                    body.remove(excess)
                    removed += 1
            consecutive_empties = []
    # Handle trailing empties
    if len(consecutive_empties) > max_consecutive:
        for excess in consecutive_empties[max_consecutive:]:
            body.remove(excess)
            removed += 1

    if removed:
        print(f"Collapsed empty paragraphs: removed {removed} excess empties (max {max_consecutive} consecutive)")


def copy_cell_formatting(src_cell, dst_cell):
    """Copy paragraph and run formatting from source cell to destination cell."""
    for p in dst_cell.paragraphs:
        p._element.getparent().remove(p._element)
    for p in src_cell.paragraphs:
        dst_cell._element.append(copy.deepcopy(p._element))


def add_revision_rows(doc: Document, new_rows: list):
    """
    Normalize existing date cells, then append one or more revision history rows.
    new_rows: list of dicts with keys: version, date, change, editor
    """
    table = doc.tables[0]

    # Normalize existing dates
    for row in table.rows[1:]:
        date_cell = row.cells[1]
        old_text = date_cell.text.strip()
        new_text = normalize_month_abbr(old_text)
        if new_text != old_text:
            for p in date_cell.paragraphs:
                for run in p.runs:
                    run.text = normalize_month_abbr(run.text)

    template_row = table.rows[-1]

    for row_data in new_rows:
        values = [row_data["version"], row_data["date"], row_data["change"], row_data["editor"]]
        new_row = table.add_row()
        for i, (src_cell, dst_cell) in enumerate(zip(template_row.cells, new_row.cells)):
            copy_cell_formatting(src_cell, dst_cell)
            if dst_cell.paragraphs:
                runs = dst_cell.paragraphs[0].runs
                if runs:
                    runs[0].text = values[i]
                    for extra in runs[1:]:
                        extra._element.getparent().remove(extra._element)
                else:
                    dst_cell.paragraphs[0].text = values[i]
            else:
                dst_cell.text = values[i]
        print(f"Added row: {values}")


# ---------------------------------------------------------------------------
# Generation
# ---------------------------------------------------------------------------

def generate(cfg: dict):
    prod_cat = cfg["prod_cat"]
    ctn = cfg["ctn"]
    dtr_num = cfg["dtr_num"]
    out_path = cfg["out_path"]

    # DTR sources from Example when dtr_num matches the first known MUDG DTR for this CTN.
    # first_dtr_num MUST be set in cfg — there is no safe universal default because it is
    # CTN-specific (SBC/CTN2026003 starts at DTR005; a new CTN may start at DTR001).
    if "first_dtr_num" not in cfg:
        raise ValueError(
            "cfg['first_dtr_num'] is required — set it to the DTR number of the first "
            "MUDG draft for this CTN (e.g. 5 for SBC/CTN2026003, 1 for a new CTN)."
        )
    if dtr_num == cfg["first_dtr_num"]:
        source_path = mudg_example_path(prod_cat, ctn)
    else:
        source_path = cfg["source_path"]

    if not source_path.exists():
        raise FileNotFoundError(f"Source not found: {source_path}")

    doc = Document(str(source_path))
    add_revision_rows(doc, cfg["new_rows"])
    collapse_empty_paragraphs(doc, max_consecutive=1)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


# ---------------------------------------------------------------------------
# Profiles — first MUDG DTR seed only; subsequent DTRs are fully prompted
# ---------------------------------------------------------------------------

SBC_CTN2026003_MUDG_PROFILES = {
    5: {
        "dtr_num": 5,
        "first_dtr_num": 5,  # sources from Example (not a prior draft)
        # Dates below are intentionally frozen for this one-time DTR005 seed run.
        # DTR4 row backdates to May 2026 (the actual approval month); DTR5 row is Jun 2026.
        # If regenerating this profile at a later date, verify these dates are still correct.
        "new_rows": [
            {"version": "5.0", "date": "May 2026", "change": "Update for DTR 4", "editor": "GCT DP Collaboration"},
            {"version": "6.0", "date": "Jun 2026", "change": "Update for DTR 5", "editor": "GCT DP Collaboration"},
        ],
    },
    # DTR006+ have no pre-baked profiles — prompted at runtime.
}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def run(prod_cat: str, ctn: str, dtr_num: int) -> None:
    """Callable entry point — invoked by the newd prompt sequence or directly via main()."""
    profiles = {}
    if prod_cat == "SBC" and ctn == "CTN2026003":
        profiles = SBC_CTN2026003_MUDG_PROFILES

    if dtr_num in profiles:
        cfg = dict(profiles[dtr_num])
        print(f"\nLoaded profile for {prod_cat}/{ctn} DTR{dtr_num:03d}:")
        for row in cfg["new_rows"]:
            print(f"  + Rev row: {row['version']} | {row['date']} | {row['change']}")
        confirm = input("\nProceed with these settings? [y/N]: ").strip().lower()
        if confirm != "y":
            print("Aborted.")
            return
    else:
        print(f"\nNo profile found for {prod_cat}/{ctn} DTR{dtr_num:03d}. Manual input required.")
        print("Enter new revision row(s) to append. Leave version blank to finish.")
        new_rows = []
        while True:
            ver = input("  Revision version (e.g. 7.0) or blank to finish: ").strip()
            if not ver:
                break
            date = input("  Date (e.g. Jul 2026): ").strip()
            change = input("  Change description (e.g. Update for DTR 6): ").strip()
            editor = input("  Editor (default: GCT DP Collaboration): ").strip() or "GCT DP Collaboration"
            new_rows.append({"version": ver, "date": date, "change": change, "editor": editor})
        cfg = {
            "dtr_num": dtr_num,
            "first_dtr_num": int(input("First DTR number for this CTN's MUDG (usually 1, or 5 for SBC/CTN2026003): ").strip()),
            "new_rows": new_rows,
        }

    cfg["prod_cat"] = prod_cat
    cfg["ctn"] = ctn
    cfg["out_path"] = mudg_draft_path(prod_cat, ctn, dtr_num)

    # Always set source_path so generate() never hits a KeyError on cfg["source_path"].
    # When dtr_num == first_dtr_num the example path is used; otherwise the prior draft.
    if dtr_num == cfg.get("first_dtr_num"):
        cfg["source_path"] = mudg_example_path(prod_cat, ctn)
    else:
        cfg["source_path"] = mudg_draft_path(prod_cat, ctn, dtr_num - 1)

    generate(cfg)

    print("\nRunning validate_doc.py...")
    run_validate(cfg["out_path"])

    engineer = input("\nEngineer username (for Draft Log): ").strip() or get_git_username() or "unknown"
    append_draft_log(
        engineer=engineer,
        action="Generated",
        ctn=ctn,
        doc_type="MUDG",
        dtr=f"DTR{dtr_num:03d}",
        version="—",
        reason="Via run_mudg.py parameterized runner",
    )

    print("\nDone.")


def main():
    print("\n=== Military Unique Deployment Guide — Parameterized Runner ===\n")
    prod_cat = input("Product Category (e.g. SBC): ").strip()
    ctn = input("CTN (e.g. CTN2026003): ").strip()
    dtr_num = int(input("DTR Number (e.g. 5): ").strip())
    run(prod_cat, ctn, dtr_num)


if __name__ == "__main__":
    main()
