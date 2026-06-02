#!/usr/bin/env python3
"""
run_sysdesc.py — Parameterized System Description runner.

Replaces all gen_dtr###_sbc_sysdesc_v2.py one-off scripts.
Invoked by the newd prompt sequence after Product Category, CTN, and Document Type
have been selected.

Usage (direct):
    python3 _Tools/run_sysdesc.py
"""

import copy
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from runner_core import (
    qn, clone_element, get_para_text, make_body_para,
    run_validate, append_draft_log, BASE, get_git_username,
    sysdesc_dir, sysdesc_draft_path, sysdesc_initial_path, sysdesc_example_path,
    normalize_month_abbr,
)

from lxml import etree
from docx import Document


# ---------------------------------------------------------------------------
# Shared helpers (normalize_month_abbr imported from runner_core)
# ---------------------------------------------------------------------------

def add_page_break_before(para_el):
    pPr = para_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(para_el, qn("w:pPr"))
        para_el.insert(0, pPr)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pb = etree.SubElement(pPr, qn("w:pageBreakBefore"))
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pStyle.addnext(pb)
        else:
            pPr.insert(0, pb)


def remove_page_break_before(para_el):
    pPr = para_el.find(qn("w:pPr"))
    if pPr is not None:
        pb = pPr.find(qn("w:pageBreakBefore"))
        if pb is not None:
            pPr.remove(pb)


def set_all_runs_text(para_el, text: str):
    runs = para_el.findall(f".//{qn('w:r')}")
    for i, r in enumerate(runs):
        for t in r.findall(qn("w:t")):
            if i == 0:
                t.text = text
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            else:
                t.text = ""


def add_revision_row(rev_table, version_str: str, approval: str, dtr_num: int):
    """Normalize existing dates, then append a new revision history row."""
    for row in rev_table.rows[1:]:
        date_cell = row.cells[1]
        for p in date_cell.paragraphs:
            for run in p.runs:
                run.text = normalize_month_abbr(run.text)

    last_row = rev_table.rows[-1]
    new_row_el = clone_element(last_row._tr)
    cells = new_row_el.findall(qn("w:tc"))
    values = [version_str, approval, f"Update for DTR {dtr_num}", "GCT DP Collaboration"]
    for cell_el, val in zip(cells, values):
        for p in cell_el.findall(qn("w:p")):
            runs = p.findall(f".//{qn('w:r')}")
            for i, r in enumerate(runs):
                t_els = r.findall(qn("w:t"))
                if i == 0:
                    for t in t_els:
                        t.text = val
                        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        break
                else:
                    r.getparent().remove(r)
            break
    rev_table._tbl.append(new_row_el)


def make_hyperlink_para(doc, body_template, prefix_text: str, display_text: str, url: str):
    """Build a paragraph with plain prefix text + a Word hyperlink, styled from body_template."""
    r_id = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    new_para = copy.deepcopy(body_template)
    for r in new_para.findall(f".//{qn('w:r')}"):
        r.getparent().remove(r)
    for hl in new_para.findall(f".//{qn('w:hyperlink')}"):
        hl.getparent().remove(hl)

    orig_runs = body_template.findall(f".//{qn('w:r')}")
    rpr_el = None
    if orig_runs:
        rpr_el = orig_runs[0].find(qn("w:rPr"))

    # Plain text prefix run
    prefix_r = etree.SubElement(new_para, qn("w:r"))
    if rpr_el is not None:
        prefix_r.append(copy.deepcopy(rpr_el))
    prefix_t = etree.SubElement(prefix_r, qn("w:t"))
    prefix_t.text = prefix_text
    prefix_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Hyperlink element
    hl = etree.SubElement(new_para, qn("w:hyperlink"))
    hl.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id)
    hl.set(qn("w:history"), "1")
    hl_r = etree.SubElement(hl, qn("w:r"))
    hl_rpr = etree.SubElement(hl_r, qn("w:rPr"))
    hl_rStyle = etree.SubElement(hl_rpr, qn("w:rStyle"))
    hl_rStyle.set(qn("w:val"), "Hyperlink")
    hl_t = etree.SubElement(hl_r, qn("w:t"))
    hl_t.text = display_text
    hl_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    return new_para


# ---------------------------------------------------------------------------
# DTR001 — generate from INITIAL + Example
# ---------------------------------------------------------------------------

def generate_dtr001(cfg: dict):
    prod_cat = cfg["prod_cat"]
    ctn = cfg["ctn"]
    dtr_num = cfg["dtr_num"]
    new_ver = cfg["new_ver"]
    old_ver = cfg["old_ver"]
    out_path = cfg["out_path"]

    initial_path = sysdesc_initial_path(prod_cat, ctn)
    example_path = sysdesc_example_path(prod_cat, ctn)

    if not initial_path.exists():
        raise FileNotFoundError(f"INITIAL doc not found: {initial_path}")

    doc = Document(str(initial_path))
    example = Document(str(example_path))
    body = doc.element.body

    # 1. Revision history row
    add_revision_row(doc.tables[0], cfg["rev_version"], cfg["approval"], dtr_num)

    # 2. Find Management Description insertion point
    mgmt_idx = None
    for i, el in enumerate(body):
        if el.tag == qn("w:p") and "Management Description" in get_para_text(el):
            mgmt_idx = i
            break
    if mgmt_idx is None:
        raise ValueError("Could not find Management Description heading")

    # 3. Extract DTR section elements from Example
    ex_body = example.element.body
    dtr_heading_el = dtr_detail_heading_el = dtr_table_el = None
    dtr_body_paras = []
    in_dtr_section = False

    for el in list(ex_body):
        if el.tag == qn("w:p"):
            text = get_para_text(el)
            if "Desktop Review (DTR)" in text:
                dtr_heading_el = el
                in_dtr_section = True
                continue
            if "DTR Detailed Component Information" in text:
                dtr_detail_heading_el = el
                continue
            if "Management Description" in text:
                break
            if in_dtr_section and dtr_detail_heading_el is None:
                dtr_body_paras.append(el)
        elif el.tag == qn("w:tbl"):
            if dtr_detail_heading_el is not None and dtr_table_el is None:
                dtr_table_el = el

    if any(x is None for x in [dtr_heading_el, dtr_detail_heading_el, dtr_table_el]):
        raise ValueError("Could not extract all required elements from Example doc")

    # 4. Build new DTR section
    body_template = dtr_body_paras[0] if dtr_body_paras else None
    if body_template is None:
        raise ValueError("No body paragraph template found in example")

    elements_to_insert = []

    # DTR heading
    new_heading = clone_element(dtr_heading_el)
    set_all_runs_text(new_heading, f"Desktop Review (DTR) {dtr_num}")
    add_page_break_before(new_heading)
    elements_to_insert.append(new_heading)

    # Updating paragraph
    if cfg.get("updating_paras"):
        hw = ", ".join(cfg["updating_paras"])
        update_text = (
            f"This DTR updates the Session Border Controller (SBC) IOS XE software on "
            f"{hw} of SBC router platforms. The IOS XE software version is being "
            f"updated from {old_ver} to {new_ver}."
        )
        elements_to_insert.append(make_body_para(body_template, update_text))

    # Sustained paragraph(s)
    for plat, ver in cfg.get("sustained_paras", []):
        sustained_text = f"The {plat} will be sustained on the current software load of {ver}."
        elements_to_insert.append(make_body_para(body_template, sustained_text))

    # Similarity
    if cfg.get("similarity_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["similarity_text"]))

    # POA&M
    if cfg.get("poam_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["poam_text"]))

    # Release notes — entries (hyperlink or text_only) take priority over plain release_notes_text
    for rn_entry in cfg.get("release_notes_entries", []):
        # rn_entry = {"label": "...", "url": "..."} for hyperlink
        # rn_entry = {"label": "...", "text_only": True} for plain text (e.g. 404 fallback)
        if rn_entry.get("text_only"):
            elements_to_insert.append(make_body_para(body_template, rn_entry["label"]))
        else:
            elements_to_insert.append(
                make_hyperlink_para(doc, body_template, "Release Notes Link: ", rn_entry["label"], rn_entry["url"])
            )
    if not cfg.get("release_notes_entries") and cfg.get("release_notes_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["release_notes_text"]))

    # DTR Detailed Component Information heading
    new_detail_heading = clone_element(dtr_detail_heading_el)
    runs = new_detail_heading.findall(f".//{qn('w:r')}")
    for i, r in enumerate(runs):
        for t in r.findall(qn("w:t")):
            t.text = "DTR Detailed Component Information" if i == 0 else ""
    remove_page_break_before(new_detail_heading)
    elements_to_insert.append(new_detail_heading)

    # Component table — update by row index only (sustained rows untouched)
    new_table = clone_element(dtr_table_el)
    updating_indices = cfg.get("updating_row_indices", [])
    for idx, row_el in enumerate(new_table.findall(qn("w:tr"))):
        if idx == 0:
            continue
        cells = row_el.findall(qn("w:tc"))
        if len(cells) < 2:
            continue
        rel_cell = cells[1]
        if idx in updating_indices:
            for t in rel_cell.findall(f".//{qn('w:t')}"):
                if t.text and re.search(r"IOS XE\s+\S+", t.text):
                    t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
    elements_to_insert.append(new_table)

    # 5. Remove blank Heading1 spacer immediately before Management Description
    # The INITIAL doc has an empty Heading1 paragraph at the element just before
    # Management Description (between the INITIAL component table and Mgmt heading).
    # If left in place it renders as a spurious blank heading line after insertion.
    # NOTE: obtain mgmt_el by direct reference (not index) so it stays valid after removal.
    mgmt_el = None
    for el in body:
        if el.tag == qn("w:p") and "Management Description" in get_para_text(el):
            mgmt_el = el
            break
    if mgmt_el is None:
        raise ValueError("Could not find Management Description element (second pass)")

    prev_el = mgmt_el.getprevious()
    if prev_el is not None and prev_el.tag == qn("w:p"):
        prev_text = get_para_text(prev_el).strip()
        pStyle = prev_el.find(f".//{qn('w:pStyle')}")
        prev_style = pStyle.get(qn("w:val")) if pStyle is not None else "Normal"
        if not prev_text and prev_style == "Heading1":
            body.remove(prev_el)
            print("  Removed blank Heading1 spacer before Management Description")

    # 6. Insert all elements before Management Description (mgmt_el is a live reference)
    for el in elements_to_insert:
        mgmt_el.addprevious(el)

    # 7. Ensure Management Description has page break
    add_page_break_before(mgmt_el)

    # 8. Copy numbering definitions from example
    ex_numbering = example.part.numbering_part
    doc_numbering = doc.part.numbering_part
    if ex_numbering is not None and doc_numbering is not None:
        ex_num_el = ex_numbering._element
        doc_num_el = doc_numbering._element
        for abstract in ex_num_el.findall(qn("w:abstractNum")):
            aid = abstract.get(qn("w:abstractNumId"))
            if doc_num_el.find(f".//{qn('w:abstractNum')}[@{qn('w:abstractNumId')}='{aid}']") is None:
                doc_num_el.append(clone_element(abstract))
        for num in ex_num_el.findall(qn("w:num")):
            nid = num.get(qn("w:numId"))
            if doc_num_el.find(f".//{qn('w:num')}[@{qn('w:numId')}='{nid}']") is None:
                doc_num_el.append(clone_element(num))

    # 9. Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    # 10. Consistency cross-check
    verify_sysdesc_consistency(out_path, cfg)


# ---------------------------------------------------------------------------
# DTR002+ — incremental from previous draft
# ---------------------------------------------------------------------------

def generate_dtr_incremental(cfg: dict):
    source_path = cfg["source_path"]
    out_path = cfg["out_path"]
    dtr_num = cfg["dtr_num"]
    new_ver = cfg["new_ver"]
    old_ver = cfg["old_ver"]

    if not source_path.exists():
        raise FileNotFoundError(f"Source draft not found: {source_path}")

    doc = Document(str(source_path))
    body = doc.element.body

    # 1. Revision history row
    add_revision_row(doc.tables[0], cfg["rev_version"], cfg["approval"], dtr_num)

    # 2. Find body paragraph template (clone from last existing DTR section)
    #    (Management Description insertion point is re-scanned live in step 6)
    body_template = None
    for el in list(body):
        if el.tag == qn("w:p"):
            text = get_para_text(el)
            if "This DTR updates" in text or "will be sustained" in text:
                body_template = el

    if body_template is None:
        raise ValueError("Could not find body paragraph template")

    # 4. Find last DTR heading + detail heading + table to clone
    last_dtr_heading = last_detail_heading = last_dtr_table = None
    for el in list(body):
        if el.tag == qn("w:p"):
            text = get_para_text(el)
            if re.match(r"Desktop Review \(DTR\) \d+", text.strip()):
                last_dtr_heading = el
            if "DTR Detailed Component Information" in text:
                last_detail_heading = el
        elif el.tag == qn("w:tbl") and last_detail_heading is not None:
            last_dtr_table = el

    if any(x is None for x in [last_dtr_heading, last_detail_heading, last_dtr_table]):
        raise ValueError("Could not find DTR section elements to clone")

    # 5. Build new DTR section
    elements_to_insert = []

    new_heading = clone_element(last_dtr_heading)
    set_all_runs_text(new_heading, f"Desktop Review (DTR) {dtr_num}")
    add_page_break_before(new_heading)
    elements_to_insert.append(new_heading)

    if cfg.get("updating_paras"):
        hw = ", ".join(cfg["updating_paras"])
        update_text = (
            f"This DTR updates the Session Border Controller (SBC) IOS XE software on "
            f"{hw} of SBC router platforms. The IOS XE software version is being "
            f"updated from {old_ver} to {new_ver}."
        )
        elements_to_insert.append(make_body_para(body_template, update_text))

    for plat, ver in cfg.get("sustained_paras", []):
        sustained_text = f"The {plat} will be sustained on the current software load of {ver}."
        elements_to_insert.append(make_body_para(body_template, sustained_text))

    if cfg.get("similarity_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["similarity_text"]))

    if cfg.get("poam_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["poam_text"]))

    # Release notes — entries (hyperlink or text_only) take priority over plain release_notes_text
    for rn_entry in cfg.get("release_notes_entries", []):
        # rn_entry = {"label": "...", "url": "..."} for hyperlink
        # rn_entry = {"label": "...", "text_only": True} for plain text (e.g. 404 fallback)
        if rn_entry.get("text_only"):
            elements_to_insert.append(make_body_para(body_template, rn_entry["label"]))
        else:
            elements_to_insert.append(
                make_hyperlink_para(doc, body_template, "Release Notes Link: ", rn_entry["label"], rn_entry["url"])
            )
    if not cfg.get("release_notes_entries") and cfg.get("release_notes_text"):
        elements_to_insert.append(make_body_para(body_template, cfg["release_notes_text"]))

    new_detail_heading = clone_element(last_detail_heading)
    runs = new_detail_heading.findall(f".//{qn('w:r')}")
    for i, r in enumerate(runs):
        for t in r.findall(qn("w:t")):
            t.text = "DTR Detailed Component Information" if i == 0 else ""
    remove_page_break_before(new_detail_heading)
    elements_to_insert.append(new_detail_heading)

    # Clone and update component table
    new_table = clone_element(last_dtr_table)
    updating_indices = cfg.get("updating_row_indices", [])
    for idx, row_el in enumerate(new_table.findall(qn("w:tr"))):
        if idx == 0:
            continue
        cells = row_el.findall(qn("w:tc"))
        if len(cells) < 2:
            continue
        rel_cell = cells[1]
        if idx in updating_indices:
            for t in rel_cell.findall(f".//{qn('w:t')}"):
                if t.text and re.search(r"IOS XE\s+\S+", t.text):
                    t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
    elements_to_insert.append(new_table)

    # 6. Insert before Management Description — re-scan live to avoid stale index
    #    (prior insertions/removals via addprevious invalidate cached integer indices)
    mgmt_el = None
    for el in body:
        if el.tag == qn("w:p") and "Management Description" in get_para_text(el):
            mgmt_el = el
            break
    if mgmt_el is None:
        raise ValueError("Could not find Management Description element for insertion")
    for el in elements_to_insert:
        mgmt_el.addprevious(el)

    # 7. Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    # 8. Consistency cross-check
    verify_sysdesc_consistency(out_path, cfg)


# ---------------------------------------------------------------------------
# Post-save consistency cross-check
# ---------------------------------------------------------------------------

def verify_sysdesc_consistency(out_path: Path, cfg: dict):
    """
    After doc.save(), read back the output and verify that:
      - Updating platform rows in the last component table show the new version
      - Sustained platform rows show the sustained version (from sustained_paras)
    Prints [CHECK PASS] or [CHECK FAIL] per platform. Does not auto-fix.
    """
    print("\n--- Consistency Cross-Check ---")

    doc = Document(str(out_path))
    tables = doc.tables
    if not tables:
        print("[CHECK FAIL] No tables found in output document.")
        return

    # Find the most recently added component table by scanning for the last table with
    # at least `component_table_min_rows` rows. Default = 14 for SBC/CTN2026003:
    #   12 platform rows + 1 header row + 1 Notes row.
    # Override via cfg["component_table_min_rows"] when onboarding a CTN with a different
    # number of platforms (e.g. 8 platforms → 8×3 rows + 1 header + 1 Notes = 26 rows min).
    min_rows = cfg.get("component_table_min_rows", 14)
    comp_table = None
    for tbl in reversed(tables):
        if len(tbl.rows) >= min_rows:
            comp_table = tbl
            break
    if comp_table is None:
        print(f"[CHECK FAIL] Could not find a {min_rows}-row component table in output.")
        return
    rows = comp_table.rows

    new_ver = cfg.get("new_ver", "")
    updating_indices = cfg.get("updating_row_indices", [])
    sustained_paras = cfg.get("sustained_paras", [])  # list of (platform_label, version)

    # Build a map of sustained platform label -> expected version
    # e.g. ("ASR 1006-X", "IOS XE 17.18")
    # We also need to know which row indices belong to sustained platforms.
    # Derive sustained indices as all data row indices NOT in updating_indices.
    # Data rows = indices 1..len(rows)-2 (skip header row 0 and Notes row last).
    if len(rows) > 2:
        all_data_indices = list(range(1, len(rows) - 1))  # exclude header and Notes
    else:
        all_data_indices = []
    sustained_indices = [i for i in all_data_indices if i not in updating_indices]

    any_fail = False

    # Check updating platforms
    if updating_indices:
        for idx in updating_indices:
            if idx >= len(rows):
                continue
            cell_text = rows[idx].cells[1].text.strip()
            # Normalize: cfg new_ver may or may not include "IOS XE" prefix
            expected = new_ver if new_ver.startswith("IOS XE") else f"IOS XE {new_ver}"
            if expected in cell_text:
                print(f"[CHECK PASS] Row {idx} (updating): '{cell_text}' matches expected '{expected}'")
            else:
                print(f"[CHECK FAIL] Row {idx} (updating): table shows '{cell_text}' but expected '{expected}'")
                any_fail = True

    # Check sustained platforms — for DTR001 the sustained rows come from the Example doc
    # template which may carry a different version than the cfg sustained version.
    # Only check sustained rows for DTR002+ (incremental) where the table is cloned from
    # the prior draft. For DTR001 (dtr_num == 1), skip the table check for sustained rows
    # and rely on the body paragraph check below.
    # KNOWN LIMITATION: if the Example DTR001 file has stale sustained versions (e.g. IOS XE 17.15
    # instead of 17.18), this check will not catch it for DTR001. Verify rows 1-3 manually after
    # DTR001 generation. See Runbook Known Issue #12.
    dtr_num_check = cfg.get("dtr_num", 0)
    if sustained_paras and dtr_num_check > 1:
        for plat_label, expected_ver in sustained_paras:
            # expected_ver may or may not have "IOS XE" prefix
            expected = expected_ver if expected_ver.startswith("IOS XE") else f"IOS XE {expected_ver}"
            # Check all sustained rows — we can't map platform label -> exact row indices
            # without a full platform map, so we check ALL sustained rows carry expected_ver.
            # If any row shows a different version, flag it.
            row_results = []
            for idx in sustained_indices:
                if idx >= len(rows):
                    continue
                cell_text = rows[idx].cells[1].text.strip()
                row_results.append((idx, cell_text))

            if not row_results:
                print(f"[CHECK SKIP] {plat_label}: no sustained rows to check")
                continue

            mismatches = [(idx, txt) for idx, txt in row_results if expected not in txt]
            if mismatches:
                for idx, txt in mismatches:
                    print(f"[CHECK FAIL] Row {idx} (sustained/{plat_label}): table shows '{txt}' but expected '{expected}'")
                any_fail = True
            else:
                print(f"[CHECK PASS] {plat_label}: all {len(row_results)} sustained row(s) show '{expected}'")

    # Cross-check body paragraphs mention the same versions
    body_texts = " ".join(p.text for p in doc.paragraphs)
    if new_ver:
        # Accept either "IOS XE 26.4" or bare "26.4" in body text
        ver_check = new_ver if new_ver.startswith("IOS XE") else f"IOS XE {new_ver}"
        bare_ver = ver_check.replace("IOS XE ", "")
        if ver_check in body_texts or bare_ver in body_texts:
            print(f"[CHECK PASS] Body paragraphs reference new version '{ver_check}'")
        else:
            print(f"[CHECK FAIL] Body paragraphs do NOT reference new version '{ver_check}'")
            any_fail = True

    for plat_label, expected_ver in sustained_paras:
        expected = expected_ver if expected_ver.startswith("IOS XE") else f"IOS XE {expected_ver}"
        bare_expected = expected.replace("IOS XE ", "")
        if expected in body_texts or bare_expected in body_texts:
            print(f"[CHECK PASS] Body paragraphs reference sustained version '{expected}' for {plat_label}")
        else:
            print(f"[CHECK FAIL] Body paragraphs do NOT reference sustained version '{expected}' for {plat_label}")
            any_fail = True

    if any_fail:
        print("[CROSS-CHECK] WARN: One or more checks FAILED — review the draft before committing.")
    else:
        print("[CROSS-CHECK] OK: All checks passed.")
    print("--- End Cross-Check ---\n")


# ---------------------------------------------------------------------------
# Platform prompt helpers (used by main())
# ---------------------------------------------------------------------------

# Platform definitions specific to SBC/CTN2026003 component table structure
# When onboarding a new CTN with different platforms, define a new <CTN>_PLATFORMS list here
SBC_CTN2026003_PLATFORMS = [
    # row_indices = component table data rows (0=header, 13=Notes row)
    # Each platform has 3 rows: IWG (role col), SBC (role col), IWG/SBC (role col)
    {"label": "ASR 1006-X",    "body_label": "the ASR 1006-X",    "row_indices": [1, 2, 3]},   # rows 1–3:  ASR 1006-X IWG / SBC / IWG+SBC
    {"label": "C8300 series",  "body_label": "the C8300 series",  "row_indices": [4, 5, 6]},   # rows 4–6:  C8300 IWG / SBC / IWG+SBC
    {"label": "C8200 series",  "body_label": "the C8200 series",  "row_indices": [7, 8, 9]},   # rows 7–9:  C8200 IWG / SBC / IWG+SBC
    {"label": "C8000v series", "body_label": "the C8000v series", "row_indices": [10, 11, 12]}, # rows 10–12: C8000v IWG / SBC / IWG+SBC
]


def prompt_platform_statuses(platforms: list, old_ver: str) -> dict:
    """
    Interactively prompt the engineer for each platform's status (updating/sustained).

    Fast-path rules:
    - After any Updating selection: offer bulk shortcut for all remaining platforms
      (all updating to same version → one version prompt, done)
    - After any Sustained selection: offer bulk shortcut for all remaining platforms
      (all updating to same version → one version prompt, done)
    - Both shortcuts re-offered after every status selection while platforms remain

    Returns a dict with:
      - updating_paras: list of body_label strings for updating platforms
      - sustained_paras: list of (label, version) tuples for sustained platforms
      - updating_row_indices: flat list of row indices for all updating platforms
      - new_ver: the single new version (if bulk shortcut used); else None
    """
    updating_paras = []
    sustained_paras = []
    updating_row_indices = []
    new_ver = None

    remaining = list(platforms)

    while remaining:
        plat = remaining[0]
        remaining = remaining[1:]

        status = input(f"  {plat['label']} — [u]pdating or [s]ustained? ").strip().lower()
        while status not in ("u", "s", "updating", "sustained"):
            status = input(f"  Enter 'u' for updating or 's' for sustained: ").strip().lower()

        if status.startswith("s"):
            ver = input(f"  Sustained version for {plat['label']} (e.g. IOS XE 17.18): ").strip()
            sustained_paras.append((plat["label"], ver))
            print(f"  → {plat['label']}: sustained at {ver}")
        else:
            updating_paras.append(plat["body_label"])
            updating_row_indices.extend(plat["row_indices"])
            print(f"  → {plat['label']}: updating")

        # After every selection, offer bulk shortcut for remaining platforms
        if remaining:
            remaining_names = ", ".join(p["label"] for p in remaining)
            print(f"\n  Remaining: {remaining_names}")
            bulk = input(
                "  Are ALL remaining platforms updating to the same new version? [y/N]: "
            ).strip().lower()

            if bulk == "y":
                print(f"\n  This will mark ALL of the following as updating:")
                for p in remaining:
                    print(f"    - {p['label']}")
                confirm = input("  Confirm? [y/N]: ").strip().lower()
                if confirm == "y":
                    bulk_ver = input("  New IOS XE version for all (e.g. IOS XE 26.4): ").strip()
                    new_ver = bulk_ver
                    for p in remaining:
                        updating_paras.append(p["body_label"])
                        updating_row_indices.extend(p["row_indices"])
                        print(f"  → {p['label']}: updating to {bulk_ver}")
                    remaining = []  # all handled
                else:
                    print("  Bulk shortcut cancelled — continuing individually.")

    return {
        "updating_paras": updating_paras,
        "sustained_paras": sustained_paras,
        "updating_row_indices": updating_row_indices,
        "new_ver": new_ver,
    }


# ---------------------------------------------------------------------------
# Profiles
# ---------------------------------------------------------------------------

SBC_CTN2026003_SYSDESC_PROFILES = {
    # DTR001 seed profile — first-production version for SBC/CTN2026003.
    # CONFIRM new_ver, approval, and release_notes_text before generating.
    # This profile is presented to the engineer with a y/N confirm prompt before use.
    1: {
        "dtr_num": 1,
        "new_ver": "IOS XE 26.1",
        "old_ver": "IOS XE 17.18",
        "rev_version": "2.0",
        "approval": "May 2026",
        "updating_paras": ["the C8300 series", "the C8200 series", "the C8000v series"],
        "sustained_paras": [("ASR 1006-X", "IOS XE 17.18")],
        # Sustained version 17.18 confirmed from INITIAL source doc (CTN2026003 - DTR000 - INITIAL).
        # Note: the Example_ reference file is named IOS XE 17.15 but is never used as source.
        "updating_row_indices": list(range(4, 13)),  # rows 4–12; ASR rows 1–3 sustained
        "similarity_text": "",
        "poam_text": "",
        "release_notes_text": "Release Notes for all devices will be provided once they become available, expected Jul 2026.",
    },
    # DTR002+ have no pre-baked profiles — prompted at runtime.
}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def run(prod_cat: str, ctn: str, dtr_num: int, cfg_overrides: dict = None, engineer: str = "") -> None:
    """Callable entry point — invoked by the newd prompt sequence or directly via main().

    cfg_overrides: dict of values collected during the newd prompt sequence (e.g. platform
    statuses, new_ver, release_notes_entries). These ALWAYS take precedence over seed profile
    values for any key they supply. The profile only provides fallback defaults for fields
    the engineer did not explicitly answer (old_ver, rev_version, approval).
    """
    interactive = cfg_overrides is None  # True when called directly from main() — prompts for all fields
    if cfg_overrides is None:
        cfg_overrides = {}

    profiles = {}
    if prod_cat == "SBC" and ctn == "CTN2026003":
        profiles = SBC_CTN2026003_SYSDESC_PROFILES

    if dtr_num in profiles:
        # Start from profile as base for non-platform fields only
        profile = dict(profiles[dtr_num])
        print(f"\nLoaded profile defaults for {prod_cat}/{ctn} DTR{dtr_num:03d}:")
        print(f"  old_ver    : {profile['old_ver']}")
        print(f"  rev_version: {profile['rev_version']}")
        print(f"  approval   : {profile['approval']}")

        # Build cfg: profile base, then override with any engineer-provided values
        cfg = {
            "dtr_num": dtr_num,
            "old_ver":     cfg_overrides.get("old_ver",     profile["old_ver"]),
            "new_ver":     cfg_overrides.get("new_ver",     profile["new_ver"]),
            "rev_version": cfg_overrides.get("rev_version", profile["rev_version"]),
            "approval":    cfg_overrides.get("approval",    profile["approval"]),
            # Platform fields: ALWAYS from overrides — never from profile
            "updating_paras":       cfg_overrides.get("updating_paras",       []),
            "sustained_paras":      cfg_overrides.get("sustained_paras",      []),
            "updating_row_indices": cfg_overrides.get("updating_row_indices", []),
            "similarity_text":      cfg_overrides.get("similarity_text",      ""),
            "poam_text":            cfg_overrides.get("poam_text",            ""),
            "release_notes_text":   cfg_overrides.get("release_notes_text",   ""),
            "release_notes_entries": cfg_overrides.get("release_notes_entries", []),
        }

        # If platform fields were not supplied via overrides AND this is an interactive
        # (direct main()) call, prompt for them now rather than silently using empty lists.
        if interactive and not cfg["updating_paras"] and not cfg["sustained_paras"]:
            platforms = SBC_CTN2026003_PLATFORMS
            print("\nPlatform status (u = updating, s = sustained):")
            plat_result = prompt_platform_statuses(platforms, cfg["old_ver"])
            cfg["updating_paras"]       = plat_result["updating_paras"]
            cfg["sustained_paras"]      = plat_result["sustained_paras"]
            cfg["updating_row_indices"] = plat_result["updating_row_indices"]
            if plat_result["new_ver"]:
                cfg["new_ver"] = plat_result["new_ver"]
            elif not cfg["new_ver"]:
                cfg["new_ver"] = input("\nNew IOS XE version (e.g. IOS XE 26.4): ").strip()

        # If not interactive and platform fields still empty — hard error.
        if not interactive and not cfg_overrides.get("updating_paras") and not cfg_overrides.get("sustained_paras"):
            print("\n[ERROR] Platform statuses were not passed to run(). "
                  "Collect platform answers during the newd prompt sequence and pass via cfg_overrides.")
            return

    else:
        print(f"\nNo profile found for {prod_cat}/{ctn} DTR{dtr_num:03d}. Manual input required.")

        old_ver = cfg_overrides.get("old_ver") or input("Old IOS XE version (e.g. IOS XE 17.18): ").strip()
        rev_version = cfg_overrides.get("rev_version") or input("Revision history version number (e.g. 3.0): ").strip()

        if cfg_overrides.get("updating_paras") is not None:
            updating_paras       = cfg_overrides["updating_paras"]
            sustained_paras      = cfg_overrides["sustained_paras"]
            updating_row_indices = cfg_overrides["updating_row_indices"]
            new_ver              = cfg_overrides.get("new_ver", "")
        else:
            platforms = SBC_CTN2026003_PLATFORMS
            print("\nPlatform status (u = updating, s = sustained):")
            plat_result = prompt_platform_statuses(platforms, old_ver)
            updating_paras       = plat_result["updating_paras"]
            sustained_paras      = plat_result["sustained_paras"]
            updating_row_indices = plat_result["updating_row_indices"]
            new_ver = plat_result["new_ver"] or input("\nNew IOS XE version (e.g. IOS XE 26.4): ").strip()

        cfg = {
            "dtr_num": dtr_num,
            "new_ver": new_ver,
            "old_ver": old_ver,
            "rev_version": rev_version,
            "approval": cfg_overrides.get("approval", ""),
            "updating_paras":        updating_paras,
            "sustained_paras":       sustained_paras,
            "updating_row_indices":  updating_row_indices,
            "similarity_text":       cfg_overrides.get("similarity_text", ""),
            "poam_text":             cfg_overrides.get("poam_text", ""),
            "release_notes_text":    cfg_overrides.get("release_notes_text", ""),
            "release_notes_entries": cfg_overrides.get("release_notes_entries", []),
        }

    cfg["prod_cat"] = prod_cat
    cfg["ctn"] = ctn
    cfg["out_path"] = sysdesc_draft_path(prod_cat, ctn, dtr_num, cfg["new_ver"])

    if dtr_num == 1:
        generate_dtr001(cfg)
    else:
        cfg["source_path"] = sysdesc_draft_path(prod_cat, ctn, dtr_num - 1, cfg["old_ver"])
        generate_dtr_incremental(cfg)

    print("\nRunning validate_doc.py...")
    run_validate(cfg["out_path"])

    if not engineer:
        engineer = input("\nEngineer username (for Draft Log): ").strip() or get_git_username() or "unknown"
    append_draft_log(
        engineer=engineer,
        action="Generated",
        ctn=ctn,
        doc_type="System Description",
        dtr=f"DTR{dtr_num:03d}",
        version=cfg["new_ver"],
        reason="Via run_sysdesc.py parameterized runner",
    )

    print("\nDone.")


def main():
    print("\n=== System Description — Parameterized Runner ===\n")
    prod_cat = input("Product Category (e.g. SBC): ").strip()
    ctn = input("CTN (e.g. CTN2026003): ").strip()
    dtr_num = int(input("DTR Number (e.g. 1): ").strip())
    # main() passes no cfg_overrides so run() will prompt interactively for all fields.
    # This is intentional for direct script invocation — the newd flow passes overrides instead.
    run(prod_cat, ctn, dtr_num, cfg_overrides=None)


if __name__ == "__main__":
    main()
