#!/usr/bin/env python3
"""
run_fa.py — Parameterized Functionality Attestation runner.

Replaces all gen_dtr###_fa_*.py one-off scripts.
Invoked by the newd prompt sequence after Product Category, CTN, and Document Type
have been selected. Prompts for all FA-specific inputs, then generates the draft.

Usage (direct):
    python3 _Tools/run_fa.py

Or triggered automatically by newd after skill_base.md routes to this runner.
"""

import copy
import re
import sys
from datetime import datetime
from pathlib import Path

# runner_core.py lives in the same _Tools/ directory
sys.path.insert(0, str(Path(__file__).parent))
from runner_core import (
    qn, clone_element, get_para_text, make_body_para, set_para_list_paragraph,
    make_spacer, _set_cell_text, update_notes_row, apply_notes_by_label,
    apply_keep_next, apply_acronym_page_break, strip_acronym_cell_borders,
    fa_dir, fa_draft_path, fa_initial_path, fa_example_path,
    run_validate, append_draft_log, BASE,
)

from lxml import etree
from docx import Document


# ---------------------------------------------------------------------------
# DTR body paragraph template text
# Each entry: (dtr_num_match_str, hw_list, component_list)
# hw_list and component_list are used to construct the body paragraph.
# These are the values that were hardcoded per-script — now passed as params.
# ---------------------------------------------------------------------------

def _oxford(items: list) -> str:
    """Return an Oxford-comma joined string from a list of strings."""
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + f", and {items[-1]}"


def build_dtr_id_suffix(updating_display: list, new_platforms: list) -> str:
    """Build the DTR ID suffix for Table 2-1 Row 6 from updating display names + new platform display names.

    Format: Oxford-comma list without 'the' prefix (e.g. 'C8300 series, C8200 series, and C8000v series').
    Only updating platforms and new platforms are listed — sustained platforms are excluded.
    updating_display: list of human-readable platform names (not short match identifiers).
    """
    parts = list(updating_display)
    seen = {p.lower() for p in parts}
    for np in new_platforms:
        dn = np.get("display_name", np["name"])
        if dn.lower() not in seen:
            parts.append(dn)
            seen.add(dn.lower())
    return _oxford(parts)


def build_hw_list(updating_display: list, new_platforms: list) -> str:
    """Build the hw_list string from updating display names and new platform display names.

    Each entry is prefixed with 'the' (e.g. 'the C8300 series').
    updating_display: list of human-readable platform names (not short match identifiers).
    """
    parts = [f"the {p}" for p in updating_display]
    seen = {p.lower() for p in updating_display}
    for np in new_platforms:
        dn = np.get("display_name", np["name"])
        if dn.lower() not in seen:
            parts.append(f"the {dn}")
            seen.add(dn.lower())
    return _oxford(parts)


def build_component_list(components: list) -> str:
    """Build the component_list string from a list of component names.

    e.g. ['IWBC', 'SBC'] -> 'IWBC and SBC'
    """
    return _oxford(components)


def build_dtr_body_text(dtr_num: int, old_ver: str, new_ver: str,
                        hw_list: str, component_list: str,
                        new_platforms: list) -> str:
    """Build the standard DTR body paragraph text."""
    dtr_label = f"DTR{dtr_num:03d}"  # zero-padded, no space (e.g. DTR001)
    text = (
        f"{dtr_label} was requested to update the IOS XE software version from "
        f"{old_ver} to {new_ver} for the product component {component_list} on "
        f"{hw_list} of router platforms."
    )
    if new_platforms:
        # Deduplicate by name and use display_name if provided, else name
        seen = {}
        for p in new_platforms:
            n = p["name"]
            if n not in seen:
                seen[n] = p.get("display_name", n)
        unique_labels = list(seen.values())
        if len(unique_labels) == 1:
            plat_str = unique_labels[0]
        elif len(unique_labels) == 2:
            plat_str = f"{unique_labels[0]} and {unique_labels[1]}"
        else:
            plat_str = ", ".join(unique_labels[:-1]) + f", and {unique_labels[-1]}"
        text += (
            f" {dtr_label} also adds the {plat_str} of router platforms."
            f" The IOS XE software version is {new_ver}."
        )
    return text


# ---------------------------------------------------------------------------
# Step A: DTR001 — INITIAL source generation (table merges required)
# ---------------------------------------------------------------------------

def generate_dtr001(cfg: dict):
    """Generate FA DTR001 from INITIAL source document."""
    prod_cat = cfg["prod_cat"]
    ctn = cfg["ctn"]
    dtr_num = cfg["dtr_num"]
    new_ver = cfg["new_ver"]
    old_ver = cfg["old_ver"]
    out_path = cfg["out_path"]

    initial_path = fa_initial_path(prod_cat, ctn)
    example_path = fa_example_path(prod_cat, ctn)

    if not initial_path.exists():
        raise FileNotFoundError(f"INITIAL doc not found: {initial_path}")
    if not example_path.exists():
        raise FileNotFoundError(f"Example doc not found: {example_path}")

    doc = Document(str(initial_path))
    example = Document(str(example_path))
    body = doc.element.body
    ver_num = new_ver.split()[-1]

    # 1. UPDATE DOCUMENT DATE + GCT BANNER (clone from example paragraph 0)
    ex_body = example.element.body
    ex_p0 = ex_body[0]
    new_p0 = clone_element(ex_p0)
    new_runs = new_p0.findall(f".//{qn('w:r')}")
    text_run_count = 0
    for r in list(new_runs):
        if r.findall(qn("w:drawing")):
            continue
        t_els = r.findall(qn("w:t"))
        if t_els:
            text_run_count += 1
            if text_run_count == 2:
                t_els[0].text = f" {datetime.now().strftime('%d %B %Y')}"
                t_els[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                for extra_t in t_els[1:]:
                    r.remove(extra_t)
            elif text_run_count > 2:
                r.getparent().remove(r)
    # Reuse banner image rId11
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    for blip in new_p0.findall(f".//{{{a_ns}}}blip"):
        blip.set(f"{{{r_ns}}}embed", "rId11")
    old_p0 = body[0]
    old_p0.addnext(new_p0)
    body.remove(old_p0)

    # 2. CONDITIONS OF ATTESTATION — fix "Certification" → "Attestation"
    for p in doc.paragraphs:
        if "Conditions of Certification" in p.text:
            for r in p.runs:
                if "Conditions of Certification" in r.text:
                    r.text = r.text.replace("Conditions of Certification", "Conditions of Attestation")
            break

    # 3. CONDITIONS OF ATTESTATION — update version (Run 4 + Run 5)
    for p in doc.paragraphs:
        if "System Under Test" in p.text and "XE" in p.text:
            runs = p.runs
            runs[4].text = f" XE {ver_num}"
            runs[5].text = ""
            break

    # 4. REFERENCE (c) IOS XE VERSION (Paragraph 4) — content-search replace
    p4 = doc.paragraphs[4]
    p4_runs = p4.runs
    in_ref_c = False
    ref_c_start = ref_c_end = None
    for ri, r in enumerate(p4_runs):
        txt = r.text
        if "(c" in txt or ("(" in txt and ri + 1 < len(p4_runs) and p4_runs[ri + 1].text.startswith("c")):
            in_ref_c = True
        if in_ref_c and ref_c_start is None and "IOS XE" in txt:
            ref_c_start = ri
        if in_ref_c and ref_c_start is not None and ("\n" in txt or "(d" in txt):
            ref_c_end = ri
            break
    if ref_c_start is not None:
        if ref_c_end is None:
            ref_c_end = min(ref_c_start + 5, len(p4_runs))
        combined = "".join(p4_runs[i].text for i in range(ref_c_start, ref_c_end))
        updated = re.sub(r"(IOS XE\s+)[\d.]+(\s*IOS XE\s+[\d.]+)*", rf"\g<1>{ver_num}", combined)
        p4_runs[ref_c_start].text = updated
        for ci in range(ref_c_start + 1, ref_c_end):
            p4_runs[ci].text = ""

    # 5. FUNCTIONALITY TESTING DATES
    _update_testing_dates(doc, cfg["testing_start_date"], cfg["testing_start_year"],
                          cfg["testing_end_date"], cfg["testing_end_year"])

    # 6. SYSTEM REQUIREMENTS PARAGRAPH VERSION (Run 9 + Run 10)
    for p in doc.paragraphs:
        if "System Requirements" in p.text and "IOS XE" in p.text:
            runs = p.runs
            runs[9].text = f" IOS XE {ver_num}"
            runs[10].text = ""
            # Page break before
            ppr = p._element.find(qn("w:pPr"))
            if ppr is None:
                ppr = etree.SubElement(p._element, qn("w:pPr"))
            if ppr.find(qn("w:pageBreakBefore")) is None:
                etree.SubElement(ppr, qn("w:pageBreakBefore"))
            break

    # 7. TABLE 1 — TDR CONDITIONS (DTR001: clone 4-cell row from example if TDR provided)
    if cfg.get("tdr_row"):
        ex_tdr_tbl = example.tables[0]._tbl
        ex_tdr_rows = ex_tdr_tbl.findall(qn("w:tr"))
        new_tdr_row = clone_element(ex_tdr_rows[1])
        new_tdr_tcs = new_tdr_row.findall(qn("w:tc"))
        tdr = cfg["tdr_row"]
        tdr_vals = (tdr["num"], tdr["desc"], tdr["impact"], tdr["remarks"]) if isinstance(tdr, dict) else tdr
        for tc_el, val in zip(new_tdr_tcs, tdr_vals):
            for t in tc_el.findall(f".//{qn('w:t')}"):
                t.text = val
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                break
        doc_tdr_tbl = doc.tables[0]._tbl
        doc_tdr_rows = doc_tdr_tbl.findall(qn("w:tr"))
        old_row = doc_tdr_rows[1]
        doc_tdr_tbl.replace(old_row, new_tdr_row)

    # 8. TABLE 2-1 SYSTEM REQUIREMENTS ID
    sysreq_table = doc.tables[1]
    _set_sysreq_version(sysreq_table, new_ver)
    _set_sysreq_dtr_id(sysreq_table, dtr_num, cfg["dtr_id_suffix"])

    # 9. DTR BODY PARAGRAPHS
    _insert_dtr_body_paragraphs_dtr001(doc, body, cfg)

    # 10. TABLE 4 — UPDATE VERSIONS
    _update_table4_versions_dtr001(doc, cfg)

    # 10a. INSERT NEW PLATFORM ROWS (DTR001)
    if cfg.get("new_platforms"):
        _insert_new_platforms_dtr001(doc, cfg["new_platforms"])

    # 11. SCALE FIGURE 2-2 DIAGRAM
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    SCALE = 0.90
    for p in doc.paragraphs:
        for ext in p._element.findall(f".//{{{WP_NS}}}extent"):
            cx = int(ext.get("cx", "0"))
            cy = int(ext.get("cy", "0"))
            if cy > 3000000:  # large system diagram threshold (EMU)
                ext.set("cx", str(int(cx * SCALE)))
                ext.set("cy", str(int(cy * SCALE)))

    # 12. TABLE 3 SUT INTERFACE STATUS — keep as 2 separate tables, set tblHeader on each
    #     INITIAL table 2 has both sections: Network Mgmt (rows 0-3) + Network Interfaces (rows 4-7)
    #     INITIAL table 3 is continuation of Network Interfaces (rows 0-5 with header+data+NOTES)
    #     Move Network Interfaces rows (4+) from table 2 into table 3 (before row 2, after header rows)
    #     so table 2 = Network Mgmt only, table 3 = all Network Interfaces
    tbl2_el = doc.tables[2]._tbl
    tbl3_el = doc.tables[3]._tbl
    tbl2_rows = tbl2_el.findall(qn("w:tr"))
    tbl3_rows = tbl3_el.findall(qn("w:tr"))
    # Rows 4+ in table 2 are Network Interfaces data (after section header + col headers in table 3)
    # Move them: insert after row 1 (col headers) in table 3, before row 2 (first data row)
    net_intf_rows = tbl2_rows[4:]  # rows 4-7: section header, col headers, 2 data rows
    # Remove the section header + col header rows (4,5) since table 3 already has them
    # Only move data rows (6+)
    data_rows_to_move = tbl2_rows[6:]  # actual data rows from Network Interfaces in table 2
    insert_before = tbl3_rows[2] if len(tbl3_rows) > 2 else None  # insert before first data row in table 3
    for row in data_rows_to_move:
        if insert_before is not None:
            insert_before.addprevious(row)  # moves the element (no need to remove separately)
        else:
            tbl3_el.append(row)
    # Remove the now-orphaned section header + col header rows (4,5) from table 2
    for row in tbl2_rows[4:6]:
        tbl2_el.remove(row)
    # Set tblHeader on rows 0+1 of each table
    for ti in (2, 3):
        tbl_el = doc.tables[ti]._tbl
        tbl_rows = tbl_el.findall(qn("w:tr"))
        for row_idx in (0, 1):  # section header + column headers
            if row_idx < len(tbl_rows):
                trpr = tbl_rows[row_idx].find(qn("w:trPr"))
                if trpr is None:
                    trpr = etree.SubElement(tbl_rows[row_idx], qn("w:trPr"))
                    tbl_rows[row_idx].insert(0, trpr)
                hdr_el = trpr.find(qn("w:tblHeader"))
                if hdr_el is None:
                    hdr_el = etree.SubElement(trpr, qn("w:tblHeader"))
                hdr_el.set(qn("w:val"), "1")
    # Remove empty spacer between tables 2 and 3
    for el in list(body):
        if el is tbl2_el:
            nxt = el.getnext()
            while nxt is not None and nxt is not tbl3_el:
                if nxt.tag == qn("w:p") and not get_para_text(nxt).strip():
                    to_del = nxt
                    nxt = nxt.getnext()
                    body.remove(to_del)
                else:
                    break
            break
    # Fix double border between tables: remove bottom border of table 2 last row
    # and top border of table 3 first row so only one line renders
    tbl2_rows_final = tbl2_el.findall(qn("w:tr"))
    last_row = tbl2_rows_final[-1]
    for cell in last_row.findall(qn("w:tc")):
        tcPr = cell.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = etree.SubElement(cell, qn("w:tcPr"))
            cell.insert(0, tcPr)
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
        bottom = tcBorders.find(qn("w:bottom"))
        if bottom is None:
            bottom = etree.SubElement(tcBorders, qn("w:bottom"))
        bottom.set(qn("w:val"), "nil")

    # 13. MERGE TABLE 4 HW/SW/FW VERSION ID + SPLIT PRODUCT COMPONENTS
    #     first_idx=4 (Product ID table), second_idx=5 (Product Components table)
    #     after prior merges, these are the correct 0-based table indices in body.
    _merge_and_split_table4(doc, body, first_idx=4, second_idx=5)

    # 14. SET tblHeader
    #     Post-split: 0=Cond,1=SysReq,2=IntfMgmt,3=IntfNet,
    #     4=ProdID,5=ProdComp(part1),6=ProdComp(part2),7=CR/FR,8=Acronym
    #     Table 4 (Product ID) — no tblHeader (short table)
    #     Tables 5 & 6 (Product Components) — no tblHeader per user preference (continuous table, no header repeat)
    _add_tbl_header_to_table(doc, 7)   # CR/FR header row

    # 15. (removed — IWBC is the correct col 0 name in DTR001+; no rename needed)

    # 16. BOLD+UNDERLINE TESTED COMPONENTS (tables 5 and 6 = Product Components parts 1 and 2)
    _apply_bold_underline(doc, cfg["bold_underline_models"], table_indices=[5, 6])

    # 17. REMOVE EMPTY PARAGRAPHS BETWEEN PRODUCT COMPONENTS TABLES (5 and 6)
    #     AND BEFORE "System Description" — these cause blank pages
    body_els = list(body)
    tbl5_el = doc.tables[5]._tbl
    tbl6_el = doc.tables[6]._tbl
    i5 = body_els.index(tbl5_el)
    i6 = body_els.index(tbl6_el)
    for i in range(i5 + 1, i6):
        el = body_els[i]
        if el.tag == qn("w:p") and not get_para_text(el).strip():
            body.remove(el)

    body_els = list(body)
    for i, el in enumerate(body_els):
        if el.tag == qn("w:p") and get_para_text(el).strip() == "System Description":
            removed = 0
            idx = i - 1
            while removed < 2 and idx >= 0:
                prev = body_els[idx]
                if prev.tag == qn("w:p") and not get_para_text(prev).strip():
                    body.remove(prev)
                    removed += 1
                idx -= 1
            break

    # 18–19. KEEP TABLE TITLES + ACRONYM PAGE BREAK
    apply_keep_next(body)
    apply_acronym_page_break(body)
    strip_acronym_cell_borders(doc)

    # 20. NOTES ROWS (Table 3 and Table 5 — by label)
    # Table 4 (Product Components) notes applied BEFORE step 20a removes Table 6's
    # header row, because find_notes_tables() labels by first-cell text which becomes
    # 'IWG' after the header is removed.
    if cfg.get("table4_notes") is not None:
        apply_notes_by_label(doc, "Product Components", cfg["table4_notes"])
    if cfg.get("table3_notes") is not None:
        apply_notes_by_label(doc, "Network Interfaces", cfg["table3_notes"])
    if cfg.get("table5_notes") is not None:
        apply_notes_by_label(doc, "CR/FR ID", cfg["table5_notes"])

    # 20a. REMOVE DUPLICATE COLUMN HEADER from Table 6 (INITIAL Table 5 continuation).
    #      Done after _apply_notes so the "Product Components" label is still present
    #      when find_notes_tables() scans for the NOTES row.
    #      INITIAL Table 5 starts with its own "Product Components / Component Name / ..."
    #      header row (row 0). Since both Product Components tables are visually continuous
    #      and tblHeader is intentionally NOT set, this row appears as a spurious
    #      mid-table header on page breaks. Remove it here.
    tbl6 = doc.tables[6]._tbl
    tbl6_rows = tbl6.findall(qn("w:tr"))
    if tbl6_rows:
        row0_text = "".join(t.text or "" for t in tbl6_rows[0].findall(f".//{qn('w:t')}"))
        if "Product Components" in row0_text or "Component Name" in row0_text:
            tbl6.remove(tbl6_rows[0])

    # 20b. NORMALIZE Table 6 tblPr and tblGrid to match Table 5 exactly.
    #      INITIAL Table 5 has tblW=0/auto and no tblLayout — columns drift out of
    #      alignment with Table 5 (Product Components part 1). Copy tblW, tblLayout,
    #      and tblGrid from Table 5 into Table 6 so column lines align perfectly.
    tbl5_el = doc.tables[5]._tbl
    tbl5_tblpr = tbl5_el.find(qn("w:tblPr"))
    tbl5_tblgrid = tbl5_el.find(qn("w:tblGrid"))
    tbl6_tblpr = tbl6.find(qn("w:tblPr"))
    tbl6_tblgrid = tbl6.find(qn("w:tblGrid"))

    if tbl5_tblpr is not None and tbl6_tblpr is not None:
        # Copy tblW
        tbl5_w = tbl5_tblpr.find(qn("w:tblW"))
        tbl6_w = tbl6_tblpr.find(qn("w:tblW"))
        if tbl5_w is not None and tbl6_w is not None:
            tbl6_w.set(qn("w:w"), tbl5_w.get(qn("w:w")))
            tbl6_w.set(qn("w:type"), tbl5_w.get(qn("w:type")))
        # Copy tblLayout (add if missing)
        tbl5_layout = tbl5_tblpr.find(qn("w:tblLayout"))
        tbl6_layout = tbl6_tblpr.find(qn("w:tblLayout"))
        if tbl5_layout is not None:
            if tbl6_layout is None:
                tbl6_layout = etree.SubElement(tbl6_tblpr, qn("w:tblLayout"))
            tbl6_layout.set(qn("w:type"), tbl5_layout.get(qn("w:type")))
        # Replace tblGrid with Table 5's grid
        if tbl5_tblgrid is not None and tbl6_tblgrid is not None:
            new_grid = copy.deepcopy(tbl5_tblgrid)
            tbl6.replace(tbl6_tblgrid, new_grid)

    # 21. SAVE
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


# ---------------------------------------------------------------------------
# Step B: DTR002+ — incremental generation from previous draft
# ---------------------------------------------------------------------------

def generate_dtr_incremental(cfg: dict):
    """Generate FA DTR002+ from the previous draft."""
    source_path = cfg["source_path"]
    out_path = cfg["out_path"]
    dtr_num = cfg["dtr_num"]
    new_ver = cfg["new_ver"]
    ver_num = new_ver.split()[-1]

    if not source_path.exists():
        raise FileNotFoundError(f"Source draft not found: {source_path}")

    doc = Document(str(source_path))
    body = doc.element.body

    # 1. UPDATE DOCUMENT DATE
    p0 = body[0]
    text_run_count = 0
    for r in list(p0.findall(f".//{qn('w:r')}")):
        if r.findall(qn("w:drawing")):
            continue
        t_els = r.findall(qn("w:t"))
        if t_els:
            text_run_count += 1
            if text_run_count == 2:
                t_els[0].text = f" {datetime.now().strftime('%d %B %Y')}"
                t_els[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                for extra_t in t_els[1:]:
                    r.remove(extra_t)
            elif text_run_count > 2:
                r.getparent().remove(r)

    # 2. FUNCTIONALITY TESTING DATES
    _update_testing_dates(doc, cfg["testing_start_date"], cfg["testing_start_year"],
                          cfg["testing_end_date"], cfg["testing_end_year"])

    # 3. REFERENCE (c) IOS XE VERSION
    p4 = doc.paragraphs[4]
    p4_runs = p4.runs
    in_ref_c = False
    ref_c_start = ref_c_end = None
    for ri, r in enumerate(p4_runs):
        txt = r.text
        if "(c" in txt or ("(" in txt and ri + 1 < len(p4_runs) and p4_runs[ri + 1].text.startswith("c")):
            in_ref_c = True
        if in_ref_c and ref_c_start is None and "IOS XE" in txt:
            ref_c_start = ri
        if in_ref_c and ref_c_start is not None and ("\n" in txt or "(d" in txt):
            ref_c_end = ri
            break
    if ref_c_start is not None:
        if ref_c_end is None:
            ref_c_end = min(ref_c_start + 5, len(p4_runs))
        combined = "".join(p4_runs[i].text for i in range(ref_c_start, ref_c_end))
        updated = re.sub(r"(IOS XE\s+)[\d.]+(\s*IOS XE\s+[\d.]+)*", rf"\g<1>{ver_num}", combined)
        p4_runs[ref_c_start].text = updated
        for ci in range(ref_c_start + 1, ref_c_end):
            p4_runs[ci].text = ""

    # 4. CONDITIONS OF ATTESTATION VERSION
    for p in doc.paragraphs:
        if "System Under Test" in p.text and "XE" in p.text:
            for r in p.runs:
                if r.text and re.search(r"XE\s+\d+\.\d+", r.text):
                    r.text = re.sub(r"XE\s+\d+\.\d+", f"XE {ver_num}", r.text)
            break

    # 5. SYSTEM REQUIREMENTS PARAGRAPH VERSION
    for p in doc.paragraphs:
        if "System Requirements" in p.text and "IOS XE" in p.text:
            for r in p.runs:
                if r.text and re.search(r"IOS XE\s+\d+\.\d+", r.text):
                    r.text = re.sub(r"IOS XE\s+\d+\.\d+", new_ver, r.text)
            break

    # 6. TABLE 2-1 SYSTEM REQUIREMENTS ID
    sysreq_table = doc.tables[1]
    _set_sysreq_version(sysreq_table, new_ver)
    _set_sysreq_dtr_id(sysreq_table, dtr_num, cfg["dtr_id_suffix"])

    # 7. TABLE 1 — TDR CONDITIONS (incremental: add/update rows)
    _update_tdr_table_incremental(doc, cfg)

    # 8. DTR BODY PARAGRAPHS (incremental)
    _insert_dtr_body_paragraphs_incremental(doc, body, cfg)

    # 9. TABLE 4 — UPDATE VERSIONS (post-gen table indices 4, 5)
    _update_table4_versions_incremental(doc, cfg)

    # 9a. INSERT NEW PLATFORM ROWS
    if cfg.get("new_platforms"):
        _insert_new_platforms_incremental(doc, cfg["new_platforms"])

    # 10. BOLD+UNDERLINE TESTED MODELS
    _apply_bold_underline(doc, cfg["bold_underline_models"], table_indices=[4, 5])

    # 10a. NO tblHeader on Table 4 (Product Components) — continuous table, no header repeat on page breaks

    # 11. KEEP TABLE TITLES + ACRONYM PAGE BREAK
    body = doc.element.body
    apply_keep_next(body)
    apply_acronym_page_break(body)
    strip_acronym_cell_borders(doc)

    # 12. NOTES ROWS
    _apply_notes(doc, cfg)

    # 13. SAVE
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


# ---------------------------------------------------------------------------
# Shared sub-operations
# ---------------------------------------------------------------------------

def _update_testing_dates(doc, start_date, start_year, end_date, end_year):
    for p in doc.paragraphs:
        if "Test Details" in p.text and "functionality testing from" in p.text:
            runs = p.runs
            runs[12].text = start_date       # run 12: start date string (e.g. "14 May")
            runs[13].text = f" {start_year}" # run 13: start year with leading space
            if len(runs) > 14:
                runs[14].text = ""           # run 14: clear any leftover text between dates
            runs[16].text = end_date         # run 16: end date string
            runs[17].text = f" {end_year}"   # run 17: end year with leading space
            break


def _set_sysreq_version(sysreq_table, new_ver: str):
    ver_cell = sysreq_table.rows[2].cells[1]
    for p in ver_cell.paragraphs:
        if p.runs:
            p.runs[0].text = new_ver
            for r in p.runs[1:]:
                r.text = ""
        break


def _set_sysreq_dtr_id(sysreq_table, dtr_num: int, dtr_id_suffix: str):
    dtr_id_cell = sysreq_table.rows[6].cells[1]
    # Strip outer parentheses from suffix if present to avoid double-wrapping
    suffix = dtr_id_suffix.strip()
    if suffix.startswith("(") and suffix.endswith(")"):
        suffix = suffix[1:-1]
    new_dtr_id = f"00{dtr_num} ({suffix})"
    for p in dtr_id_cell.paragraphs:
        if p.runs:
            p.runs[0].text = new_dtr_id
            for r in p.runs[1:]:
                r.text = ""
        break


def _update_tdr_table_incremental(doc, cfg: dict):
    """Handle TDR table updates for DTR002+."""
    tdr_table = doc.tables[0]

    def _set_cell_value(row, col_idx, value):
        cell = row.cells[col_idx]
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].text = value
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = value
            break

    # Full row updates (replace all 4 fields for a given TDR number)
    for update in cfg.get("tdr_updates_full", []):
        tdr_num_str, new_desc, new_op, new_remarks = update
        for row in tdr_table.rows[1:]:
            if row.cells[0].text.strip() == tdr_num_str:
                _set_cell_value(row, 1, new_desc)
                _set_cell_value(row, 2, new_op)
                _set_cell_value(row, 3, new_remarks)
                break

    # Add new TDR row
    tdr_row = cfg.get("tdr_row")
    if tdr_row:
        # Normalise: tdr_row may be a dict {num, desc, impact, remarks} or a 4-tuple
        if isinstance(tdr_row, dict):
            tdr_row = (tdr_row["num"], tdr_row["desc"], tdr_row["impact"], tdr_row["remarks"])

        def _set_row_values(row, values):
            for i, val in enumerate(values):
                _set_cell_value(row, i, val)

        tdr_num_new = tdr_row[0]
        existing_row = None
        for row in tdr_table.rows[1:]:
            if row.cells[0].text.strip() == tdr_num_new:
                existing_row = row
                break

        if existing_row is not None:
            _set_row_values(existing_row, tdr_row)
        else:
            last_row = tdr_table.rows[-1]
            last_row_text = last_row.cells[0].text.strip()
            if last_row_text.lower() == "none" or last_row_text == "":
                _set_row_values(last_row, tdr_row)
            else:
                tbl_el = tdr_table._tbl
                last_tr = tbl_el[-1]
                new_tr = copy.deepcopy(last_tr)
                tbl_el.append(new_tr)
                _set_row_values(tdr_table.rows[-1], tdr_row)


def _insert_dtr_body_paragraphs_dtr001(doc, body, cfg: dict):
    """Insert DTR body paragraphs for DTR001 (first DTR, no prior DTR paragraphs)."""
    func_status_idx = None
    initial_request_idx = None
    for i, el in enumerate(body):
        if el.tag == qn("w:p"):
            text = get_para_text(el)
            if "Functionality Status" in text:
                func_status_idx = i
            if "initial request" in text.lower():
                initial_request_idx = i

    if func_status_idx is None:
        raise ValueError("Could not find 'Functionality Status' paragraph")
    if initial_request_idx is None:
        raise ValueError("Could not find 'initial request' paragraph for template")

    body_template = body[initial_request_idx]

    # Find empty List Paragraph for spacer template
    spacer_template = None
    for el in body:
        if el.tag == qn("w:p"):
            pstyle_els = el.findall(f".//{qn('w:pStyle')}")
            if not get_para_text(el).strip() and pstyle_els and pstyle_els[0].get(qn("w:val")) == "ListParagraph":
                spacer_template = el
                break

    func_status_el = body[func_status_idx]
    paras_to_insert = _build_dtr_paragraphs(body_template, cfg)

    if spacer_template is not None:
        spacer = make_spacer(spacer_template)
        func_status_el.addprevious(spacer)

    anchor = func_status_el.getprevious() if spacer_template is not None else func_status_el
    for para in paras_to_insert:
        anchor.addprevious(para)


def _insert_dtr_body_paragraphs_incremental(doc, body, cfg: dict):
    """Insert DTR body paragraphs for DTR002+ (append after existing DTR blocks)."""
    func_status_idx = None
    body_template_idx = None
    dtr_num = cfg["dtr_num"]

    # Find the most recent DTR body paragraph to use as template
    for i, el in enumerate(body):
        if el.tag == qn("w:p"):
            text = get_para_text(el)
            if "Functionality Status" in text:
                func_status_idx = i
            # Look for any prior "DTR N was requested" paragraph
            if re.search(r"DTR\d+ was requested", text):
                body_template_idx = i

    if func_status_idx is None:
        raise ValueError("Could not find 'Functionality Status' paragraph")
    if body_template_idx is None:
        raise ValueError("Could not find body paragraph template (DTR N was requested)")

    body_template = body[body_template_idx]
    func_status_el = body[func_status_idx]

    paras_to_insert = _build_dtr_paragraphs(body_template, cfg)

    # Anchor: the existing spacer immediately before Functionality Status, or Functionality Status itself
    spacer_before_func = func_status_el.getprevious()
    anchor = (spacer_before_func
              if spacer_before_func is not None and not get_para_text(spacer_before_func).strip()
              else func_status_el)

    # Empty spacer between previous DTR block and this one
    dtr_spacer = make_body_para(body_template, "")
    set_para_list_paragraph(dtr_spacer)
    for r in dtr_spacer.findall(f".//{qn('w:r')}"):
        r.getparent().remove(r)
    anchor.addprevious(dtr_spacer)

    for para in paras_to_insert:
        anchor.addprevious(para)


def _build_dtr_paragraphs(body_template, cfg: dict) -> list:
    """Build the list of paragraph elements to insert for this DTR."""
    dtr_num = cfg["dtr_num"]
    old_ver = cfg["old_ver"]
    new_ver = cfg["new_ver"]
    hw_list = cfg["hw_list"]
    component_list = cfg["component_list"]
    new_platforms = cfg.get("new_platforms", [])
    sustained_platforms = cfg.get("sustained_platforms", [])
    similarity_text = cfg.get("similarity_text", "")
    poam_text = cfg.get("poam_text", "")

    dtr_text = build_dtr_body_text(dtr_num, old_ver, new_ver, hw_list, component_list, new_platforms)
    paras = []

    dtr_para = make_body_para(body_template, dtr_text)
    set_para_list_paragraph(dtr_para)
    paras.append(dtr_para)

    for plat_name, plat_ver in sustained_platforms:
        sustain_text = f"The {plat_name} will be sustained on the current software load of {plat_ver}."
        sustain_para = make_body_para(body_template, sustain_text)
        set_para_list_paragraph(sustain_para)
        paras.append(sustain_para)

    if similarity_text:
        sim_para = make_body_para(body_template, similarity_text)
        set_para_list_paragraph(sim_para)
        paras.append(sim_para)

    if poam_text:
        poam_para = make_body_para(body_template, poam_text)
        set_para_list_paragraph(poam_para)
        paras.append(poam_para)

    return paras


def _update_table4_versions_dtr001(doc, cfg: dict):
    """Update table version strings for DTR001.

    Pre-merge table indices (before step 12 Table 3 split):
      4 = Product ID + first Product Components batch (INITIAL table 4)
      5 = Product Components part 2 (INITIAL table 5)
      6 = Product Components part 3 (INITIAL table 6)
    """
    new_ver = cfg["new_ver"]
    updating = cfg["updating_platforms_table"]
    sustained_names = [s[0] for s in cfg.get("sustained_platforms", [])]

    for ti in [4, 5, 6]:
        tbl_el = doc.tables[ti]._tbl
        for row_el in tbl_el.findall(qn("w:tr")):
            cells = row_el.findall(qn("w:tc"))
            if len(cells) < 2:
                continue
            col0_text = "".join(t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")).strip()
            if col0_text == "Software Release":
                for cell in cells[1:]:
                    for t in cell.findall(f".//{qn('w:t')}"):
                        if t.text and re.search(r"IOS XE\s+\S+", t.text):
                            t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
                continue
            if len(cells) < 4:
                continue
            comp_name = "".join(t.text or "" for t in cells[1].findall(f".//{qn('w:t')}"))
            comp_name_nospace = comp_name.replace(" ", "")
            if any(plat.replace(" ", "") in comp_name_nospace for plat in sustained_names):
                continue
            if any(plat.replace(" ", "") in comp_name_nospace for plat in updating):
                for t in cells[3].findall(f".//{qn('w:t')}"):
                    if t.text and re.search(r"IOS XE\s+\S+", t.text):
                        t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
                continue
            if len(cells) >= 5:
                for col_idx in [2, 3]:
                    for t in cells[col_idx].findall(f".//{qn('w:t')}"):
                        if t.text and re.search(r"IOS XE\s+\S+", t.text):
                            t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)


def _update_table4_versions_incremental(doc, cfg: dict):
    """Update table version strings for DTR002+ (post-gen table indices 4, 5)."""
    new_ver = cfg["new_ver"]
    updating = cfg["updating_platforms_table"]
    # sustained_by_group: list of (comp_name_substr, group_name_substr)
    sustained_by_group = cfg.get("sustained_by_group", [])

    for ti in [4, 5]:
        tbl_el = doc.tables[ti]._tbl
        for row_el in tbl_el.findall(qn("w:tr")):
            cells = row_el.findall(qn("w:tc"))
            if len(cells) < 2:
                continue
            col0_text = "".join(t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")).strip()
            if col0_text == "Software Release":
                for cell in cells[1:]:
                    for t in cell.findall(f".//{qn('w:t')}"):
                        if t.text and re.search(r"IOS XE\s+\S+", t.text):
                            t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
                continue
            if len(cells) < 4:
                continue
            comp_name = "".join(t.text or "" for t in cells[1].findall(f".//{qn('w:t')}"))
            comp_name_nospace = comp_name.replace(" ", "")
            # Skip sustained rows (matched by component + group)
            skip = False
            for comp_substr, group_substr in sustained_by_group:
                if comp_substr.replace(" ", "") in comp_name_nospace and group_substr in col0_text:
                    skip = True
                    break
            if skip:
                continue
            col3_text = "".join(t.text or "" for t in cells[3].findall(f".//{qn('w:t')}"))
            if re.search(r"IOS XE", col3_text):
                for t in cells[3].findall(f".//{qn('w:t')}"):
                    if t.text and re.search(r"IOS XE\s+\S+", t.text):
                        t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)
                continue
            if len(cells) >= 5:
                for col_idx in [2, 3]:
                    for t in cells[col_idx].findall(f".//{qn('w:t')}"):
                        if t.text and re.search(r"IOS XE\s+\S+", t.text):
                            t.text = re.sub(r"IOS XE\s+\S+", new_ver, t.text)


def _insert_new_platforms_dtr001(doc, new_platforms: list):
    """Insert new platform rows into Table 4 for DTR001 (pre-merge, INITIAL indices).

    Each platform entry has a 'group' key — the product component group to insert into
    (e.g. 'IWBC', 'SBC', 'IWG'). The row is inserted after the last existing
    row whose col 0 matches that group exactly.
    """
    for plat in new_platforms:
        target_group = plat["group"]
        match_groups = {target_group}

        last_group_row = None
        last_group_tbl = None

        for ti in [4, 5]:
            table_el = doc.tables[ti]._tbl
            rows = table_el.findall(qn("w:tr"))
            for ri, row in enumerate(rows):
                cells = row.findall(qn("w:tc"))
                if len(cells) < 4:
                    continue
                cell0_text = "".join(t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")).strip()
                col1_text = "".join(t.text or "" for t in cells[1].findall(f".//{qn('w:t')}")).strip()
                if "NOTE" in col1_text or "Component Name" in col1_text:
                    continue
                if cell0_text in match_groups:
                    last_group_row = row
                    last_group_tbl = table_el

        if last_group_row is not None:
            new_row = copy.deepcopy(last_group_row)
            new_cells = new_row.findall(qn("w:tc"))
            # Col 0: set group name — write target_group text directly (no vMerge in updated INITIAL)
            cell0_tcpr = new_cells[0].find(qn("w:tcPr"))
            vm = cell0_tcpr.find(qn("w:vMerge")) if cell0_tcpr is not None else None
            if vm is not None:
                # Legacy vMerge path — clear continuation, set text
                if qn("w:val") in vm.attrib:
                    del vm.attrib[qn("w:val")]
                for p in new_cells[0].findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        p.remove(r)
            else:
                _set_cell_text(new_cells[0], target_group)
            # Col 1: component name
            _set_cell_text(new_cells[1], plat["name"])
            # Col 2: sub-component
            _set_cell_text(new_cells[2], plat.get("sub_component", "NA"))
            # Col 3: tested version
            _set_cell_text(new_cells[3], plat.get("tested_version", ""))
            # Col 4: function
            if len(new_cells) >= 5:
                _set_cell_text(new_cells[4], plat.get("function", ""))
            last_group_row.addnext(new_row)


def _insert_new_platforms_incremental(doc, new_platforms: list):
    """Insert new platform rows into Table 4 for DTR002+ (post-gen, index 5 = Product Components)."""
    pc_tbl = doc.tables[5]._tbl
    for plat in new_platforms:
        prod_comp_group = plat["group"]
        all_rows = pc_tbl.findall(qn("w:tr"))
        current_group = None
        last_group_row = None
        for ri, row_el in enumerate(all_rows):
            cells = row_el.findall(qn("w:tc"))
            if len(cells) < 2:
                continue
            cell0 = cells[0]
            tcpr = cell0.find(qn("w:tcPr"))
            vm = tcpr.find(qn("w:vMerge")) if tcpr is not None else None
            vm_val = vm.get(qn("w:val"), "continue") if vm is not None else "none"
            if vm_val in ("restart", "none"):
                current_group = "".join(t.text or "" for t in cell0.findall(f".//{qn('w:t')}")).strip()
            if current_group == prod_comp_group:
                last_group_row = row_el

        if last_group_row is None:
            print(f"WARNING: Could not find product component group '{prod_comp_group}' in Table 4")
            continue

        new_row = clone_element(last_group_row)
        new_cells = new_row.findall(qn("w:tc"))
        # Col 0: vMerge continue
        cell0 = new_cells[0]
        tcpr = cell0.find(qn("w:tcPr"))
        if tcpr is None:
            tcpr = etree.SubElement(cell0, qn("w:tcPr"))
            cell0.insert(0, tcpr)
        vm = tcpr.find(qn("w:vMerge"))
        if vm is None:
            vm = etree.SubElement(tcpr, qn("w:vMerge"))
        if qn("w:val") in vm.attrib:
            del vm.attrib[qn("w:val")]
        for p in cell0.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
        # Col 1: remove vMerge if present
        if len(new_cells) > 1:
            col1_tcpr = new_cells[1].find(qn("w:tcPr"))
            if col1_tcpr is not None:
                col1_vm = col1_tcpr.find(qn("w:vMerge"))
                if col1_vm is not None:
                    col1_tcpr.remove(col1_vm)
            _set_cell_text(new_cells[1], plat["name"])
        if len(new_cells) > 2:
            _set_cell_text(new_cells[2], plat.get("sub_component", "NA"))
        if len(new_cells) > 3:
            _set_cell_text(new_cells[3], plat.get("tested_version", ""))
        if len(new_cells) > 4:
            _set_cell_text(new_cells[4], plat.get("function", ""))
        last_group_row.addnext(new_row)


def _apply_bold_underline(doc, bold_underline_models: list, table_indices: list):
    for ti in table_indices:
        if ti >= len(doc.tables):
            continue
        pc_tbl = doc.tables[ti]
        for row in pc_tbl.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            col1_text = cells[1].text.strip()
            if "Component Name" in col1_text or "NOTE" in col1_text or "Product Components" in cells[0].text:
                continue
            for p in cells[1].paragraphs:
                for r in p.runs:
                    r.bold = False
                    r.underline = False
                for r in p.runs:
                    run_text = r.text.strip()
                    should_bold = any(run_text == model for model in bold_underline_models)
                    r.bold = should_bold
                    r.underline = should_bold


def _apply_notes(doc, cfg: dict):
    """Apply TABLE3_NOTES, TABLE4_NOTES, TABLE5_NOTES if configured.

    Uses label-based lookup via find_notes_tables() — no hardcoded indices.
    """
    if cfg.get("table3_notes") is not None:
        apply_notes_by_label(doc, "Network Interfaces", cfg["table3_notes"])
    if cfg.get("table4_notes") is not None:
        apply_notes_by_label(doc, "Product Components", cfg["table4_notes"])
    if cfg.get("table5_notes") is not None:
        apply_notes_by_label(doc, "CR/FR ID", cfg["table5_notes"])


def _merge_tables(doc, body, src_idx: int, dst_idx: int, skip_rows: int):
    """Merge dst table rows into src table and remove dst table + intervening elements."""
    tbl_src = doc.tables[src_idx]._tbl
    tbl_dst = doc.tables[dst_idx]._tbl
    dst_rows = tbl_dst.findall(qn("w:tr"))
    for row in dst_rows[skip_rows:]:
        tbl_src.append(copy.deepcopy(row))
    found_src = False
    to_remove = []
    for el in list(body):
        if el is tbl_src:
            found_src = True
            continue
        if el is tbl_dst:
            to_remove.append(el)
            break
        if found_src:
            to_remove.append(el)
    for el in to_remove:
        body.remove(el)


def _add_tbl_header_to_table(doc, table_idx: int):
    tbl_merged = doc.tables[table_idx]._tbl
    tbl_rows = tbl_merged.findall(qn("w:tr"))
    if tbl_rows:
        trpr = tbl_rows[0].find(qn("w:trPr"))
        if trpr is None:
            trpr = etree.SubElement(tbl_rows[0], qn("w:trPr"))
            tbl_rows[0].insert(0, trpr)
        hdr_el = trpr.find(qn("w:tblHeader"))
        if hdr_el is None:
            hdr_el = etree.SubElement(trpr, qn("w:tblHeader"))
        hdr_el.set(qn("w:val"), "1")


def _merge_and_split_table4(doc, body, first_idx=3, second_idx=4):
    """Merge Table 4 HW/SW/FW Version ID (doc tables first_idx+second_idx), then split Product Components into its own table."""
    tbl4_first = doc.tables[first_idx]._tbl
    tbl4_second = doc.tables[second_idx]._tbl
    second_rows = tbl4_second.findall(qn("w:tr"))
    for row in second_rows[1:]:
        tbl4_first.append(copy.deepcopy(row))
    found_first = False
    to_remove = []
    for el in list(body):
        if el is tbl4_first:
            found_first = True
            continue
        if el is tbl4_second:
            to_remove.append(el)
            break
        if found_first:
            to_remove.append(el)
    for el in to_remove:
        body.remove(el)
    # Split at row 5 (Product Components)
    merged_rows = tbl4_first.findall(qn("w:tr"))
    new_tbl = copy.deepcopy(tbl4_first)
    for row in new_tbl.findall(qn("w:tr")):
        new_tbl.remove(row)
    for row in merged_rows[5:]:
        new_tbl.append(copy.deepcopy(row))
    for row in merged_rows[5:]:
        tbl4_first.remove(row)
    tbl4_first.addnext(new_tbl)
    # Fix boundary between Product ID and Product Components tables:
    # Product ID keeps its original sz=6 borders (intentional thicker weight).
    # Set bottom border of Product ID last row cells to nil so there's no
    # double-thick line where the two tables meet.
    prod_id_rows = tbl4_first.findall(qn("w:tr"))
    if prod_id_rows:
        last_row = prod_id_rows[-1]
        for cell in last_row.findall(qn("w:tc")):
            tcPr = cell.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = etree.SubElement(cell, qn("w:tcPr"))
                cell.insert(0, tcPr)
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
            bottom = tcBorders.find(qn("w:bottom"))
            if bottom is None:
                bottom = etree.SubElement(tcBorders, qn("w:bottom"))
            bottom.set(qn("w:val"), "nil")
    # Add tblLayout fixed to Product ID table so column widths don't auto-adjust
    prod_id_tblpr = tbl4_first.find(qn("w:tblPr"))
    if prod_id_tblpr is not None and prod_id_tblpr.find(qn("w:tblLayout")) is None:
        layout = etree.SubElement(prod_id_tblpr, qn("w:tblLayout"))
        layout.set(qn("w:type"), "fixed")
    # Fix tblPr on new Product Components table
    new_tblpr = new_tbl.find(qn("w:tblPr"))
    if new_tblpr is not None:
        inherited_borders = new_tblpr.find(qn("w:tblBorders"))
        if inherited_borders is not None:
            new_tblpr.remove(inherited_borders)
        if new_tblpr.find(qn("w:tblLayout")) is None:
            layout = etree.SubElement(new_tblpr, qn("w:tblLayout"))
            layout.set(qn("w:type"), "fixed")
    # Do NOT set tblHeader on Product Components table — continuous table, no header repeat on page breaks
    # The duplicate column header row from INITIAL Table 5 (Table 6) is removed in step 13a above.


# ---------------------------------------------------------------------------
# Configuration profiles for SBC / CTN2026003
# (These replace the hardcoded values in the one-off scripts.)
# New CTNs: add a new profile section or load from a config file.
# ---------------------------------------------------------------------------

SBC_CTN2026003_FA_PROFILES = {
    # DTR001 seed profile — first-production version for SBC/CTN2026003.
    # CONFIRM new_ver and testing dates before generating.
    # This profile is presented to the engineer with a y/N confirm prompt before use.
    # hw_list, component_list, and dtr_id_suffix are derived from updating_platforms_table
    # and new_platforms using the same helpers as the manual path — do NOT hardcode them.
}

def _build_sbc_ctn2026003_profiles() -> dict:
    """Build SBC/CTN2026003 FA profiles with auto-derived hw_list, component_list, dtr_id_suffix."""
    _dtr1_updating_table  = ["ASR", "ISR", "C8300", "C8200", "C8000v", "C8000V", "8300 Secure"]  # short match identifiers for Table 4
    _dtr1_updating_display = ["ASR 1006-X", "ISR 4461", "C8300 series", "C8200 series", "C8000v series"]  # human-readable for body sentence
    _dtr1_new_platforms = [
        {"group": "IWBC", "name": "C8300", "display_name": "C8300 series", "sub_component": "NA", "tested_version": "IOS XE 26.2", "function": "IWBC"},
        {"group": "SBC",  "name": "C8300", "display_name": "C8300 series", "sub_component": "NA", "tested_version": "IOS XE 26.2", "function": "SBC"},
        {"group": "IWG",  "name": "C8300", "display_name": "C8300 series", "sub_component": "NA", "tested_version": "IOS XE 26.2", "function": "IWG"},
    ]
    _dtr1_components = ["IWBC", "IWG", "SBC"]
    return {
        1: {
            "dtr_num": 1,
            "new_ver": "IOS XE 26.2",
            "old_ver": "IOS XE 17.18",
            "dtr_id_suffix": build_dtr_id_suffix(_dtr1_updating_display, _dtr1_new_platforms),
            "hw_list": build_hw_list(_dtr1_updating_display, _dtr1_new_platforms),
            "component_list": build_component_list(_dtr1_components),
            "testing_start_date": "05 Jun",
            "testing_start_year": "2026",
            "testing_end_date": "08 Jul",
            "testing_end_year": "2026",
            "updating_platforms_table": _dtr1_updating_table,
            "sustained_platforms": [],
            "sustained_by_group": [],
            "bold_underline_models": [
                "ASR 1006-X", "ISR 4461", "C8200-1N-4T", "C8200L-1N-4T",
                "C8300-1N1S-4T2X", "C8300-2N2S-4T2X", "C8300-1N1S-6T", "C8300-2N2S-6T",
                "C8000v", "8300 Secure Router", "ESXi Server Host", "Management Workstation",
            ],
            "similarity_text": "",
            "poam_text": "",
            "tdr_row": None,
            "tdr_updates_full": [],
            "new_platforms": _dtr1_new_platforms,
            "table3_notes": None,
            "table4_notes": None,
            "table5_notes": None,
        },
        # DTR002, DTR003, DTR004+ have no pre-baked profiles.
        # When newd prompts for DTR number, if no profile exists the runner
        # will prompt for all inputs manually. This is intentional — the
        # parameterized runner handles any DTR on demand.
    }

SBC_CTN2026003_FA_PROFILES = _build_sbc_ctn2026003_profiles()


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def _detect_old_ver(prod_cat: str, ctn: str, dtr_num: int) -> str:
    """Detect old_ver from the previous DTR's draft filename.

    Scans the Drafts/ folder for a file matching DTR{dtr_num-1} and extracts
    the version string from the filename. Returns empty string if not found.
    """
    import re
    prev_dtr_str = f"DTR{dtr_num - 1:03d}"
    drafts_dir = fa_dir(prod_cat, ctn) / "Drafts"
    pattern = re.compile(
        rf"Draft_{re.escape(ctn)} - {re.escape(prev_dtr_str)} - {re.escape(prod_cat)} - (.+?) - Cisco Functionality Attestation\.docx",
        re.IGNORECASE,
    )
    for f in drafts_dir.glob("*.docx"):
        m = pattern.match(f.name)
        if m:
            return m.group(1)
    return ""


def execute_cfg(cfg: dict) -> None:
    """Execute generation from a fully-built cfg dict — no prompts.

    Called by OpenCode after building cfg via question tool prompts.
    Also called by run() after its own prompting is complete.
    """
    prod_cat = cfg["prod_cat"]
    ctn      = cfg["ctn"]
    dtr_num  = cfg["dtr_num"]

    if dtr_num == 1:
        out_path = fa_draft_path(prod_cat, ctn, dtr_num, cfg["new_ver"])
        cfg["out_path"] = out_path
        generate_dtr001(cfg)
    else:
        prev_ver    = cfg["old_ver"]
        source_path = fa_draft_path(prod_cat, ctn, dtr_num - 1, prev_ver)
        out_path    = fa_draft_path(prod_cat, ctn, dtr_num, cfg["new_ver"])
        cfg["source_path"] = source_path
        cfg["out_path"]    = out_path
        generate_dtr_incremental(cfg)

    # Post-generation validation
    print("\nRunning validate_doc.py...")
    run_validate(out_path)

    # Draft Log
    import subprocess as _sp
    _git_user = _sp.run(["git", "config", "user.name"], capture_output=True, text=True,
                        cwd=str(BASE)).stdout.strip()
    engineer = cfg.get("engineer") or _git_user or "unknown"
    append_draft_log(
        engineer=engineer,
        action="Generated",
        ctn=ctn,
        doc_type="FA",
        dtr=f"DTR{dtr_num:03d}",
        version=cfg["new_ver"],
        reason="Via run_fa.py parameterized runner",
    )

    print("\nDone.")


def run(prod_cat: str, ctn: str, dtr_num: int) -> None:
    """Callable entry point — invoked by the newd prompt sequence or directly via main()."""
    # Load profile for this CTN if available, else require manual input
    profiles = {}
    if prod_cat == "SBC" and ctn == "CTN2026003":
        profiles = SBC_CTN2026003_FA_PROFILES

    if dtr_num in profiles:
        cfg = dict(profiles[dtr_num])
        print(f"\nLoaded profile for {prod_cat}/{ctn} DTR{dtr_num:03d}.")
        print("  (Platform config, components, and bold models pre-loaded.)\n")
        # Always prompt for version and dates — these change every DTR
        cfg["new_ver"] = input(f"New IOS XE version [{cfg.get('new_ver','')}]: ").strip() or cfg.get("new_ver", "")
        cfg["old_ver"] = _detect_old_ver(prod_cat, ctn, dtr_num) or cfg.get("old_ver", "")
        if cfg["old_ver"]:
            print(f"  Previous version: {cfg['old_ver']}")
        else:
            cfg["old_ver"] = input("  Previous IOS XE version (enter manually): ").strip()
        cfg["testing_start_date"] = input(f"Testing start date [{cfg.get('testing_start_date','')}]: ").strip() or cfg.get("testing_start_date", "")
        cfg["testing_start_year"] = input(f"Testing start year [{cfg.get('testing_start_year','')}]: ").strip() or cfg.get("testing_start_year", "")
        cfg["testing_end_date"]   = input(f"Testing end date [{cfg.get('testing_end_date','')}]: ").strip() or cfg.get("testing_end_date", "")
        cfg["testing_end_year"]   = input(f"Testing end year [{cfg.get('testing_end_year','')}]: ").strip() or cfg.get("testing_end_year", "")
    else:
        print(f"\nNo profile found for {prod_cat}/{ctn} DTR{dtr_num:03d}. Manual input required.")
        new_ver  = input("New IOS XE version (e.g. IOS XE 26.2): ").strip()

        # Detect old_ver from previous DTR draft filename — no prompt needed
        old_ver = _detect_old_ver(prod_cat, ctn, dtr_num)
        if old_ver:
            print(f"  Previous version detected from DTR{dtr_num-1:03d} draft: {old_ver}")
        else:
            old_ver = input("  Previous IOS XE version (could not detect — enter manually): ").strip()

        # Updating platforms — two lists:
        #   updating       : short identifiers matched against Component Name col in Table 4
        #   updating_display: human-readable names used in hw_list and dtr_id_suffix sentences
        print("\nUpdating platforms — Table 4 match identifiers (e.g. C8300, C8200, C8000v):")
        print("  One per line, blank to finish")
        updating = []
        while True:
            v = input("  Identifier (blank=done): ").strip()
            if not v:
                break
            updating.append(v)

        print("\nUpdating platforms — display names for DTR sentence (e.g. C8300 series, C8200 series):")
        print("  One per line, same order as above, blank to finish")
        updating_display = []
        while True:
            v = input("  Display name (blank=done): ").strip()
            if not v:
                break
            updating_display.append(v)

        # Sustained platforms — name + version pairs
        print("\nSustained platforms (keep current version, not updating):")
        print("  e.g. ASR 1006-X at IOS XE 17.18 — blank name to finish")
        sustained_platforms = []
        sustained_by_group  = []
        while True:
            sname = input("  Sustained platform name (blank=done): ").strip()
            if not sname:
                break
            sver  = input(f"    Sustained version for {sname}: ").strip()
            sustained_platforms.append((sname, sver))
            groups_raw = input(f"    Groups for {sname} (e.g. IWG,SBC): ").strip()
            for g in [g.strip() for g in groups_raw.split(",") if g.strip()]:
                sustained_by_group.append((sname.split()[0], g))

        # Bold+underline models
        print("\nBold+underline tested models (one per line, blank to finish):")
        bold_models = []
        while True:
            m = input("  Model (blank=done): ").strip()
            if not m:
                break
            bold_models.append(m)

        # New platforms
        print("\nNew platforms to add to Table 4 (blank name to finish):")
        new_platforms = []
        while True:
            pname = input("  New platform name (blank=done): ").strip()
            if not pname:
                break
            pdisplay = input(f"    Display name for DTR sentence (e.g. {pname} series): ").strip() or pname
            pgroups  = input(f"    Groups (e.g. IWBC,SBC,IWG): ").strip()
            psub     = input(f"    Sub-component (blank=NA): ").strip() or "NA"
            pfunc_default = pgroups.split(",")[0].strip() if pgroups else ""
            for g in [g.strip() for g in pgroups.split(",") if g.strip()]:
                new_platforms.append({
                    "group": g, "name": pname, "display_name": pdisplay,
                    "sub_component": psub, "tested_version": new_ver, "function": g,
                })

        # Product components covered in this DTR (drives component_list in body sentence)
        print("\nProduct Component(s) covered in this DTR:")
        print("  e.g. IWBC, SBC, IWG — one per line, blank to finish")
        components = []
        while True:
            c = input("  Component (blank=done): ").strip()
            if not c:
                break
            components.append(c)

        cfg = {
            "dtr_num": dtr_num,
            "new_ver": new_ver,
            "old_ver": old_ver,
            "dtr_id_suffix": build_dtr_id_suffix(updating_display, new_platforms),
            "hw_list": build_hw_list(updating_display, new_platforms),
            "component_list": build_component_list(components),
            "testing_start_date": input("Testing start date (e.g. 05 Jun): ").strip(),
            "testing_start_year": input("Testing start year (e.g. 2026): ").strip(),
            "testing_end_date": input("Testing end date (e.g. 08 Jul): ").strip(),
            "testing_end_year": input("Testing end year (e.g. 2026): ").strip(),
            "updating_platforms_table": updating,
            "sustained_platforms": sustained_platforms,
            "sustained_by_group": sustained_by_group,
            "bold_underline_models": bold_models,
            "similarity_text": "",
            "poam_text": "",
            "tdr_row": None,
            "tdr_updates_full": [],
            "new_platforms": new_platforms,
            "table3_notes": None,
            "table4_notes": None,
            "table5_notes": None,
        }

    # Prompt for Table 1 TDR (step 9) — applies to both profile and manual paths
    # Handles both: adding a new TDR row, and updating an existing TDR row in place.
    # If the entered TDR number matches an existing row, the generation step updates it;
    # if it does not match, a new row is appended. tdr_updates_full is driven from here.
    print("\nTable 1. Conditions — TDR update?")
    tdr_action = input("  Add or update a TDR row? [y/N]: ").strip().lower()
    if tdr_action == "y":
        tdr_num_str = input("    TDR Number (e.g. 2026003-1): ").strip()
        tdr_desc    = input("    TDR Description: ").strip()
        tdr_impact  = input("    Operational Impact (e.g. OPEN, CLOSED): ").strip()
        tdr_remarks = input("    Remarks: ").strip()
        # Store in tdr_row — generation checks existing rows first:
        # if TDR number matches an existing row it updates in place (via tdr_updates_full logic),
        # otherwise it appends a new row.
        cfg["tdr_row"] = {
            "num": tdr_num_str,
            "desc": tdr_desc,
            "impact": tdr_impact,
            "remarks": tdr_remarks,
        }
        # Also wire into tdr_updates_full so the update-in-place path is triggered
        cfg["tdr_updates_full"] = [(tdr_num_str, tdr_desc, tdr_impact, tdr_remarks)]
    else:
        cfg["tdr_row"] = None
        cfg["tdr_updates_full"] = []

    # Prompt for table notes (step 11) — applies to both profile and manual paths
    print("\nTable NOTES update (leave blank to skip):")
    t3 = input("  Table 3 (Network Interfaces) notes text [blank=skip]: ").strip()
    cfg["table3_notes"] = t3 if t3 else None
    t4 = input("  Table 4 (Product Components) notes text [blank=skip]: ").strip()
    cfg["table4_notes"] = t4 if t4 else None
    t5 = input("  Table 5 (CR/FR) notes text [blank=skip]: ").strip()
    cfg["table5_notes"] = t5 if t5 else None

    # Similarity statement — applies to both profile and manual paths
    sim = input('\nSimilarity statement? (blank=none, or: SBC CTN2026002 DTR2): ').strip()
    if sim:
        # Expect format: "<ProdCat> <CTN> <DTR>" — build canonical sentence
        parts = sim.split()
        if len(parts) == 3:
            sim_prodcat, sim_ctn, sim_dtr_raw = parts
            sim_dtr_num = int(''.join(filter(str.isdigit, sim_dtr_raw)))
            cfg["similarity_text"] = (
                f'Request certification through similarity based on '
                f'"{sim_prodcat} TN: {sim_ctn}, DTR{sim_dtr_num:02d}."'
            )
        else:
            cfg["similarity_text"] = sim  # use as-is if non-standard
    else:
        cfg["similarity_text"] = ""

    # Add paths to cfg
    cfg["prod_cat"] = prod_cat
    cfg["ctn"] = ctn

    import subprocess as _sp
    _git_user = _sp.run(["git", "config", "user.name"], capture_output=True, text=True,
                        cwd=str(BASE)).stdout.strip()
    cfg["engineer"] = input("\nEngineer username (for Draft Log): ").strip() or _git_user or "unknown"

    execute_cfg(cfg)


def main():
    """Interactive entry point — accepts optional CLI args: prod_cat ctn dtr_num.

    When called from the newd sequence, OpenCode passes all three as arguments
    so the engineer is dropped straight into the content prompts without re-entering
    selections already made.

    Usage:
        python3 run_fa.py                        # fully interactive
        python3 run_fa.py SBC CTN2026003 2       # newd handoff — skips header prompts
    """
    import sys
    args = sys.argv[1:]
    if len(args) == 3:
        prod_cat = args[0].strip()
        ctn      = args[1].strip()
        dtr_num  = int(args[2].strip())
        print(f"\n=== Functionality Attestation — {prod_cat} / {ctn} / DTR{dtr_num:03d} ===\n")
    else:
        print("\n=== Functionality Attestation — Parameterized Runner ===\n")
        prod_cat = input("Product Category (e.g. SBC): ").strip()
        ctn      = input("CTN (e.g. CTN2026003): ").strip()
        dtr_num  = int(input("DTR Number (e.g. 1): ").strip())
    run(prod_cat, ctn, dtr_num)


if __name__ == "__main__":
    main()
