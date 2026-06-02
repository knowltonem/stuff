#!/usr/bin/env python3
"""
run_tdr.py — Parameterized Test Discrepancy Report (TDR) runner.

Replaces gen_tdr26003_02.py and all future one-off TDR scripts.
Invoked by the newd prompt sequence after Product Category, CTN, and Document Type
have been selected. Prompts for all TDR-specific inputs, then generates the draft.

Usage (direct):
    python3 _Tools/run_tdr.py

Or triggered automatically by newd after skill_base.md routes to this runner.

TDR numbers are standalone per-finding documents — they do not chain per DTR.
Each TDR is generated from the Template file; product description is reused
from the most recent existing TDR (or the DTR001 example if none exist).
"""

import copy
import re
import sys
from pathlib import Path

# runner_core.py lives in the same _Tools/ directory
sys.path.insert(0, str(Path(__file__).parent))
from runner_core import (
    qn, run_validate, append_draft_log, BASE, tdr_dir, tdr_template_path,
    tdr_draft_path, get_git_username, get_para_text,
    set_run_text, set_para_single_run_text, set_label_value,
    find_para_by_label, find_para_after_label,
    set_run_font_size, set_all_runs_font_size, strip_highlighting,
)

from lxml import etree
from docx import Document


# ---------------------------------------------------------------------------
# Path helpers (TDR-specific — not in runner_core)
# ---------------------------------------------------------------------------

def tdr_existing_path(prod_cat: str, ctn: str) -> Path:
    """Return the most recent non-draft TDR for product description reuse."""
    d = tdr_dir(prod_cat, ctn)
    # Prefer files directly in Examples & Templates (not drafts)
    candidates = [
        p for p in sorted(d.glob("Examples & Templates/*.docx"))
        if not p.name.startswith("Template_") and not p.name.startswith("Draft_")
    ]
    if candidates:
        return candidates[-1]
    # Fall back to any non-template, non-draft docx anywhere under TDR dir
    candidates = [
        p for p in sorted(d.rglob("*.docx"))
        if not p.name.startswith("Template_") and not p.name.startswith("Draft_")
    ]
    if candidates:
        return candidates[-1]
    return None


def _next_tdr_sequence(prod_cat: str, ctn: str) -> int:
    """Auto-detect the next TDR sequence number by scanning Examples & Templates and Drafts."""
    d = tdr_dir(prod_cat, ctn)
    ctn_num_match = re.search(r"(\d+)$", ctn)  # extract trailing numeric suffix (e.g. "26003" from "CTN2026003")
    if not ctn_num_match:
        raise ValueError(f"Cannot parse numeric suffix from CTN: {ctn!r}")
    ctn_num = ctn_num_match.group(1)
    max_seq = 0
    pattern = re.compile(rf"TDR{ctn_num}-(\d+)", re.IGNORECASE)
    for p in d.rglob("*.docx"):
        m = pattern.search(p.name)
        if m:
            seq = int(m.group(1))
            if seq > max_seq:
                max_seq = seq
    return max_seq + 1



# ---------------------------------------------------------------------------
# Generation
# ---------------------------------------------------------------------------

def generate_tdr(cfg: dict):
    """Generate a TDR draft from the template."""
    prod_cat  = cfg["prod_cat"]
    ctn       = cfg["ctn"]
    dtr_num   = cfg["dtr_num"]
    tdr_number = cfg["tdr_number"]
    version   = cfg["version"]
    ver_num   = version.replace("IOS XE ", "")
    prod_name = cfg["product_name"]
    sw_ver    = f"with Software Release IOS XE {ver_num}"
    out_path  = cfg["out_path"]

    template_path = tdr_template_path(prod_cat, ctn)
    if not template_path.exists():
        raise FileNotFoundError(f"TDR template not found: {template_path}")

    existing_path = tdr_existing_path(prod_cat, ctn)

    doc = Document(str(template_path))
    body = doc.element.body
    paras = doc.paragraphs

    # -- Product description (reuse from existing TDR via content scan, or blank placeholder)
    # Use find_para_after_label() to locate the paragraph immediately after the
    # "Description of product:" label — robust to template layout changes.
    # Avoids the fragile paragraphs[17] absolute index used previously.
    product_desc = ""
    if existing_path and existing_path.exists():
        try:
            existing_doc = Document(str(existing_path))
            existing_elements = list(existing_doc.element.body)
            desc_para = find_para_after_label(existing_elements, "Description of product:")
            if desc_para is not None:
                product_desc = get_para_text(desc_para).strip()
            print(f"Reusing product description from: {existing_path.name}")
        except Exception as e:
            print(f"WARNING: Could not read existing TDR for product description: {e}")
    if not product_desc:
        product_desc = f"[Product description — copy from existing TDR or fill in manually]"
        print("WARNING: No existing TDR found; product description placeholder inserted.")

    # -------------------------------------------------------------------------
    # P00: Fix logo paragraph — remove oversized font
    # -------------------------------------------------------------------------
    for r in paras[0]._element.findall(f".//{qn('w:r')}"):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            for tag in [qn("w:sz"), qn("w:szCs")]:
                el = rPr.find(tag)
                if el is not None:
                    rPr.remove(el)
    print("Fixed logo paragraph font size")

    # -------------------------------------------------------------------------
    # P02: Title — fix font size to 12pt
    # -------------------------------------------------------------------------
    set_all_runs_font_size(paras[2]._element, 12)
    print("Fixed title font size to 12pt")

    # -------------------------------------------------------------------------
    # P04: Product Name — replace placeholder, fix font size
    # -------------------------------------------------------------------------
    runs_p4 = paras[4]._element.findall(f".//{qn('w:r')}")
    if runs_p4:
        set_run_text(runs_p4[0], prod_name)
        for r in runs_p4[1:]:
            r.getparent().remove(r)
    set_all_runs_font_size(paras[4]._element, 12)

    # -------------------------------------------------------------------------
    # P05: Insert Software Version as new paragraph after P04
    # -------------------------------------------------------------------------
    ver_para = copy.deepcopy(paras[4]._element)
    runs_ver = ver_para.findall(f".//{qn('w:r')}")
    if runs_ver:
        set_run_text(runs_ver[0], sw_ver)
        for r in runs_ver[1:]:
            r.getparent().remove(r)
    set_all_runs_font_size(ver_para, 12)
    paras[4]._element.addnext(ver_para)
    print(f"Set product name: {prod_name}")
    print(f"Inserted software version: {sw_ver}")

    # Re-read body elements after insertion
    elements = list(body)

    # -------------------------------------------------------------------------
    # TDR Number
    # -------------------------------------------------------------------------
    tdr_el = find_para_by_label(elements, "TDR Number:")
    if tdr_el is not None:
        set_label_value(tdr_el, tdr_number)
        print(f"Set TDR Number: {tdr_number}")

    # -------------------------------------------------------------------------
    # Remove TDR numbering note (*Note: TDR number is...)
    # -------------------------------------------------------------------------
    note_el = find_para_by_label(elements, "*Note: TDR number is")
    if note_el is not None:
        note_el.getparent().remove(note_el)
        elements = list(body)
        print("Removed TDR numbering note")

    # -------------------------------------------------------------------------
    # Requirement
    # -------------------------------------------------------------------------
    req_el = find_para_by_label(elements, "Requirement:")
    if req_el is not None:
        set_label_value(req_el, cfg["requirement"])
        print(f"Set Requirement: {cfg['requirement']}")

    # -------------------------------------------------------------------------
    # Criticality
    # -------------------------------------------------------------------------
    crit_el = find_para_by_label(elements, "Criticality:")
    if crit_el is not None:
        set_label_value(crit_el, cfg["criticality"])
        print(f"Set Criticality: {cfg['criticality']}")

    # -------------------------------------------------------------------------
    # Finding
    # -------------------------------------------------------------------------
    finding_el = find_para_by_label(elements, "Finding:")
    if finding_el is not None:
        set_label_value(finding_el, cfg["finding"])
        print(f"Set Finding: {cfg['finding']}")

    # -------------------------------------------------------------------------
    # Description of product (P17)
    # -------------------------------------------------------------------------
    desc_el = find_para_after_label(elements, "Description of product:")
    if desc_el is not None:
        set_para_single_run_text(desc_el, product_desc)
        print("Set Description of product")

    # -------------------------------------------------------------------------
    # Problem Description
    # -------------------------------------------------------------------------
    prob_el = find_para_after_label(elements, "Problem Description:")
    if prob_el is not None:
        steps = cfg["problem_description"]
        set_para_single_run_text(prob_el, steps[0] if steps else "")
        prev = prob_el
        for step in steps[1:]:
            new_p = copy.deepcopy(prob_el)
            set_para_single_run_text(new_p, step)
            prev.addnext(new_p)
            prev = new_p
        elements = list(body)
        print(f"Set Problem Description ({len(steps)} line(s))")

    # -------------------------------------------------------------------------
    # Test Scenario
    # -------------------------------------------------------------------------
    scenario_el = find_para_after_label(elements, "Test Scenario:")
    if scenario_el is not None:
        steps = cfg["test_scenario_steps"]
        set_para_single_run_text(scenario_el, steps[0] if steps else "")
        prev = scenario_el
        for step in steps[1:]:
            new_p = copy.deepcopy(scenario_el)
            set_para_single_run_text(new_p, step)
            prev.addnext(new_p)
            prev = new_p
        elements = list(body)
        print(f"Set Test Scenario ({len(steps)} step(s))")

    # -------------------------------------------------------------------------
    # Test Result
    # -------------------------------------------------------------------------
    result_el = find_para_after_label(elements, "Test Result:")
    if result_el is not None:
        steps = cfg["test_result_steps"]
        set_para_single_run_text(result_el, steps[0] if steps else "")
        prev = result_el
        for step in steps[1:]:
            new_p = copy.deepcopy(result_el)
            set_para_single_run_text(new_p, step)
            prev.addnext(new_p)
            prev = new_p
        elements = list(body)
        print(f"Set Test Result ({len(steps)} step(s))")

    # -------------------------------------------------------------------------
    # Note (optional)
    # -------------------------------------------------------------------------
    if cfg.get("note"):
        note_field_el = find_para_by_label(elements, "Note:")
        if note_field_el is not None:
            set_label_value(note_field_el, cfg["note"])
            print(f"Set Note: {cfg['note']}")

    # -------------------------------------------------------------------------
    # Expected Behavior
    # -------------------------------------------------------------------------
    expected_el = find_para_after_label(elements, "Expected behavior:")
    if expected_el is not None:
        set_para_single_run_text(expected_el, cfg["expected_behavior"])
        print(f"Set Expected behavior: {cfg['expected_behavior']}")

    # -------------------------------------------------------------------------
    # Components Affected — label pluralizes per skill_tdr.md spec
    # -------------------------------------------------------------------------
    comp_value = cfg["components_affected_value"]
    comp_count = cfg.get("components_affected_count", 1)
    if comp_count == 1:
        comp_label = "Component Affected:"
    else:
        comp_label = f"Components Affected (x{comp_count}):"

    comp_el = find_para_by_label(elements, "Component")
    if comp_el is not None:
        # Rewrite label+value manually since label text changes
        runs = comp_el.findall(f".//{qn('w:r')}")
        if runs:
            # Set run 0 to bold label
            set_run_text(runs[0], f"{comp_label}  ")
            rPr = runs[0].find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.SubElement(runs[0], qn("w:rPr"))
                runs[0].insert(0, rPr)
            if rPr.find(qn("w:b")) is None:
                etree.SubElement(rPr, qn("w:b"))
            # Remove remaining runs
            for r in runs[1:]:
                r.getparent().remove(r)
            # New non-bold value run
            value_run = copy.deepcopy(runs[0])
            val_rPr = value_run.find(qn("w:rPr"))
            if val_rPr is not None:
                val_b = val_rPr.find(qn("w:b"))
                if val_b is not None:
                    val_rPr.remove(val_b)
                b_off = etree.SubElement(val_rPr, qn("w:b"))
                b_off.set(qn("w:val"), "0")
            set_run_text(value_run, comp_value)
            runs[0].addnext(value_run)
            print(f"Set {comp_label} {comp_value}")

    # -------------------------------------------------------------------------
    # Strip all highlighting and shading
    # -------------------------------------------------------------------------
    strip_highlighting(body)
    print("Stripped all highlighting and shading")

    # -------------------------------------------------------------------------
    # Save
    # -------------------------------------------------------------------------
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")


# ---------------------------------------------------------------------------
# Product Name map — extend as new product categories are onboarded
# ---------------------------------------------------------------------------

PRODUCT_NAME_MAP = {
    "SBC": "Cisco GCT DP Session Border Controller (SBC)",
    # ESC, SS: add here when example docs are placed
}

PROD_CAT_ABBR_MAP = {
    "SBC": "SBC",
    # ESC: "ESC", SS: "SS" — add when onboarded
}


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def run(prod_cat: str, ctn: str, dtr_num: int) -> None:
    """Callable entry point — invoked by the newd prompt sequence or directly via main()."""
    prod_name = PRODUCT_NAME_MAP.get(prod_cat, f"[Product name for {prod_cat}]")
    prod_cat_abbr = PROD_CAT_ABBR_MAP.get(prod_cat, prod_cat)

    # IOS XE version
    version_raw = input("IOS XE version (e.g. 26.2 or IOS XE 26.2): ").strip()
    if not version_raw.startswith("IOS XE "):
        version_raw = f"IOS XE {version_raw}"

    # TDR number — auto-suggest, allow override
    next_seq = _next_tdr_sequence(prod_cat, ctn)
    ctn_suffix = re.search(r"\d+$", ctn).group() if re.search(r"\d+$", ctn) else ctn  # trailing digits only, e.g. "2026003" from "CTN2026003"
    suggested_tdr = f"{ctn_suffix}-{next_seq:02d}"
    tdr_input = input(f"TDR number [suggested: {suggested_tdr}] (press Enter to accept): ").strip()
    tdr_number = tdr_input if tdr_input else suggested_tdr

    # Requirement
    req_input = input(
        "UCR section(s) (e.g. Section 2.7.1 — separate multiple with comma): "
    ).strip()
    requirement = f"UCR 2013 Change 2, {req_input}"

    # Criticality
    criticality = input("Criticality (e.g. SCM-012040 Required): ").strip()

    # Finding
    finding = input("Finding summary (short description): ").strip()

    # Problem Description — multi-line (one line per entry, blank to stop)
    print("Problem description — enter lines one at a time; blank line to finish:")
    problem_lines = []
    while True:
        line = input("  > ").strip()
        if not line:
            break
        problem_lines.append(line)
    if not problem_lines:
        problem_lines = ["[Problem description]"]

    # Test Scenario Steps
    print("Test scenario steps — enter steps one at a time; blank line to finish:")
    scenario_steps = []
    while True:
        line = input("  > ").strip()
        if not line:
            break
        scenario_steps.append(line)
    if not scenario_steps:
        scenario_steps = ["[Test scenario step]"]

    # Test Diagram — prompt but just note; image insertion not automated
    print("\nTest diagram: Image insertion is not automated.")
    print("Options: (1) Add image manually after generation  (2) Note file path to add later")
    input("Press Enter to continue (diagram placeholder will remain in document): ")

    # Test Result Steps
    print("Test result steps — enter steps one at a time; blank line to finish:")
    result_steps = []
    while True:
        line = input("  > ").strip()
        if not line:
            break
        result_steps.append(line)
    if not result_steps:
        result_steps = ["[Test result step]"]

    # Note (optional)
    note = input("Optional note (press Enter to skip): ").strip()

    # Expected Behavior
    expected_behavior = input("Expected behavior: ").strip()

    # Components Affected — multi-select + classify
    print("\nComponents Affected — enter components one at a time; blank to finish.")
    print("After all components are entered, you will classify each as IWBC, SBC, or Both.")
    components = []
    while True:
        c = input("  Component (e.g. c8200, ASR-1006): ").strip()
        if not c:
            break
        components.append(c)
    if not components:
        components = ["[Component]"]

    comp_entries = []
    for comp in components:
        cls = input(f"  Classification for {comp} [IWBC / SBC / Both]: ").strip()
        if cls.lower() in ("both", "iwbc, sbc", "iwbc/sbc"):
            comp_entries.append(f"Cisco {comp} (IWBC, SBC)")
        else:
            comp_entries.append(f"Cisco {comp} ({cls.upper()})")

    comp_value = ", ".join(comp_entries)
    comp_count = len(comp_entries)

    # Build output path
    out_path = tdr_draft_path(prod_cat, ctn, dtr_num, tdr_number, prod_cat_abbr, version_raw)

    cfg = {
        "prod_cat": prod_cat,
        "ctn": ctn,
        "dtr_num": dtr_num,
        "tdr_number": tdr_number,
        "version": version_raw,
        "product_name": prod_name,
        "requirement": requirement,
        "criticality": criticality,
        "finding": finding,
        "problem_description": problem_lines,
        "test_scenario_steps": scenario_steps,
        "test_result_steps": result_steps,
        "note": note,
        "expected_behavior": expected_behavior,
        "components_affected_value": comp_value,
        "components_affected_count": comp_count,
        "out_path": out_path,
    }

    print(f"\nGenerating TDR: {out_path.name}")
    generate_tdr(cfg)

    # Post-generation validation
    print("\nRunning validate_doc.py...")
    run_validate(out_path)

    # Draft Log
    engineer = input("\nEngineer username (for Draft Log): ").strip() or get_git_username() or "unknown"
    append_draft_log(
        engineer=engineer,
        action="Generated",
        ctn=ctn,
        doc_type="TDR",
        dtr=f"DTR{dtr_num:03d}",
        version=version_raw,
        reason=f"Via run_tdr.py parameterized runner — TDR{tdr_number}",
    )

    print("\nDone.")


def main():
    """Interactive entry point — called by newd after Product Category / CTN / Doc Type selected."""
    print("\n=== Test Discrepancy Report (TDR) — Parameterized Runner ===\n")
    prod_cat = input("Product Category (e.g. SBC): ").strip()
    ctn = input("CTN (e.g. CTN2026003): ").strip()
    dtr_num = int(input("DTR number when finding was discovered (e.g. 1): ").strip())
    run(prod_cat, ctn, dtr_num)


if __name__ == "__main__":
    main()
