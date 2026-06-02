#!/usr/bin/env python3
"""
validate_doc.py — ICR Automation post-generation document linter.

Usage:
    python3 _Tools/validate_doc.py <path_to_docx>
    python3 _Tools/validate_doc.py --all       # validate all drafts in repo

Checks:
    1. File is a valid ZIP / well-formed Word document
    2. XML document.xml is well-formed
    3. CTN appears in document body
    4. No unreplaced placeholder text ([[ ]], {{ }}, [PLACEHOLDER])
    5. Revision history table exists and last row is populated
    6. At least one DTR heading present (Desktop Review (DTR))
    7. All tables have more than one row (not empty)
    8. Required headings present for document type
    9. Hyperlinks resolve (rel targets are not empty)
   10. Core properties set (author / subject)
"""

import sys
import os
import re
import zipfile
import argparse
from pathlib import Path
from datetime import datetime

# Named constants — update here if signatory changes (must also update _Skills/skill_icr_memo.md)
# If this value changes, also update _Skills/skill_icr_memo.md → Signature Block section.
# QAC will flag any occurrence of ICR_MEMO_SIGNATORY in validate_doc.py not matching skill_icr_memo.md.
ICR_MEMO_SIGNATORY = "Robbie Horgan"

try:
    from docx import Document
    from lxml import etree
except ImportError:
    print("ERROR: Missing dependencies. Run: pip install python-docx lxml")
    sys.exit(1)

# ── Constants ─────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent.parent

REQUIRED_HEADINGS = {
    "System Description": [
        "Desktop Review (DTR)",
        "DTR Detailed Component Information",
    ],
    "Military Unique Deployment Guide": [
        "Conditions of Fielding",
        "Configuration Checklist",
    ],
    "Test Discrepancy Report": [
        "TDR Number:",
        "Finding:",
        "Problem Description:",
        "Expected behavior:",
    ],
}

PLACEHOLDER_PATTERNS = [
    r'\[\[.*?\]\]',       # [[PLACEHOLDER]]
    r'\{\{.*?\}\}',       # {{placeholder}}
    r'\[YOUR ',           # [YOUR NAME]
    r'\[INSERT ',         # [INSERT VALUE]
    r'\[TBD\]',
    r'\[TBD ',
]

REQUIRED_REVISION_COLS = 4  # Version | Date | Description | Editor

# Document types that skip generic table/heading checks
TABLE_EXEMPT_TYPES = {"ICR Summary Memorandum"}

# ── Helpers ───────────────────────────────────────────────────────────────────

def parse_filename(path):
    """Extract CTN and doc type from filename."""
    name = Path(path).stem
    ctn_match = re.search(r'(CTN\d+)', name, re.IGNORECASE)
    ctn = ctn_match.group(1) if ctn_match else None

    doc_type = None
    if "System Description" in name:
        doc_type = "System Description"
    elif "CSR" in name:
        doc_type = "Cybersecurity Summary Report"
    elif "Functionality Attestation" in name or re.search(r"\bFA\b", name):
        doc_type = "Functionality Attestation"
    elif "ICR Memo" in name:
        doc_type = "ICR Summary Memorandum"
    elif "LoC" in name:
        doc_type = "Letter Of Compliance"
    elif "MUDG" in name or "Military Unique Deployment Guide" in name:
        doc_type = "Military Unique Deployment Guide"
    elif "POA" in name:
        doc_type = "Plan of Action & Milestone"
    elif "System Diagram" in name:
        doc_type = "System Diagram"
    elif "TDR" in name:
        doc_type = "Test Discrepancy Report"

    return ctn, doc_type

def all_drafts():
    """Find all Draft_*.docx files in the repo."""
    return list(BASE_DIR.glob("Product Category/**/Drafts/Draft_*.docx"))

# ── Checks ────────────────────────────────────────────────────────────────────

def check_valid_zip(path, errors):
    try:
        with zipfile.ZipFile(path) as z:
            z.namelist()
        return True
    except Exception as e:
        errors.append(f"[FAIL] Not a valid ZIP/docx file: {e}")
        return False

def check_xml_wellformed(path, errors):
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml")
            etree.fromstring(xml)
        return True
    except Exception as e:
        errors.append(f"[FAIL] Malformed document.xml: {e}")
        return False

def check_ctn_present(doc, ctn, errors):
    if ctn is None:
        errors.append("[WARN] Could not parse CTN from filename — skipping CTN presence check")
        return
    # CTN may appear in paragraphs or table cells
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    combined = full_text + "\n" + table_text
    if ctn not in combined:
        errors.append(f"[WARN] CTN '{ctn}' not found in document body or tables — verify it appears in header/title page")

def check_no_placeholders(doc, errors):
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for pattern in PLACEHOLDER_PATTERNS:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            errors.append(f"[FAIL] Unreplaced placeholder text found: {matches[:3]}")

def check_revision_history(doc, errors, required_cols=None):
    """Check revision history table (Table 0) for required columns and populated cells.

    required_cols: override column count (default REQUIRED_REVISION_COLS=4).
    CSR uses 3 cols (Version / Date / Change Description) — pass required_cols=3.
    """
    if required_cols is None:
        required_cols = REQUIRED_REVISION_COLS
    if not doc.tables:
        errors.append("[FAIL] No tables found — revision history table missing")
        return
    rev_table = doc.tables[0]
    if len(rev_table.rows) < 2:
        errors.append("[FAIL] Revision history table has no data rows")
        return
    last_row = rev_table.rows[-1]
    cells = [c.text.strip() for c in last_row.cells]
    if len(cells) < required_cols:
        errors.append(f"[FAIL] Revision history last row has {len(cells)} cells, expected {required_cols}")
        return
    if not cells[0]:
        errors.append("[FAIL] Revision history last row — Version cell is empty")
    if not cells[1]:
        errors.append("[FAIL] Revision history last row — Date cell is empty")
    # Only check Editor cell (index 3) when 4-col revision history is expected
    if required_cols >= 4 and not cells[3]:
        errors.append("[FAIL] Revision history last row — Editor cell is empty")

def check_dtr_heading(doc, errors):
    headings = [p.text for p in doc.paragraphs if "Desktop Review" in p.text]
    if not headings:
        errors.append("[FAIL] No 'Desktop Review (DTR)' heading found in document")

def check_tables_not_empty(doc, errors):
    for i, table in enumerate(doc.tables):
        if len(table.rows) <= 1:
            errors.append(f"[WARN] Table {i} has only {len(table.rows)} row(s) — may be empty")

def check_required_headings(doc, doc_type, errors):
    if doc_type not in REQUIRED_HEADINGS:
        return
    all_text = [p.text for p in doc.paragraphs]
    for required in REQUIRED_HEADINGS[doc_type]:
        if not any(required in t for t in all_text):
            errors.append(f"[FAIL] Required heading not found: '{required}'")

def check_hyperlinks(doc, errors):
    rels = doc.part.rels
    hyperlinks = {r: rels[r].target_ref for r in rels if "hyperlink" in rels[r].reltype}
    for rid, url in hyperlinks.items():
        if not url or url.strip() == "":
            errors.append(f"[FAIL] Hyperlink relationship '{rid}' has empty target URL")

def check_core_properties(doc, errors):
    props = doc.core_properties
    if not props.author:
        errors.append("[WARN] Document core property 'author' is not set")

# ── ICR Memo-specific checks ─────────────────────────────────────────────────

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _get_para_xml(para):
    """Return the lxml element for a paragraph."""
    return para._element

def check_icr_memo_date(doc, errors):
    """Check that P1 (date paragraph) is right-aligned with 194-twip right indent."""
    if len(doc.paragraphs) < 2:
        errors.append("[FAIL] ICR Memo has fewer than 2 paragraphs — cannot locate date")
        return
    # Date is typically the first non-empty text paragraph (P0 is image, P1 is date)
    date_para = None
    for p in doc.paragraphs[:5]:
        text = p.text.strip()
        if re.match(r'\d{1,2}\s+\w+\s+\d{4}', text):
            date_para = p
            break
    if date_para is None:
        errors.append("[FAIL] ICR Memo date paragraph not found (expected 'DD Month YYYY' in first 5 paragraphs)")
        return

    el = _get_para_xml(date_para)
    pPr = el.find(f'{{{WNS}}}pPr')
    if pPr is None:
        errors.append("[FAIL] ICR Memo date paragraph has no formatting properties")
        return

    # Check right alignment
    jc = pPr.find(f'{{{WNS}}}jc')
    jc_val = jc.get(f'{{{WNS}}}val') if jc is not None else None
    if jc_val != 'right':
        errors.append(f"[FAIL] ICR Memo date alignment is '{jc_val}' — expected 'right'")

    # Check right indent (194 twips)
    ind = pPr.find(f'{{{WNS}}}ind')
    if ind is not None:
        right_ind = ind.get(f'{{{WNS}}}right', '0')
        if right_ind != '194':
            errors.append(f"[WARN] ICR Memo date right indent is {right_ind} twips — expected 194")
    else:
        errors.append("[FAIL] ICR Memo date paragraph has no indentation — expected right indent of 194 twips")

def check_icr_memo_subject(doc, errors):
    """Check SUBJECT line exists and has no version."""
    subject_para = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("SUBJECT"):
            subject_para = p
            break
    if subject_para is None:
        errors.append("[FAIL] ICR Memo SUBJECT line not found")
        return

    text = subject_para.text
    # Check for version — should NOT be present
    if re.search(r'Software\s+Rel(?:ease)?.*?IOS\s+XE\s+[\d.]+', text):
        errors.append("[FAIL] ICR Memo SUBJECT line contains a version — should match INITIAL format (no version)")

def check_icr_memo_approval_version(doc, errors):
    """Check that the approval summary paragraph contains a version."""
    for p in doc.paragraphs:
        if "completed approval" in p.text or "has passed CS scanning" in p.text:
            if re.search(r'IOS\s+XE\s+[\d.]+', p.text):
                return  # version found — good
            errors.append("[FAIL] ICR Memo approval summary paragraph missing IOS XE version")
            return
    errors.append("[WARN] ICR Memo approval summary paragraph not found")

def check_icr_memo_dtr_paragraphs(doc, errors):
    """Check DTR approval paragraphs exist and have spacers between them."""
    dtr_indices = []
    for i, p in enumerate(doc.paragraphs):
        if re.search(r'DTR\s*#\d{3}', p.text):
            dtr_indices.append(i)

    # Also check for INITIAL approval paragraph (different format)
    for i, p in enumerate(doc.paragraphs):
        if 'initial request was approved' in p.text.lower() and i not in dtr_indices:
            dtr_indices.append(i)
    dtr_indices.sort()

    if not dtr_indices:
        errors.append("[FAIL] No DTR approval paragraphs found in ICR Memo")
        return

    # Check spacers between consecutive DTR paragraphs
    for j in range(len(dtr_indices) - 1):
        idx_current = dtr_indices[j]
        idx_next = dtr_indices[j + 1]
        # There should be at least one empty paragraph between them
        has_spacer = False
        for k in range(idx_current + 1, idx_next):
            if doc.paragraphs[k].text.strip() == '':
                has_spacer = True
                break
        if not has_spacer:
            errors.append(f"[WARN] No spacer paragraph between DTR paragraphs at P{idx_current} and P{idx_next}")

def check_icr_memo_signature(doc, errors):
    """Check signature block is present."""
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if ICR_MEMO_SIGNATORY not in full_text:
        errors.append(f"[FAIL] ICR Memo signature block missing — '{ICR_MEMO_SIGNATORY}' not found")

# ── FA-specific checks ───────────────────────────────────────────────────────

def check_fa_test_details(doc, errors):
    """Check FA Test Details section has DTR paragraphs and testing dates."""
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Check Test Details paragraph exists with testing dates
    test_details_found = False
    for p in doc.paragraphs:
        if "Test Details" in p.text and "functionality testing from" in p.text:
            test_details_found = True
            # Check dates are not placeholder/default
            if "01 February" in p.text and "28 February" in p.text:
                errors.append("[WARN] FA Test Details dates may be INITIAL defaults (01 Feb - 28 Feb) — verify they were updated")
            break
    if not test_details_found:
        errors.append("[FAIL] FA Test Details paragraph not found")

    # Check at least one DTR body paragraph exists
    dtr_paras = [p for p in doc.paragraphs if re.search(r'DTR\d+ was requested', p.text)]
    if not dtr_paras:
        errors.append("[FAIL] No DTR body paragraphs found in FA Test Details section")

    # Check Functionality Status paragraph exists
    func_status = [p for p in doc.paragraphs if "Functionality Status" in p.text]
    if not func_status:
        errors.append("[FAIL] FA 'Functionality Status' paragraph not found")

    # IWBC is the correct col 0 name for Product Components in DTR001+ — no check needed

# ── Main validator ────────────────────────────────────────────────────────────

def validate(path):
    path = Path(path)
    errors = []
    passed = []

    print(f"\nValidating: {path.name}")
    print("─" * 60)

    ctn, doc_type = parse_filename(path)
    print(f"  CTN: {ctn or 'unknown'}  |  Doc Type: {doc_type or 'unknown'}")

    # Check 1 — valid ZIP
    if not check_valid_zip(path, errors):
        print(f"  [FAIL] Not a valid docx file — aborting further checks")
        return False, errors

    # Check 2 — XML well-formed
    if not check_xml_wellformed(path, errors):
        print(f"  [FAIL] Malformed XML — aborting further checks")
        return False, errors

    # Load document
    doc = Document(path)

    # Checks 3–10 (route by doc type)
    # FA and ICR Memo have CTN in header/title, not body — skip CTN body check for them
    if doc_type not in ("Functionality Attestation", "ICR Summary Memorandum", "Military Unique Deployment Guide"):
        check_ctn_present(doc, ctn, errors)
    check_no_placeholders(doc, errors)

    if doc_type == "ICR Summary Memorandum":
        # ICR Memo-specific checks (no tables, no DTR headings)
        check_icr_memo_date(doc, errors)
        check_icr_memo_subject(doc, errors)
        check_icr_memo_approval_version(doc, errors)
        check_icr_memo_dtr_paragraphs(doc, errors)
        check_icr_memo_signature(doc, errors)
    elif doc_type == "Functionality Attestation":
        # FA has no revision history table and no DTR heading —
        # DTR updates are body paragraphs in Test Details section
        check_tables_not_empty(doc, errors)
        check_fa_test_details(doc, errors)
    elif doc_type == "Military Unique Deployment Guide":
        # MUDG has a revision history table but no DTR heading paragraphs
        check_revision_history(doc, errors)
        check_tables_not_empty(doc, errors)
        check_required_headings(doc, doc_type, errors)
    elif doc_type == "Test Discrepancy Report":
        # TDR is paragraph-only — no tables, no revision history, no DTR heading
        check_required_headings(doc, doc_type, errors)
    elif doc_type == "Cybersecurity Summary Report":
        # CSR revision history is 3 cols (Version / Date / Change Description) — no Editor col
        check_revision_history(doc, errors, required_cols=3)
        check_dtr_heading(doc, errors)
        check_tables_not_empty(doc, errors)
        check_required_headings(doc, doc_type, errors)
    else:
        # Generic checks for table-based documents (System Description, LoC, etc.)
        check_revision_history(doc, errors)
        check_dtr_heading(doc, errors)
        check_tables_not_empty(doc, errors)
        check_required_headings(doc, doc_type, errors)

    check_hyperlinks(doc, errors)
    check_core_properties(doc, errors)

    # Report
    fails  = [e for e in errors if e.startswith("[FAIL]")]
    warns  = [e for e in errors if e.startswith("[WARN]")]

    for msg in fails:  print(f"  {msg}")
    for msg in warns:  print(f"  {msg}")

    if not fails and not warns:
        print("  [PASS] All checks passed")
    elif not fails:
        print(f"  [PASS] No failures — {len(warns)} warning(s)")
    else:
        print(f"  [FAIL] {len(fails)} failure(s), {len(warns)} warning(s)")

    return len(fails) == 0, errors

# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="ICR document validator")
    parser.add_argument("path", nargs="?", help="Path to .docx file to validate")
    parser.add_argument("--all", action="store_true", help="Validate all Draft_*.docx files in repo")
    args = parser.parse_args()

    if args.all:
        drafts = all_drafts()
        if not drafts:
            print("No Draft_*.docx files found.")
            sys.exit(0)
        total = len(drafts)
        passed_count = 0
        for draft in sorted(drafts):
            ok, _ = validate(draft)
            if ok: passed_count += 1
        print(f"\n{'─'*60}")
        print(f"Results: {passed_count}/{total} passed")
        sys.exit(0 if passed_count == total else 1)

    elif args.path:
        ok, _ = validate(args.path)
        sys.exit(0 if ok else 1)

    else:
        parser.print_help()
        sys.exit(1)

if __name__ == "__main__":
    main()
