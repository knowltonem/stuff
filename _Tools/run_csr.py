#!/usr/bin/env python3
"""
run_csr.py — Parameterized Cybersecurity Summary Report (CSR) runner.

Generates CSR DTR001 from INITIAL source document, or DTR002+ from previous draft.
Invoked by the newd prompt sequence after Product Category, CTN, and Document Type
have been selected.

Usage (direct):
    python3 _Tools/run_csr.py

Or triggered automatically by newd after skill_base.md routes to this runner.
"""

import copy
import re
import sys
from datetime import datetime
from pathlib import Path

# runner_core.py lives in the same _Tools/ directory
sys.path.insert(0, str(Path(__file__).parent))
from runner_core import (
    qn, clone_element, get_para_text, _set_cell_text,
    apply_keep_next, apply_acronym_page_break, strip_acronym_cell_borders,
    csr_dir, csr_draft_path, csr_initial_path, csr_example_path,
    run_validate, append_draft_log, BASE, get_git_username,
    merge_tables,
)

from lxml import etree
from docx import Document


# ---------------------------------------------------------------------------
# Well-known port types lookup (TCP/UDP)
# ---------------------------------------------------------------------------
WELL_KNOWN_PORTS = {
    # TCP ports
    "22": "TCP",      # SSH
    "80": "TCP",      # HTTP
    "443": "TCP",     # HTTPS
    "5060": "TCP",    # SIP (unencrypted)
    "5061": "TCP",    # SIP over TLS
    "8443": "TCP",    # HTTPS alt
    "389": "TCP",     # LDAP
    "636": "TCP",     # LDAPS
    "3389": "TCP",    # RDP
    # UDP ports
    "123": "UDP",     # NTP
    "161": "UDP",     # SNMP
    "162": "UDP",     # SNMP Trap
    "500": "UDP",     # ISAKMP/IKE (IPsec)
    "514": "UDP",     # Syslog
    "4500": "UDP",    # IPsec NAT-T
    "5060": "UDP",    # SIP (can be UDP too)
    "69": "UDP",      # TFTP
}


def get_port_type(port: str, default: str = "TCP") -> str:
    """Look up port type (TCP/UDP) from well-known ports table."""
    return WELL_KNOWN_PORTS.get(str(port), default)


# ---------------------------------------------------------------------------
# Helper: set tblHeader on a row
# ---------------------------------------------------------------------------

def _set_tbl_header(row_el):
    """Set tblHeader property on a table row element."""
    trpr = row_el.find(qn("w:trPr"))
    if trpr is None:
        trpr = etree.SubElement(row_el, qn("w:trPr"))
        row_el.insert(0, trpr)
    if trpr.find(qn("w:tblHeader")) is None:
        etree.SubElement(trpr, qn("w:tblHeader"))


# ---------------------------------------------------------------------------
# Helper: remove a table from the document body + surrounding title/spacer
# ---------------------------------------------------------------------------

def _remove_table_and_title(body, tbl_el, title_text_prefix=None):
    """Remove a table element from body. Optionally also remove the title paragraph above it."""
    if title_text_prefix:
        prev = tbl_el.getprevious()
        while prev is not None:
            if prev.tag == qn("w:p"):
                txt = get_para_text(prev).strip()
                if txt.startswith(title_text_prefix):
                    body.remove(prev)
                    break
                elif txt == "":
                    # skip empty spacers above
                    prev = prev.getprevious()
                    continue
                else:
                    break
            else:
                break
    body.remove(tbl_el)


# ---------------------------------------------------------------------------
# Helper: remove a column from a table
# ---------------------------------------------------------------------------

def _remove_column(tbl_el, col_idx):
    """Remove a column (by index) from all rows in a table element."""
    for row_el in tbl_el.findall(qn("w:tr")):
        cells = row_el.findall(qn("w:tc"))
        if col_idx < len(cells):
            row_el.remove(cells[col_idx])


# ---------------------------------------------------------------------------
# Helper: add a column to a table (clone last cell of each row)
# ---------------------------------------------------------------------------

def _add_column(tbl_el, header_text=""):
    """Add a column to the end of each row, cloning the last cell. Set header text on row 0."""
    rows = tbl_el.findall(qn("w:tr"))
    for i, row_el in enumerate(rows):
        cells = row_el.findall(qn("w:tc"))
        if cells:
            new_cell = clone_element(cells[-1])
            # Clear text in new cell
            for t in new_cell.findall(f".//{qn('w:t')}"):
                t.text = ""
            if i == 0 and header_text:
                # Set header text
                t_els = new_cell.findall(f".//{qn('w:t')}")
                if t_els:
                    t_els[0].text = header_text
                    t_els[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            row_el.append(new_cell)


# ---------------------------------------------------------------------------
# Helper: parse CAT counts from a W/O-RAE or W-RAE cell text
# ---------------------------------------------------------------------------

def _parse_cat_counts(text):
    """Parse CAT I/II/III counts from cell text like '12 CAT I\\n44 CAT II\\nNo CAT III'.
    Returns dict {'I': int, 'II': int, 'III': int}."""
    cats = {'I': 0, 'II': 0, 'III': 0}
    for line in text.split('\n'):
        line = line.strip()
        m = re.match(r'(\d+)\s+CAT\s+(I{1,3})', line)
        if m:
            cats[m.group(2)] = int(m.group(1))
        # 'No CAT X' = 0, already default
    return cats


def _format_cat_counts(cats):
    """Format CAT counts dict back to cell text: '6 CAT I\\n29 CAT II\\nNo CAT III'."""
    lines = []
    for level in ('I', 'II', 'III'):
        val = cats.get(level, 0)
        if val == 0:
            lines.append(f"No CAT {level}")
        else:
            lines.append(f"{val} CAT {level}")
    return "\n".join(lines)


def _set_cell_cat_lines(cell_el, cats_dict):
    """Write CAT I/II/III counts as separate paragraphs in a cell element.
    Each CAT level gets its own paragraph so they stack vertically in Word.
    Clones formatting from the first existing paragraph in the cell."""
    lines = []
    for level in ('I', 'II', 'III'):
        val = cats_dict.get(level, 0)
        suffix = " " if level != "III" else ""  # trailing space on I/II per Example
        if val == 0:
            lines.append(f"No CAT {level}{suffix}")
        else:
            lines.append(f"{val} CAT {level}{suffix}")

    paras = cell_el.findall(qn("w:p"))
    # Use first paragraph as template for formatting
    template_p = paras[0] if paras else None

    # Remove all existing paragraphs
    for p in paras:
        cell_el.remove(p)

    # Add one paragraph per CAT line
    for line_text in lines:
        if template_p is not None:
            new_p = copy.deepcopy(template_p)
        else:
            new_p = etree.SubElement(cell_el, qn("w:p"))
        # Clear existing runs and set new text
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # Create a single run with the text
        run_el = etree.SubElement(new_p, qn("w:r"))
        # Copy rPr from template if available
        if template_p is not None:
            tmpl_runs = template_p.findall(qn("w:r"))
            if tmpl_runs:
                tmpl_rPr = tmpl_runs[0].find(qn("w:rPr"))
                if tmpl_rPr is not None:
                    run_el.insert(0, copy.deepcopy(tmpl_rPr))
        t_el = etree.SubElement(run_el, qn("w:t"))
        t_el.text = line_text
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        cell_el.append(new_p)


# ---------------------------------------------------------------------------
# Helper: recalculate Table 1a from Table 5 + update CoF paragraph count
# ---------------------------------------------------------------------------

def _recalculate_table1a_and_cof(doc):
    """
    Calculate Table 1a (CS Test Summary) values from Table 5 (STIG CoF):
    1. Each STIG group in Table 5 has one W/O-RAE and W-RAE value (same for all rows in group)
    2. Sum unique group values across all groups = Total findings
    3. Divide by number of components (from Table 1a data rows) = per-component value
    4. CoF count = Total CAT I + Total CAT II + Total CAT III (from W/O-RAE)
    """
    # Find Table 5 (STIG CoF tables) — look for "Requirement" + "CoF" in header
    stig_tables = []
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Requirement" in row0_text and ("CoF" in row0_text or "W/O-RAE" in row0_text):
            stig_tables.append(tbl)

    if not stig_tables:
        print("  Step 11: [WARN] Table 5 (STIG CoF) not found — skipped")
        return

    # Extract unique W/O-RAE and W-RAE per requirement group (col 0 text)
    # Each group has the same values repeated across rows — take first occurrence
    groups_wo_rae = {}  # group_name -> cats dict
    groups_w_rae = {}
    for tbl in stig_tables:
        for i, row in enumerate(tbl.rows):
            if i == 0:
                continue  # skip header
            cells = row.cells
            req_text = cells[0].text.strip()
            if req_text.upper().startswith("NOTE"):
                continue
            # Use first line of requirement as group key (strip "(Continued)" suffix)
            group_key = req_text.split('\n')[0].strip()
            group_key = re.sub(r'\s*\(Continued\)\s*$', '', group_key)
            if group_key not in groups_wo_rae:
                # W/O-RAE is col index -2, W-RAE is col index -1
                wo_rae_text = cells[-2].text.strip()
                w_rae_text = cells[-1].text.strip()
                groups_wo_rae[group_key] = _parse_cat_counts(wo_rae_text)
                groups_w_rae[group_key] = _parse_cat_counts(w_rae_text)

    # Sum across all groups
    total_wo_rae = {'I': 0, 'II': 0, 'III': 0}
    total_w_rae = {'I': 0, 'II': 0, 'III': 0}
    for cats in groups_wo_rae.values():
        for k in cats:
            total_wo_rae[k] += cats[k]
    for cats in groups_w_rae.values():
        for k in cats:
            total_w_rae[k] += cats[k]

    # Find Table 1a (CS Test Summary) — header has "Component" + "W/O-RAE"
    table1a = None
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Component" in row0_text and "W/O-RAE" in row0_text:
            table1a = tbl
            break

    if table1a is None:
        print("  Step 11: [WARN] Table 1a (CS Test Summary) not found — skipped")
        return

    # Count component rows (exclude header row 0 and Total row)
    tbl1a_rows = table1a.rows
    num_components = len(tbl1a_rows) - 2  # minus header and Total row
    if num_components < 1:
        num_components = 1

    # Per-component values = total / num_components (rounded; warn if uneven)
    per_comp_wo_rae = {k: round(total_wo_rae[k] / num_components) for k in total_wo_rae}
    per_comp_w_rae = {k: round(total_w_rae[k] / num_components) for k in total_w_rae}
    for k in total_wo_rae:
        if per_comp_wo_rae[k] * num_components != total_wo_rae[k]:
            print(f"  [WARN] W/O-RAE {k} total {total_wo_rae[k]} not evenly divisible by {num_components} components")

    # Update Table 1a
    tbl1a_el = table1a._tbl
    tbl1a_rows_el = tbl1a_el.findall(qn("w:tr"))
    for row_idx in range(1, len(tbl1a_rows_el)):
        row_el = tbl1a_rows_el[row_idx]
        cells = row_el.findall(qn("w:tc"))
        if len(cells) < 4:
            continue
        cell0_text = "".join(
            t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Total" in cell0_text:
            # Total Findings row — use full totals
            _set_cell_cat_lines(cells[2], total_wo_rae)
            _set_cell_cat_lines(cells[3], total_w_rae)
        else:
            # Component row — use per-component values
            _set_cell_cat_lines(cells[2], per_comp_wo_rae)
            _set_cell_cat_lines(cells[3], per_comp_w_rae)

    # Update CoF paragraph count
    cof_total = total_wo_rae['I'] + total_wo_rae['II'] + total_wo_rae['III']
    for p in doc.paragraphs:
        if "findings shown below" in p.text:
            for r in p.runs:
                if re.search(r'\d+\s+findings', r.text):
                    r.text = re.sub(r'\d+(\s+findings)', f'{cof_total}\\1', r.text)
                    break
            break

    print(f"  Step 11: Table 1a recalculated (Total: {cof_total} findings W/O-RAE)")


# ---------------------------------------------------------------------------
# Step A: DTR001 — INITIAL source generation
# ---------------------------------------------------------------------------

def generate_dtr001(cfg: dict):
    """Generate CSR DTR001 from INITIAL source document."""
    prod_cat = cfg["prod_cat"]
    ctn = cfg["ctn"]
    dtr_num = cfg["dtr_num"]  # should be 1
    new_ver = cfg["new_ver"]
    out_path = cfg["out_path"]

    initial_path = csr_initial_path(prod_cat, ctn)
    example_path = csr_example_path(prod_cat, ctn)

    if not initial_path.exists():
        raise FileNotFoundError(f"INITIAL doc not found: {initial_path}")
    if not example_path.exists():
        raise FileNotFoundError(f"Example doc not found: {example_path}")

    doc = Document(str(initial_path))
    example = Document(str(example_path))
    body = doc.element.body
    ver_num = new_ver.replace("IOS XE ", "")

    # -----------------------------------------------------------------------
    # 1. REVISION HISTORY — add new row to Table 0
    # -----------------------------------------------------------------------
    rev_tbl = doc.tables[0]
    rev_tbl_el = rev_tbl._tbl
    rev_rows = rev_tbl_el.findall(qn("w:tr"))
    # Clone last data row as template
    new_row = clone_element(rev_rows[-1])
    new_cells = new_row.findall(qn("w:tc"))
    # Version: increment from last row — read first cell text directly
    first_cell_el = rev_rows[-1].findall(qn("w:tc"))[0] if rev_rows[-1].findall(qn("w:tc")) else None
    last_version_text = "".join(
        t.text or "" for t in (first_cell_el.findall(f".//{qn('w:t')}") if first_cell_el is not None else [])
    ).strip()
    try:
        last_ver = float(last_version_text)
        new_version_num = f"{last_ver + 1.0:.1f}"
    except ValueError:
        new_version_num = "2.0"
    date_str = datetime.now().strftime("%B %Y")  # CSR uses "Month YYYY" (e.g. May 2026)
    change_desc = f"Update for DTR {dtr_num}"
    values = [new_version_num, date_str, change_desc]
    for cell_el, val in zip(new_cells, values):
        _set_cell_text(cell_el, val)
    rev_tbl_el.append(new_row)
    print("  Step 1: Revision history row added")

    # -----------------------------------------------------------------------
    # 2. UPDATE SYSTEM TITLE — replace IOS XE version
    # -----------------------------------------------------------------------
    for p in doc.paragraphs:
        if "SYSTEM TITLE" in p.text and "IOS XE" in p.text:
            for r in p.runs:
                if "IOS XE" in r.text:
                    r.text = re.sub(r"IOS XE\s+[\d.]+", f"IOS XE {ver_num}", r.text)
            break
    print("  Step 2: SYSTEM TITLE version updated")

    # -----------------------------------------------------------------------
    # 3. TABLE 1c (CMVP) — KEEP in DTR001+ (previously removed, now retained)
    # -----------------------------------------------------------------------
    print("  Step 3: Table 1c (CMVP) retained")

    # -----------------------------------------------------------------------
    # 4. ADD CONCURRENCE COLUMN to Desktop Review Table
    # -----------------------------------------------------------------------
    # Find Desktop Review table by header content
    desktop_review_tbl = None
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Request Date" in row0_text and "DTR Number" in row0_text:
            desktop_review_tbl = tbl
            break
    if desktop_review_tbl is None:
        raise RuntimeError("Desktop Review Table not found")
    _add_column(desktop_review_tbl._tbl, header_text="Concurrence")
    print("  Step 4: Concurrence column added to Desktop Review Table")

    # -----------------------------------------------------------------------
    # 5. ADD DTR ROW to Desktop Review Table
    # -----------------------------------------------------------------------
    dr_tbl_el = desktop_review_tbl._tbl
    dr_rows = dr_tbl_el.findall(qn("w:tr"))
    # Clone header row as template for structure, then set values
    new_dtr_row = clone_element(dr_rows[-1])
    dtr_cells = new_dtr_row.findall(qn("w:tc"))
    # Build description text (DTR statement + optional similarity + optional POA&M)
    description = cfg["dtr_description"]
    if cfg.get("similarity_text"):
        description += f"\n\n{cfg['similarity_text']}"
    if cfg.get("poam_text"):
        description += f"\n\n{cfg['poam_text']}"
    dtr_row_values = [
        cfg["request_date"],
        f"DTR {dtr_num}",
        description,
        cfg["results_of_testing"],
        cfg.get("concurrence", ""),
    ]
    for cell_el, val in zip(dtr_cells, dtr_row_values):
        _set_cell_text(cell_el, val)
    dr_tbl_el.append(new_dtr_row)
    print("  Step 5: DTR row added to Desktop Review Table")

    # -----------------------------------------------------------------------
    # 6. MERGE TABLE 4 (SUT HW/SW) — INITIAL has it split across 2 tables
    #    Find the two tables by looking for "System Name" header
    # -----------------------------------------------------------------------
    hwsw_tables = []
    for i, tbl in enumerate(doc.tables):
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "System Name" in row0_text or "Product Component" in row0_text:
            hwsw_tables.append(i)
    if len(hwsw_tables) >= 2:
        merge_tables(doc, body, hwsw_tables[0], hwsw_tables[1], skip_header_rows=1)
        print("  Step 6: Table 4 (HW/SW) merged from 2 tables into 1")
    else:
        print("  Step 6: Table 4 already single table — skipped merge")

    # Split merged Table 4 at "Product Component" row into two tables.
    # Word only repeats consecutive tblHeader rows from row 0, so a
    # non-contiguous header row won't repeat on page breaks. Splitting
    # gives Product Component its own table with tblHeader on row 0.
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "System Name" in row0_text:
            tbl_el = tbl._tbl
            rows = tbl_el.findall(qn("w:tr"))
            split_idx = None
            for ri, row_el in enumerate(rows):
                cell0 = "".join(
                    t.text or "" for t in row_el.findall(qn("w:tc"))[0].findall(f".//{qn('w:t')}")
                ).strip()
                if cell0.startswith("Product Component"):
                    split_idx = ri
                    break
            if split_idx is not None:
                new_tbl = copy.deepcopy(tbl_el)
                # Remove rows before split_idx from new table
                for r in new_tbl.findall(qn("w:tr"))[:split_idx]:
                    new_tbl.remove(r)
                # Remove rows from split_idx onward from original
                for r in rows[split_idx:]:
                    tbl_el.remove(r)
                # Set tblHeader on row 0 of new table
                _set_tbl_header(new_tbl.findall(qn("w:tr"))[0])
                # Add tblLayout fixed
                new_tblPr = new_tbl.find(qn("w:tblPr"))
                if new_tblPr is not None and new_tblPr.find(qn("w:tblLayout")) is None:
                    layout = etree.SubElement(new_tblPr, qn("w:tblLayout"))
                    layout.set(qn("w:type"), "fixed")
                tbl_el.addnext(new_tbl)
                # Insert minimal spacer paragraph between the two tables.
                # Without this, Word treats adjacent tables as one table
                # and repeats the System Name header instead of Product Component.
                # Build the element detached (not via SubElement on body) then insert
                # directly before new_tbl to avoid double-parenting side effects.
                spacer = etree.Element(qn("w:p"))
                pPr = etree.SubElement(spacer, qn("w:pPr"))
                sp = etree.SubElement(pPr, qn("w:spacing"))
                sp.set(qn("w:before"), "0")
                sp.set(qn("w:after"), "0")
                sp.set(qn("w:line"), "1")
                sp.set(qn("w:lineRule"), "exact")
                rPr = etree.SubElement(pPr, qn("w:rPr"))
                sz = etree.SubElement(rPr, qn("w:sz"))
                sz.set(qn("w:val"), "1")
                etree.SubElement(rPr, qn("w:vanish"))
                new_tbl.addprevious(spacer)
                # Fix boundary: set bottom border of System Name last row
                # to single sz=4 so the line is visible even across page breaks.
                last_row = tbl_el.findall(qn("w:tr"))[-1]
                for cell in last_row.findall(qn("w:tc")):
                    tcPr = cell.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = etree.SubElement(cell, qn("w:tcPr"))
                        cell.insert(0, tcPr)
                    tcBorders = tcPr.find(qn("w:tcBorders"))
                    if tcBorders is None:
                        tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
                    bottom = tcBorders.find(qn("w:bottom"))
                    if bottom is None:
                        bottom = etree.SubElement(tcBorders, qn("w:bottom"))
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "4")
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "auto")
                print("  Step 6a: Table 4 split — Product Component table with tblHeader")
            break

    # -----------------------------------------------------------------------
    # 7. UPDATE TABLE 4 — component versions
    # -----------------------------------------------------------------------
    if cfg.get("updating_components"):
        # Find both Table 4 halves (System Name + Product Component) after split
        hwsw_tables = []
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "System Name" in row0_text or "Product Component" in row0_text:
                hwsw_tables.append(tbl)
        for hwsw_tbl in hwsw_tables:
            for row in hwsw_tbl.rows[1:]:  # skip header
                cells = row.cells
                if len(cells) < 2:
                    continue
                # Check both Product Component (col 0) and Hardware (col 1) for match
                row_text = " ".join(c.text.strip() for c in cells[:2])
                for comp_cfg in cfg["updating_components"]:
                    if comp_cfg["name"] in row_text:
                        # Update version in relevant cells (merged cols 1-4 typically)
                        for cell in cells[1:]:
                            cell_text = cell.text.strip()
                            if re.search(r"IOS XE\s+[\d.]+", cell_text):
                                for para in cell.paragraphs:
                                    for r in para.runs:
                                        if "IOS XE" in r.text:
                                            r.text = re.sub(
                                                r"IOS XE\s+[\d.]+",
                                                f"IOS XE {ver_num}",
                                                r.text
                                            )
        print("  Step 7: Table 4 component versions updated")
    else:
        print("  Step 7: No component updates — skipped")

    # -----------------------------------------------------------------------
    # 7b. ADD NEW PLATFORMS to Table 4 (Product Component table)
    #     cfg["new_platforms"] = list of dicts:
    #       { "hardware": str, "card_name": str, "version": str,
    #         "similarity": str, "components": ["SBC", "IWBC", "IWG"] }
    #     Rows are inserted after the last existing row for each component.
    # -----------------------------------------------------------------------
    new_platforms = cfg.get("new_platforms", [])
    if new_platforms:
        # Find the Product Component table
        prod_comp_tbl = None
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Product Component" in row0_text:
                prod_comp_tbl = tbl
                break

        if prod_comp_tbl is None:
            print("  Step 7b: [WARN] Product Component table not found — new platforms skipped")
        else:
            tbl_el = prod_comp_tbl._tbl
            rows = tbl_el.findall(qn("w:tr"))

            for platform in new_platforms:
                hw = platform.get("hardware", "")
                card = platform.get("card_name", "NA")
                plat_ver = platform.get("version") or f"IOS XE {ver_num}"
                sim = platform.get("similarity", "NA")
                components = platform.get("components", ["SBC", "IWBC", "IWG"])

                for comp_name in components:
                    # Find the last row whose col 0 matches this component
                    last_match_idx = None
                    for idx, row_el in enumerate(rows):
                        cells = row_el.findall(qn("w:tc"))
                        if cells:
                            cell0_text = "".join(
                                t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")
                            ).strip()
                            if cell0_text == comp_name:
                                last_match_idx = idx

                    # Clone a data row from the same component group as template
                    template_row_el = None
                    for row_el in rows:
                        cells = row_el.findall(qn("w:tc"))
                        if cells:
                            cell0_text = "".join(
                                t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")
                            ).strip()
                            if cell0_text == comp_name:
                                template_row_el = row_el
                                break

                    if template_row_el is None:
                        print(f"  Step 7b: [WARN] No existing '{comp_name}' rows — cannot add platform '{hw}'")
                        continue

                    new_row_el = clone_element(template_row_el)
                    new_cells = new_row_el.findall(qn("w:tc"))

                    # Set cell values: col0=comp_name, col1=hw, col2=card, col3=version, col4=sim
                    values = [comp_name, hw, card, plat_ver, sim]
                    for ci, val in enumerate(values):
                        if ci < len(new_cells):
                            for t_el in new_cells[ci].findall(f".//{qn('w:t')}"):
                                t_el.text = val
                                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

                    # Insert after the last matching row
                    insert_after = rows[last_match_idx] if last_match_idx is not None else rows[-1]
                    insert_after.addnext(new_row_el)
                    # Refresh rows list after insertion
                    rows = tbl_el.findall(qn("w:tr"))

            added = sum(len(p.get("components", ["SBC", "IWBC", "IWG"])) for p in new_platforms)
            print(f"  Step 7b: {len(new_platforms)} new platform(s) added ({added} rows total)")
    else:
        print("  Step 7b: No new platforms — skipped")

    # -----------------------------------------------------------------------
    # 8. REMOVE CAT COLUMN from Table 5 (STIG CoF)
    #    INITIAL Table 5 has 6 cols: Requirement | CoF SRGs | CAT | Component | W/O-RAE | W-RAE
    #    DTR001+ has 5 cols: Requirement | CoF SRGs | Component | W/O-RAE | W-RAE
    #    CAT is column index 2
    # -----------------------------------------------------------------------
    stig_tables = []
    for i, tbl in enumerate(doc.tables):
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Requirement" in row0_text and "CoF" in row0_text:
            stig_tables.append(tbl)
    for stig_tbl in stig_tables:
        tbl_el = stig_tbl._tbl
        # Verify it has 6 columns (has CAT column)
        first_row = tbl_el.findall(qn("w:tr"))[0]
        if len(first_row.findall(qn("w:tc"))) == 6:
            _remove_column(tbl_el, 2)  # Remove CAT (index 2)
            # Fix tblGrid — remove the CAT gridCol (index 2) and rebuild
            old_grid = tbl_el.find(qn("w:tblGrid"))
            if old_grid is not None:
                cols = old_grid.findall(qn("w:gridCol"))
                if len(cols) == 6:
                    # Get widths excluding CAT col, merge CAT width into col 1
                    widths = [c.get(qn("w:w")) for c in cols]
                    # Remove index 2 (CAT), add its width to index 1 (CoF SRGs)
                    cat_w = int(widths[2])
                    new_widths = widths[:2] + widths[3:]
                    new_widths[1] = str(int(new_widths[1]) + cat_w)
                    new_grid = etree.SubElement(tbl_el, qn("w:tblGrid"))
                    for w in new_widths:
                        col = etree.SubElement(new_grid, qn("w:gridCol"))
                        col.set(qn("w:w"), w)
                    tbl_el.replace(old_grid, new_grid)
            # Fix cell widths to match new grid
            new_grid_el = tbl_el.find(qn("w:tblGrid"))
            target_w = [c.get(qn("w:w")) for c in new_grid_el.findall(qn("w:gridCol"))]
            for row_el in tbl_el.findall(qn("w:tr")):
                cells = row_el.findall(qn("w:tc"))
                if len(cells) == 1:
                    # Merged row (NOTES) — fix width + gridSpan
                    tcPr = cells[0].find(qn("w:tcPr"))
                    if tcPr is not None:
                        tcW = tcPr.find(qn("w:tcW"))
                        if tcW is not None:
                            tcW.set(qn("w:w"), str(sum(int(w) for w in target_w)))
                        gs = tcPr.find(qn("w:gridSpan"))
                        if gs is not None:
                            gs.set(qn("w:val"), str(len(target_w)))
                    continue
                for ci, cell in enumerate(cells):
                    if ci >= len(target_w):
                        break
                    tcPr = cell.find(qn("w:tcPr"))
                    if tcPr is not None:
                        tcW = tcPr.find(qn("w:tcW"))
                        if tcW is not None:
                            tcW.set(qn("w:w"), target_w[ci])
            # Set tblLayout fixed
            tblPr = tbl_el.find(qn("w:tblPr"))
            if tblPr is not None and tblPr.find(qn("w:tblLayout")) is None:
                layout_el = etree.SubElement(tblPr, qn("w:tblLayout"))
                layout_el.set(qn("w:type"), "fixed")
    print("  Step 8: CAT column removed from Table 5 (grid + cell widths fixed)")

    # Merge STIG CoF tables into one (INITIAL has them split across 2 tables)
    if len(stig_tables) == 2:
        tbl1_el = stig_tables[0]._tbl
        tbl2_el = stig_tables[1]._tbl
        # Append data rows from part 2 (skip header row 0)
        for row in tbl2_el.findall(qn("w:tr"))[1:]:
            tbl1_el.append(row)
        # Remove spacers between the two tables
        nxt = tbl1_el.getnext()
        while nxt is not None and nxt is not tbl2_el:
            to_del = nxt
            nxt = nxt.getnext()
            body.remove(to_del)
        body.remove(tbl2_el)
        print("  Step 8a: STIG CoF tables merged into one")

    # -----------------------------------------------------------------------
    # 8b. ADD 4th COLUMN to Internal Findings table (Known Issue #7)
    #     INITIAL has 3 grid cols; DTR001+ has 4.
    #     R0-R1 (IP/Connectivity): merged cell gridSpan 2→3
    #     R2-R4 (Test Tools rows): Functions cell gridSpan 1→2
    #     R5 (NOTE): gridSpan 3→4
    # -----------------------------------------------------------------------
    int_findings_tbl = None
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Internet Protocol Addresses" in row0_text:
            int_findings_tbl = tbl
            break
    if int_findings_tbl:
        tbl_el = int_findings_tbl._tbl
        # Add 4th gridCol
        grid = tbl_el.find(qn("w:tblGrid"))
        if grid is not None:
            grid_cols = grid.findall(qn("w:gridCol"))
            if len(grid_cols) == 3:
                # Split last col width in half for the new column
                last_w = int(grid_cols[-1].get(qn("w:w")))
                new_w = last_w // 2
                grid_cols[-1].set(qn("w:w"), str(last_w - new_w))
                new_col = etree.SubElement(grid, qn("w:gridCol"))
                new_col.set(qn("w:w"), str(new_w))
        # Update gridSpan on each row
        for row_el in tbl_el.findall(qn("w:tr")):
            cells = row_el.findall(qn("w:tc"))
            if len(cells) == 1:
                # NOTE row — span 3→4
                tcPr = cells[0].find(qn("w:tcPr"))
                if tcPr is not None:
                    gs = tcPr.find(qn("w:gridSpan"))
                    if gs is not None and gs.get(qn("w:val")) == "3":
                        gs.set(qn("w:val"), "4")
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        total_w = sum(int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol")))
                        tcW.set(qn("w:w"), str(total_w))
            elif len(cells) == 2:
                # IP Addresses / Connectivity — merged cell gridSpan 2→3
                tcPr = cells[1].find(qn("w:tcPr"))
                if tcPr is not None:
                    gs = tcPr.find(qn("w:gridSpan"))
                    if gs is not None and gs.get(qn("w:val")) == "2":
                        gs.set(qn("w:val"), "3")
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        # Width = sum of grid cols 1-3
                        gw = [int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))]
                        tcW.set(qn("w:w"), str(sum(gw[1:])))
            elif len(cells) == 3:
                # Test Tools rows — Functions cell (index 1) gridSpan 1→2
                tcPr = cells[1].find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = etree.SubElement(cells[1], qn("w:tcPr"))
                    cells[1].insert(0, tcPr)
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is None:
                    gs = etree.SubElement(tcPr, qn("w:gridSpan"))
                gs.set(qn("w:val"), "2")
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is not None:
                    gw = [int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))]
                    tcW.set(qn("w:w"), str(gw[1] + gw[2]))
        print("  Step 8b: Internal Findings expanded to 4 grid columns")
    else:
        print("  Step 8b: [WARN] Internal Findings table not found — skipped")

    # -----------------------------------------------------------------------
    # 9. RENAME IWBC → Combined SBC/IWG throughout
    # -----------------------------------------------------------------------
    # In paragraphs
    for p in doc.paragraphs:
        if p.text and "IWBC" in p.text:
            if "Combined SBC/IWG Combined (IWBC)" in p.text:
                # Special case: fix run-split pattern → "Combined SBC/IWG (IWBC)"
                # Run structure: "Combined SBC/" | "IWG Combined" | " (" | "IWBC)"
                for r in p.runs:
                    if r.text:
                        r.text = r.text.replace("IWG Combined", "IWG")
                        # Leave IWBC) as-is — already correct after removing "Combined"
            elif "Interworking Gateway (IWG) and Session Border Controller (SBC) Combined (IWBC)" in p.text:
                # Component 1 paragraph — IWBC is the abbreviation label; do not rename
                pass
            else:
                for r in p.runs:
                    if r.text and "IWBC" in r.text:
                        r.text = r.text.replace("IWBC", "Combined SBC/IWG")
    # In tables (except Acronym List)
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Component" in row0_text and "W/O-RAE" in row0_text and "Critical" in row0_text:
            # Table 1a (CS Test Summary) — strip parentheticals from component names
            for row in tbl.rows[1:]:
                cell = row.cells[0]
                for para in cell.paragraphs:
                    for r in para.runs:
                        if r.text:
                            r.text = r.text.replace("(Combined SBC/IWG)", "").replace("(Interworking Gateway)", "").strip()
                for para in cell.paragraphs:
                    if para.text.strip() == "":
                        p_el = para._element
                        p_el.getparent().remove(p_el)
            continue
        if "CMVP" in row0_text:
            # Table 1c only: IWBC → "Combined SBC/IWG (IWBC)" in Components column
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip() == "IWBC":
                        for para in cell.paragraphs:
                            for r in para.runs:
                                if r.text and r.text.strip() == "IWBC":
                                    r.text = r.text.replace("IWBC", "Combined SBC/IWG (IWBC)")
            continue
        # All other tables: leave IWBC as-is
    print("  Step 9: IWBC/rename rules applied (CMVP: Combined SBC/IWG (IWBC); elsewhere: kept as IWBC)")

    # -----------------------------------------------------------------------
    # 10. SET tblHeader ON ALL TABLES
    # -----------------------------------------------------------------------
    for tbl in doc.tables:
        tbl_el = tbl._tbl
        rows = tbl_el.findall(qn("w:tr"))
        if rows:
            _set_tbl_header(rows[0])
    print("  Step 10: tblHeader set on all tables")
    
    # -----------------------------------------------------------------------
    # 10a. SPECIAL: Internal Findings table — only "Test Tools | Functions | Results" repeats
    #      Word only repeats consecutive tblHeader rows from row 0, so we must
    #      split the table at row 2 to make that row the header of its own table.
    # -----------------------------------------------------------------------
    for tbl in doc.tables:
        tbl_el = tbl._tbl
        rows = tbl_el.findall(qn("w:tr"))
        if not rows:
            continue
        row0_text = "".join(
            t.text or "" for t in rows[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Internet Protocol Address" in row0_text and len(rows) >= 3:
            # Check row 2 is "Test Tools | Functions | Results"
            row2_text = "".join(
                t.text or "" for t in rows[2].findall(f".//{qn('w:t')}")
            ).strip()
            if "Test Tools" in row2_text:
                import copy as _copy
                # Split: rows 0-1 stay, rows 2+ become new table
                new_tbl = _copy.deepcopy(tbl_el)
                # Remove rows 0-1 from new table
                for r in new_tbl.findall(qn("w:tr"))[:2]:
                    new_tbl.remove(r)
                # Remove rows 2+ from original
                for r in rows[2:]:
                    tbl_el.remove(r)
                # Set tblHeader on row 0 of new table (Test Tools row)
                _set_tbl_header(new_tbl.findall(qn("w:tr"))[0])
                # Insert new table after original
                tbl_el.addnext(new_tbl)
                # Insert minimal spacer paragraph between tables.
                # Build detached (not via SubElement on body) then insert via addprevious
                # to avoid double-parenting side effects — matches the Table 4 split pattern.
                spacer = etree.Element(qn("w:p"))
                pPr = etree.SubElement(spacer, qn("w:pPr"))
                sp = etree.SubElement(pPr, qn("w:spacing"))
                sp.set(qn("w:before"), "0")
                sp.set(qn("w:after"), "0")
                sp.set(qn("w:line"), "1")
                sp.set(qn("w:lineRule"), "exact")
                rPr = etree.SubElement(pPr, qn("w:rPr"))
                sz = etree.SubElement(rPr, qn("w:sz"))
                sz.set(qn("w:val"), "1")
                etree.SubElement(rPr, qn("w:vanish"))
                new_tbl.addprevious(spacer)
                print("  Step 10a: Internal Findings split — Test Tools row repeats on page break")

    # -----------------------------------------------------------------------
    # 11. RECALCULATE Table 1a from Table 5 + update CoF paragraph count
    # -----------------------------------------------------------------------
    _recalculate_table1a_and_cof(doc)

    # -----------------------------------------------------------------------
    # 12. KEEP TABLE TITLES + APPENDIX PAGE BREAKS
    # -----------------------------------------------------------------------
    apply_keep_next(body)
    # Page break before APPENDIX A
    for child in body:
        if child.tag == qn("w:p"):
            txt = get_para_text(child).strip()
            if txt == "APPENDIX A":
                pPr = child.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(child, qn("w:pPr"))
                    child.insert(0, pPr)
                if pPr.find(qn("w:pageBreakBefore")) is None:
                    etree.SubElement(pPr, qn("w:pageBreakBefore"))
                break
    # Page break before APPENDIX B
    for child in body:
        if child.tag == qn("w:p"):
            txt = get_para_text(child).strip()
            if txt.startswith("APPENDIX B"):
                pPr = child.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(child, qn("w:pPr"))
                    child.insert(0, pPr)
                if pPr.find(qn("w:pageBreakBefore")) is None:
                    etree.SubElement(pPr, qn("w:pageBreakBefore"))
                break
    print("  Step 12: keepNext + page breaks applied")

    # -----------------------------------------------------------------------
    # 13. TABLE QUALITY — fix border weights, grids
    # -----------------------------------------------------------------------
    # NOTE: Acronym List table borders are preserved as-is from INITIAL
    # (outset sz=6 table borders + single sz=6 cell borders)
    # Align STIG CoF table grids (part 2 must match part 1)
    stig_tables = []
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Requirement" in row0_text and ("CoF" in row0_text or "W/O-RAE" in row0_text):
            stig_tables.append(tbl._tbl)
    if len(stig_tables) == 2:
        grid1 = stig_tables[0].find(qn("w:tblGrid"))
        grid2 = stig_tables[1].find(qn("w:tblGrid"))
        widths1 = [c.get(qn("w:w")) for c in grid1.findall(qn("w:gridCol"))]
        widths2 = [c.get(qn("w:w")) for c in grid2.findall(qn("w:gridCol"))]
        if widths1 != widths2:
            new_grid = etree.SubElement(stig_tables[1], qn("w:tblGrid"))
            for w in widths1:
                col = etree.SubElement(new_grid, qn("w:gridCol"))
                col.set(qn("w:w"), w)
            stig_tables[1].replace(grid2, new_grid)
    # Fix Desktop Review gridCol count to match actual columns
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Request Date" in row0_text:
            tbl_el = tbl._tbl
            grid = tbl_el.find(qn("w:tblGrid"))
            grid_cols = grid.findall(qn("w:gridCol"))
            actual_cols = len(tbl_el.findall(qn("w:tr"))[0].findall(qn("w:tc")))
            if len(grid_cols) < actual_cols:
                last_w = int(grid_cols[-1].get(qn("w:w")))
                new_w = last_w // 2  # split the last column evenly to accommodate the extra column
                grid_cols[-1].set(qn("w:w"), str(last_w - new_w))
                new_col = etree.SubElement(grid, qn("w:gridCol"))
                new_col.set(qn("w:w"), str(new_w))
            break
    # Center-align Product Component col 4 (Hardware certified by Similarity)
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if row0_text.startswith("Product Component"):
            for ri, row in enumerate(tbl.rows):
                if ri == 0:
                    continue
                tr_cells = row._tr.findall(qn("w:tc"))
                if len(tr_cells) < 5:
                    continue  # NOTE row
                cell = row.cells[4]
                for p in cell.paragraphs:
                    pPr = p._element.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = etree.SubElement(p._element, qn("w:pPr"))
                        p._element.insert(0, pPr)
                    jc = pPr.find(qn("w:jc"))
                    if jc is None:
                        jc = etree.SubElement(pPr, qn("w:jc"))
                    jc.set(qn("w:val"), "center")
            break
    
    # Center-align STIG CoF table cols 2-4 (Component Affected, W/O-RAE, W-RAE)
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Requirement" in row0_text and "W/O-RAE" in row0_text:
            for ri, row in enumerate(tbl.rows):
                if ri == 0:
                    continue  # Header already centered
                tr_cells = row._tr.findall(qn("w:tc"))
                if len(tr_cells) < 5:
                    continue  # Skip incomplete rows
                for ci in [2, 3, 4]:  # Component Affected, W/O-RAE, W-RAE
                    cell = row.cells[ci]
                    for p in cell.paragraphs:
                        pPr = p._element.find(qn("w:pPr"))
                        if pPr is None:
                            pPr = etree.SubElement(p._element, qn("w:pPr"))
                            p._element.insert(0, pPr)
                        jc = pPr.find(qn("w:jc"))
                        if jc is None:
                            jc = etree.SubElement(pPr, qn("w:jc"))
                        jc.set(qn("w:val"), "center")
            break
    print("  Step 13: Table quality fixes applied")

    # -----------------------------------------------------------------------
    # 14. UPDATE STIG FINDINGS (if changed)
    # -----------------------------------------------------------------------
    if cfg.get("stig_updates"):
        # cfg["stig_updates"] is a list of dicts:
        # [{"section": "VVoIP STIG", "cat_i": "None.", "cat_ii": "...", "cat_iii": "..."}]
        for update in cfg["stig_updates"]:
            section_name = update["section"]
            # Find the section paragraph and update CAT lines after it
            # Use flexible matching: check if key part of section name is in paragraph
            # e.g. "VVoIP" for "Voice/Video over Internet Protocol (VVoIP) STIG"
            section_key = section_name.replace(" STIG", "").replace(" SRG", "").strip()
            found_section = False
            cat_i_done = False
            cat_ii_done = False
            cat_iii_done = False
            for p in doc.paragraphs:
                if found_section:
                    txt = p.text.strip()
                    if txt.startswith("CAT I") and not cat_i_done:
                        # Keep "CAT I:" bold in run0, put value in run1 (non-bold)
                        p.runs[0].text = "CAT I:"
                        if len(p.runs) > 1:
                            p.runs[1].text = f"  {update['cat_i']}"
                            p.runs[1].bold = False
                            for r in p.runs[2:]:
                                r.text = ""
                        else:
                            # Add a new run for the value
                            new_run = p.add_run(f"  {update['cat_i']}")
                            new_run.bold = False
                        cat_i_done = True
                    elif txt.startswith("CAT II") and not cat_ii_done:
                        p.runs[0].text = "CAT II:"
                        if len(p.runs) > 1:
                            p.runs[1].text = f"  {update['cat_ii']}"
                            p.runs[1].bold = False
                            for r in p.runs[2:]:
                                r.text = ""
                        else:
                            new_run = p.add_run(f"  {update['cat_ii']}")
                            new_run.bold = False
                        cat_ii_done = True
                    elif txt.startswith("CAT III") and not cat_iii_done:
                        p.runs[0].text = "CAT III:"
                        if len(p.runs) > 1:
                            p.runs[1].text = f" {update['cat_iii']}"
                            p.runs[1].bold = False
                            for r in p.runs[2:]:
                                r.text = ""
                        else:
                            new_run = p.add_run(f" {update['cat_iii']}")
                            new_run.bold = False
                        cat_iii_done = True
                        break  # done with this section
                elif section_key in p.text and ("STIG" in p.text or "SRG" in p.text):
                    found_section = True
        print("  Step 14: STIG Phase I findings updated")
    else:
        print("  Step 14: No STIG changes — skipped")

    # -----------------------------------------------------------------------
    # 15. UPDATE IPV FINDINGS (if changed)
    # -----------------------------------------------------------------------
    if cfg.get("ipv_update"):
        ipv = cfg["ipv_update"]
        # Build findings text: "X High Risk" or "No High Risk", etc.
        high = ipv.get("high", 0)
        medium = ipv.get("medium", 0)
        low = ipv.get("low", 0)
        high_text = f"{high} High Risk" if high > 0 else "No High Risk"
        medium_text = f"{medium} Medium Risk" if medium > 0 else "No Medium Risk"
        low_text = f"{low} Low Risk" if low > 0 else "No Low Risk"
        
        # Update Table 1b date and findings
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Requirement" in row0_text and "Total Findings" in row0_text:
                # Table 1b — update date in row 1, col 1 and findings in row 1, col 2
                rows = tbl._tbl.findall(qn("w:tr"))
                if len(rows) > 1:
                    cells = rows[1].findall(qn("w:tc"))
                    if len(cells) > 1:
                        _set_cell_text(cells[1], ipv.get("date", ""))
                    if len(cells) > 2:
                        # Cell has 3 paragraphs (one per risk level) — update each
                        cell_el = cells[2]
                        paras = cell_el.findall(qn("w:p"))
                        risk_texts = [high_text, medium_text, low_text]
                        for pi, (p, txt) in enumerate(zip(paras, risk_texts)):
                            runs = p.findall(qn("w:r"))
                            if runs:
                                t_els = runs[0].findall(qn("w:t"))
                                if t_els:
                                    t_els[0].text = txt
                                # Clear extra runs
                                for r in runs[1:]:
                                    p.remove(r)
                        # Remove extra paragraphs if any
                        for p in paras[3:]:
                            cell_el.remove(p)
                break
        
        # Update Table 3 (IP Vulnerability Summary) — per-component and total
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Component" in row0_text and "Critical" in row0_text and "Plug-in" in row0_text:
                rows = tbl._tbl.findall(qn("w:tr"))
                # Per-component findings text (with "s" suffix: "Risks")
                high_text_s = f"{high} High Risks" if high > 0 else "No High Risks"
                medium_text_s = f"{medium} Medium Risks" if medium > 0 else "No Medium Risks"
                low_text_s = f"{low} Low Risks" if low > 0 else "No Low Risks"
                risk_texts_s = [high_text_s, medium_text_s, low_text_s]
                total_count = high + medium + low
                for ri, row in enumerate(rows[1:], start=1):
                    cells = row.findall(qn("w:tc"))
                    cell0_text = "".join(
                        t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")
                    ).strip()
                    if "Total" in cell0_text:
                        # Total row — update col 4 with count
                        if len(cells) > 4:
                            _set_cell_text(cells[4], f"{total_count} Findings")
                    else:
                        # Component row — update cols 3 and 4 (each has multiple paragraphs)
                        for col_idx in [3, 4]:
                            if len(cells) > col_idx:
                                cell_el = cells[col_idx]
                                paras = cell_el.findall(qn("w:p"))
                                for pi, (p, txt) in enumerate(zip(paras, risk_texts_s)):
                                    runs = p.findall(qn("w:r"))
                                    if runs:
                                        t_els = runs[0].findall(qn("w:t"))
                                        if t_els:
                                            t_els[0].text = txt
                                        for r in runs[1:]:
                                            p.remove(r)
                                for p in paras[3:]:
                                    cell_el.remove(p)
                break
        
        print(f"  Step 15: IPV findings updated ({high} High, {medium} Medium, {low} Low)")
    else:
        print("  Step 15: No IPV changes — skipped")

    # -----------------------------------------------------------------------
    # 16. OPEN PORTS (Table 8) — add/remove
    # -----------------------------------------------------------------------
    if cfg.get("ports_add") or cfg.get("ports_remove"):
        ports_tbl = None
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Components" in row0_text and "Open Ports" in row0_text:
                ports_tbl = tbl
                break
        if ports_tbl:
            tbl_el = ports_tbl._tbl
            # Remove ports
            if cfg.get("ports_remove"):
                rows = tbl_el.findall(qn("w:tr"))
                for row in rows[1:]:  # skip header
                    row_text = "".join(
                        t.text or "" for t in row.findall(f".//{qn('w:t')}")
                    ).strip()
                    for port_to_remove in cfg["ports_remove"]:
                        if port_to_remove in row_text:
                            tbl_el.remove(row)
                            break
            # Add ports
            if cfg.get("ports_add"):
                rows = tbl_el.findall(qn("w:tr"))
                # Find a data row to clone (not header, not NOTES)
                template_row = None
                for row in rows[2:]:  # skip header rows
                    row_text = "".join(
                        t.text or "" for t in row.findall(f".//{qn('w:t')}")
                    ).strip()
                    if "NOTES" not in row_text and "Cisco" in row_text:
                        template_row = row
                        break
                if template_row is None:
                    template_row = rows[2]  # fallback to first data row
                
                for port_entry in cfg["ports_add"]:
                    new_row = clone_element(template_row)
                    cells = new_row.findall(qn("w:tc"))
                    # port_type: use explicit value, or look up from well-known ports, or default to TCP
                    port_num = port_entry.get("port", "")
                    port_type = port_entry.get("port_type")
                    if not port_type:
                        port_type = get_port_type(port_num, default="TCP")
                    if not port_type.endswith(":"):
                        port_type += ":"
                    # CSR Table 8 columns (DTR001 path):
                    # col 0 = Component, col 1 = Protocol (e.g. "TCP:"), col 2 = Port,
                    # col 3 = FW Rule, col 4 = Purpose
                    port_values = [
                        port_entry.get("component", ""),
                        port_type,
                        port_num,
                        port_entry.get("fw_rule", ""),
                        port_entry.get("purpose", ""),
                    ]
                    for cell_el, val in zip(cells, port_values):
                        _set_cell_text(cell_el, val)
                    
                    # Find the last row for this component and insert after it
                    target_component = port_entry.get("component", "")
                    insert_after_row = None
                    current_rows = tbl_el.findall(qn("w:tr"))
                    for row in current_rows[2:]:  # skip header rows
                        row_text = "".join(
                            t.text or "" for t in row.findall(f".//{qn('w:t')}")
                        ).strip()
                        if "NOTES" in row_text:
                            break
                        # Check if this row matches the component
                        cell0_text = "".join(
                            t.text or "" for t in row.findall(qn("w:tc"))[0].findall(f".//{qn('w:t')}")
                        ).strip()
                        if target_component in cell0_text:
                            insert_after_row = row
                    
                    # Insert after the last matching component row
                    if insert_after_row is not None:
                        insert_after_row.addnext(new_row)
                    else:
                        # Fallback: insert before NOTES row
                        notes_row = None
                        for row in current_rows:
                            row_text = "".join(
                                t.text or "" for t in row.findall(f".//{qn('w:t')}")
                            ).strip()
                            if "NOTES" in row_text:
                                notes_row = row
                                break
                        if notes_row is not None:
                            notes_row.addprevious(new_row)
                        else:
                            tbl_el.append(new_row)
            print("  Step 16: Open Ports table updated")
        else:
            print("  Step 16: [WARN] Open Ports table not found — skipped")
    else:
        print("  Step 16: No port changes — skipped")

    # -----------------------------------------------------------------------
    # 17. SAVE
    # -----------------------------------------------------------------------
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")
    # Note: validation is handled by run() after this function returns — do not call run_validate() here


# ---------------------------------------------------------------------------
# Step B: DTR002+ — incremental generation from previous draft
# ---------------------------------------------------------------------------

def generate_dtr_incremental(cfg: dict):
    """Generate CSR DTR002+ from the previous draft."""
    source_path = cfg["source_path"]
    out_path = cfg["out_path"]
    dtr_num = cfg["dtr_num"]
    new_ver = cfg["new_ver"]
    ver_num = new_ver.replace("IOS XE ", "")

    if not source_path.exists():
        raise FileNotFoundError(f"Source draft not found: {source_path}")

    doc = Document(str(source_path))
    body = doc.element.body

    # -----------------------------------------------------------------------
    # 1. REVISION HISTORY — add new row
    # -----------------------------------------------------------------------
    rev_tbl = doc.tables[0]
    rev_tbl_el = rev_tbl._tbl
    rev_rows = rev_tbl_el.findall(qn("w:tr"))
    new_row = clone_element(rev_rows[-1])
    new_cells = new_row.findall(qn("w:tc"))
    # Increment version
    last_cells = rev_rows[-1].findall(qn("w:tc"))
    last_ver_text = "".join(
        t.text or "" for t in last_cells[0].findall(f".//{qn('w:t')}")
    ).strip()
    try:
        last_ver = float(last_ver_text)
        new_version_num = f"{last_ver + 1.0:.1f}"
    except ValueError:
        new_version_num = f"{dtr_num + 1}.0"
    date_str = datetime.now().strftime("%B %Y")  # CSR uses "Month YYYY"
    change_desc = f"Update for DTR {dtr_num}"
    values = [new_version_num, date_str, change_desc]
    for cell_el, val in zip(new_cells, values):
        _set_cell_text(cell_el, val)
    rev_tbl_el.append(new_row)
    print("  Step 1: Revision history row added")

    # -----------------------------------------------------------------------
    # 2. UPDATE SYSTEM TITLE
    # -----------------------------------------------------------------------
    for p in doc.paragraphs:
        if "SYSTEM TITLE" in p.text and "IOS XE" in p.text:
            for r in p.runs:
                if "IOS XE" in r.text:
                    r.text = re.sub(r"IOS XE\s+[\d.]+", f"IOS XE {ver_num}", r.text)
            break
    print("  Step 2: SYSTEM TITLE version updated")

    # -----------------------------------------------------------------------
    # 3. ADD DTR ROW to Desktop Review Table
    # -----------------------------------------------------------------------
    desktop_review_tbl = None
    for tbl in doc.tables:
        row0_text = "".join(
            t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
        ).strip()
        if "Request Date" in row0_text and "DTR Number" in row0_text:
            desktop_review_tbl = tbl
            break
    if desktop_review_tbl is None:
        raise RuntimeError("Desktop Review Table not found")
    dr_tbl_el = desktop_review_tbl._tbl
    dr_rows = dr_tbl_el.findall(qn("w:tr"))
    new_dtr_row = clone_element(dr_rows[-1])
    dtr_cells = new_dtr_row.findall(qn("w:tc"))
    description = cfg["dtr_description"]
    if cfg.get("similarity_text"):
        description += f"\n\n{cfg['similarity_text']}"
    if cfg.get("poam_text"):
        description += f"\n\n{cfg['poam_text']}"
    dtr_row_values = [
        cfg["request_date"],
        f"DTR {dtr_num}",
        description,
        cfg["results_of_testing"],
        cfg.get("concurrence", ""),
    ]
    for cell_el, val in zip(dtr_cells, dtr_row_values):
        _set_cell_text(cell_el, val)
    dr_tbl_el.append(new_dtr_row)
    print("  Step 3: DTR row added to Desktop Review Table")

    # -----------------------------------------------------------------------
    # 4. UPDATE TABLE 4 — component versions (if any)
    # -----------------------------------------------------------------------
    if cfg.get("updating_components"):
        hwsw_tbl = None
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "System Name" in row0_text:
                hwsw_tbl = tbl
                break
        if hwsw_tbl:
            for row in hwsw_tbl.rows[1:]:
                cells = row.cells
                if len(cells) < 2:
                    continue
                row_text = " ".join(c.text.strip() for c in cells[:2])
                for comp_cfg in cfg["updating_components"]:
                    if comp_cfg["name"] in row_text:
                        for cell in cells[1:]:
                            cell_text = cell.text.strip()
                            if re.search(r"IOS XE\s+[\d.]+", cell_text):
                                for para in cell.paragraphs:
                                    for r in para.runs:
                                        if "IOS XE" in r.text:
                                            r.text = re.sub(
                                                r"IOS XE\s+[\d.]+",
                                                f"IOS XE {ver_num}",
                                                r.text
                                            )
        print("  Step 4: Table 4 component versions updated")
    else:
        print("  Step 4: No component updates — skipped")

    # -----------------------------------------------------------------------
    # 4b. ADD NEW PLATFORMS to Table 4 (Product Component table)
    #     Same logic as generate_dtr001 Step 7b
    # -----------------------------------------------------------------------
    new_platforms = cfg.get("new_platforms", [])
    if new_platforms:
        prod_comp_tbl = None
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Product Component" in row0_text:
                prod_comp_tbl = tbl
                break
        if prod_comp_tbl is None:
            print("  Step 4b: [WARN] Product Component table not found — new platforms skipped")
        else:
            tbl_el = prod_comp_tbl._tbl
            rows = tbl_el.findall(qn("w:tr"))
            for platform in new_platforms:
                hw = platform.get("hardware", "")
                card = platform.get("card_name", "NA")
                plat_ver = platform.get("version") or f"IOS XE {ver_num}"
                sim = platform.get("similarity", "NA")
                components = platform.get("components", ["SBC", "IWBC", "IWG"])
                for comp_name in components:
                    last_match_idx = None
                    template_row_el = None
                    for idx, row_el in enumerate(rows):
                        cells = row_el.findall(qn("w:tc"))
                        if cells:
                            cell0_text = "".join(
                                t.text or "" for t in cells[0].findall(f".//{qn('w:t')}")
                            ).strip()
                            if cell0_text == comp_name:
                                last_match_idx = idx
                                if template_row_el is None:
                                    template_row_el = row_el
                    if template_row_el is None:
                        print(f"  Step 4b: [WARN] No '{comp_name}' rows found — skipping")
                        continue
                    new_row_el = clone_element(template_row_el)
                    new_cells = new_row_el.findall(qn("w:tc"))
                    for ci, val in enumerate([comp_name, hw, card, plat_ver, sim]):
                        if ci < len(new_cells):
                            for t_el in new_cells[ci].findall(f".//{qn('w:t')}"):
                                t_el.text = val
                                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    insert_after = rows[last_match_idx] if last_match_idx is not None else rows[-1]
                    insert_after.addnext(new_row_el)
                    rows = tbl_el.findall(qn("w:tr"))
            added = sum(len(p.get("components", ["SBC", "IWBC", "IWG"])) for p in new_platforms)
            print(f"  Step 4b: {len(new_platforms)} new platform(s) added ({added} rows total)")
    else:
        print("  Step 4b: No new platforms — skipped")

    # -----------------------------------------------------------------------
    # 5. UPDATE STIG FINDINGS (if changed)
    # -----------------------------------------------------------------------
    if cfg.get("stig_updates"):
        # cfg["stig_updates"] is a list of dicts:
        # [{"section": "Network Device Management SRG", "cat_i": "None.", "cat_ii": "None.", "cat_iii": "None."}]
        for update in cfg["stig_updates"]:
            section_name = update["section"]
            # Find the section paragraph and update CAT lines after it
            found_section = False
            for p in doc.paragraphs:
                if found_section:
                    txt = p.text.strip()
                    if txt.startswith("CAT I"):
                        p.runs[0].text = f"CAT I:  {update['cat_i']}"
                        for r in p.runs[1:]:
                            r.text = ""
                    elif txt.startswith("CAT II"):
                        p.runs[0].text = f"CAT II:  {update['cat_ii']}"
                        for r in p.runs[1:]:
                            r.text = ""
                    elif txt.startswith("CAT III"):
                        p.runs[0].text = f"CAT III: {update['cat_iii']}"
                        for r in p.runs[1:]:
                            r.text = ""
                        break  # done with this section
                elif section_name in p.text:
                    found_section = True
        print("  Step 5: STIG findings updated")
    else:
        print("  Step 5: No STIG changes — skipped")

    # -----------------------------------------------------------------------
    # 6. UPDATE IPV FINDINGS (if changed)
    # -----------------------------------------------------------------------
    if cfg.get("ipv_update"):
        # Update Table 1b date
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Requirement" in row0_text and "Total Findings" in row0_text:
                # Table 1b — update date in row 1, col 1
                rows = tbl._tbl.findall(qn("w:tr"))
                if len(rows) > 1:
                    cells = rows[1].findall(qn("w:tc"))
                    if len(cells) > 1:
                        _set_cell_text(cells[1], cfg["ipv_update"]["scan_date"])
                break
        print("  Step 6: IPV findings updated")
    else:
        print("  Step 6: No IPV changes — skipped")

    # -----------------------------------------------------------------------
    # 7. OPEN PORTS (Table 8) — add/remove
    # -----------------------------------------------------------------------
    if cfg.get("ports_add") or cfg.get("ports_remove"):
        ports_tbl = None
        for tbl in doc.tables:
            row0_text = "".join(
                t.text or "" for t in tbl._tbl.findall(qn("w:tr"))[0].findall(f".//{qn('w:t')}")
            ).strip()
            if "Components" in row0_text and "Open Ports" in row0_text:
                ports_tbl = tbl
                break
        if ports_tbl:
            tbl_el = ports_tbl._tbl
            # Remove ports
            if cfg.get("ports_remove"):
                rows = tbl_el.findall(qn("w:tr"))
                for row in rows[1:]:  # skip header
                    row_text = "".join(
                        t.text or "" for t in row.findall(f".//{qn('w:t')}")
                    ).strip()
                    for port_to_remove in cfg["ports_remove"]:
                        if port_to_remove in row_text:
                            tbl_el.remove(row)
                            break
            # Add ports
            if cfg.get("ports_add"):
                rows = tbl_el.findall(qn("w:tr"))
                template_row = rows[-1]  # clone last row as template
                for port_entry in cfg["ports_add"]:
                    new_row = clone_element(template_row)
                    cells = new_row.findall(qn("w:tc"))
                    # CSR Table 8 columns (DTR002+ incremental path):
                    # col 0 = Component, col 1 = Protocol (e.g. "TCP:"), col 2 = Port,
                    # col 3 = FW Rule, col 4 = Purpose
                    # TODO: Verify col layout against a real DTR002+ CSR once one is on disk.
                    # If identical to DTR001 path, replace col 1 value with get_port_type() lookup.
                    port_num = port_entry.get("port", "")
                    port_type_inc = port_entry.get("port_type")
                    if not port_type_inc:
                        port_type_inc = get_port_type(port_num, default="TCP")
                    if not port_type_inc.endswith(":"):
                        port_type_inc += ":"
                    port_values = [
                        port_entry.get("component", ""),
                        port_type_inc,
                        port_num,  # col 2: port number
                        port_entry.get("fw_rule", ""),
                        port_entry.get("purpose", ""),
                    ]
                    if len(cells) < len(port_values):
                        print(f"[WARNING] Port row has {len(cells)} cells but expected {len(port_values)} — skipping row for {port_entry.get('port', '?')}")
                        continue
                    for cell_el, val in zip(cells, port_values):
                        _set_cell_text(cell_el, val)
                    tbl_el.append(new_row)
        print("  Step 7: Open Ports table updated")
    else:
        print("  Step 7: No port changes — skipped")

    # -----------------------------------------------------------------------
    # 8. SAVE
    # -----------------------------------------------------------------------
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        resp = input(f"WARNING: {out_path.name} already exists. Overwrite? [y/N] ").strip().lower()
        if resp != "y":
            print("Aborted — file not overwritten.")
            return
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")
    # Note: validation is handled by run() after this function returns — do not call run_validate() here


# ---------------------------------------------------------------------------
# Entry point (for direct invocation / testing)
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Seed profiles — add DTR001 profile per CTN when onboarding
# ---------------------------------------------------------------------------

# Structure mirrors other runners:
# { dtr_num (int): { cfg fields } }
# DTR002+ are never pre-baked — fully prompted at runtime.
CSR_PROFILES: dict = {
    "SBC/CTN2026003": {
        1: {
            "new_ver": "IOS XE 26.0",
            "old_ver": "IOS XE 17.18",
            "request_date": "01 June 2026",
            "dtr_description": (
                "DTR 1 was requested to update the IOS XE software version from IOS XE 17.18 "
                "to IOS XE 26.0 for the product components IWG and SBC on the C8200 series, "
                "the C8300 series, and the C8000v of router platforms. "
                "The ASR 1006-X, ISR 4461, ESXi Server Host, and Management Workstation will "
                "be sustained on current software load."
            ),
            "results_of_testing": "Testing was successful.",
            "concurrence": "Tester: J. Misal",
            "similarity_text": 'Request certification through similarity based on "ESC TN: CTN2026001, DTR01."',
            "poam_text": "This DTR Clears POA&M/TDR Number: 26003-01, http2.",
            "updating_components": [
                {"name": "C8000v"},   # updating to IOS XE 26.0
                {"name": "C8200"},    # C8200-1N-4T → IOS XE 26.0
                {"name": "C8300"},    # C8300 series → IOS XE 26.0
            ],
            "new_platforms": [],     # list of {"hardware","card_name","version","similarity","components"}
            "stig_updates": [],
            "ipv_update": None,
            "ports_add": [],
            "ports_remove": [],
        }
    },
}


def _get_profiles(prod_cat: str, ctn: str) -> dict:
    return CSR_PROFILES.get(f"{prod_cat}/{ctn}", {})


def run(prod_cat: str, ctn: str, dtr_num: int) -> None:
    """Callable entry point — invoked by the newd prompt sequence or directly via main()."""
    print("\n=== Cybersecurity Summary Report (CSR) — Parameterized Runner ===\n")

    profiles = _get_profiles(prod_cat, ctn)

    if dtr_num in profiles:
        cfg = dict(profiles[dtr_num])
        print(f"Loaded profile for {prod_cat}/{ctn} DTR{dtr_num:03d}:")
        print(f"  Version: {cfg.get('old_ver', '?')} → {cfg.get('new_ver', '?')}")
        confirm = input("\nProceed with these settings? [y/N]: ").strip().lower()
        if confirm != "y":
            print("Aborted.")
            return
    else:
        print(f"No profile found for {prod_cat}/{ctn} DTR{dtr_num:03d}. Manual input required.")
        new_ver = input("New IOS XE version (e.g. IOS XE 26.2): ").strip()
        old_ver = input("Old IOS XE version (e.g. IOS XE 17.18): ").strip()
        request_date = input("Request date (e.g. 01 May 2026): ").strip()
        dtr_description = input("DTR description (brief summary of this review): ").strip()
        results_of_testing = input("Results of testing (e.g. Passed): ").strip()
        cfg = {
            "dtr_num": dtr_num,
            "new_ver": new_ver,
            "old_ver": old_ver,
            "request_date": request_date,
            "dtr_description": dtr_description,
            "results_of_testing": results_of_testing,
            "similarity_text": "",
            "poam_text": "",
            "concurrence": "",
            "updating_components": [],
            "new_platforms": [],
            "stig_updates": [],
            "ipv_update": None,
            "ports_add": [],
            "ports_remove": [],
        }

    cfg["prod_cat"] = prod_cat
    cfg["ctn"] = ctn
    cfg["out_path"] = csr_draft_path(prod_cat, ctn, dtr_num, cfg["new_ver"])

    if dtr_num == 1:
        generate_dtr001(cfg)
    else:
        cfg["source_path"] = csr_draft_path(prod_cat, ctn, dtr_num - 1, cfg["old_ver"])
        generate_dtr_incremental(cfg)

    print("\nRunning validate_doc.py...")
    run_validate(cfg["out_path"])

    engineer = input("\nEngineer username (for Draft Log): ").strip() or get_git_username() or "unknown"
    append_draft_log(
        engineer=engineer,
        action="Generated",
        ctn=ctn,
        doc_type="CSR",
        dtr=f"DTR{dtr_num:03d}",
        version=cfg["new_ver"],
        reason="Via run_csr.py parameterized runner",
    )

    print("\nDone.")


def main():
    """Interactive entry point — called by newd after Product Category / CTN / Doc Type selected."""
    print("\n=== Cybersecurity Summary Report (CSR) — Parameterized Runner ===\n")
    prod_cat = input("Product Category (e.g. SBC): ").strip()
    ctn = input("CTN (e.g. CTN2026003): ").strip()
    dtr_num = int(input("DTR Number (e.g. 1): ").strip())
    run(prod_cat, ctn, dtr_num)


if __name__ == "__main__":
    main()
